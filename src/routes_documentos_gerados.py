import unicodedata

from flask import current_app
import io
import math
from zoneinfo import ZoneInfo
from flask_wtf.csrf import validate_csrf
from flask_login import login_required
from flask import abort, json, request, jsonify, make_response, current_app
from random import shuffle
import os
import zipfile
import qrcode
import pytz
import pandas as pd
import base64
import matplotlib.pyplot as plt
import requests
import urllib
from src.decorators.utils_acumulo import b2_bucket_name, b2_client, b2_delete_all_versions, b2_upload_fileobj
from src.identificacao import buscar_pessoa_por_cpf, normaliza_matricula
from src.formatar_cpf import cadete_restantes, formatar_cpf, get_militar_por_user, is_cadete
from flask import render_template, redirect, url_for, request, flash, jsonify, session, send_file, make_response, \
    Response, stream_with_context
from flask_login import login_user, logout_user, login_required, current_user
from flask_wtf.csrf import validate_csrf, generate_csrf
from werkzeug.utils import secure_filename
from src import app, database, bcrypt
from src.forms import (AtualizacaoCadastralForm, ControleConvocacaoForm, CriarSenhaForm, FichaAlunosForm, FormEsqueciSenha, FormFiltroMilitar, FormMilitarInativo, FormResetarSenhaPublica, FormViatura,
                       IdentificacaoForm, ImpactoForm, FormLogin, FormMilitar, FormCriarUsuario, FormMotoristas, FormFiltroMotorista, LtsAlunoForm, RecompensaAlunoForm,
                       RestricaoAlunoForm, SancaoAlunoForm, TabelaVencimentoForm, InativarAlunoForm, MatriculaConfirmForm)
from src.models import (ControleConvocacao, Convocacao, DocumentoMilitar, EfetivoDiarioOBM, HistoricoEfetivoDiario, ImportacaoMilitarHistorico, LogAcesso, LtsAlunos, Militar, MilitaresInativos, NomeConvocado, PostoGrad, Quadro, Obm, Localidade, Funcao, RecompensaAluno, RestricaoAluno, SancaoAluno, SegundoVinculo, SituacaoConvocacao, User, FuncaoUser, PublicacaoBg,
                        EstadoCivil, Especialidade, Destino, Motivo, Modalidade, Punicao, Comportamento, MilitarObmFuncao,
                        FuncaoGratificada,
                        MilitaresAgregados, MilitaresADisposicao, LicencaEspecial, LicencaParaTratamentoDeSaude, Paf,
                        Meses, Motoristas, Categoria, TabelaVencimento, ValorDetalhadoPostoGrad, FichaAlunos, AlunoInativo, Viaturas, ViaturaMilitar, MilitarGraduacao, MilitarContatoEmergencia, MilitarConjuge, LogExportacaoExcel, Curso, MilitarCurso, AuditoriaAtualizacaoCadastral, now_manaus_naive)
from src.querys import dados_para_mapa, obter_estatisticas_militares, login_usuario
from src.decorators.control import checar_ocupacao, militar_esta_no_escopo, obms_permitidas_para_usuario, sync_user_admin_obms_from_militar
from src.decorators.business_logic import processar_militares_a_disposicao, processar_militares_agregados, \
    processar_militares_le, processar_militares_lts
from datetime import datetime, date, timedelta
from io import BytesIO
from sqlalchemy.orm import joinedload, selectinload, load_only, aliased
from sqlalchemy import case, distinct, extract, func, not_, or_, cast, String, and_
from sqlalchemy.exc import IntegrityError
from openpyxl import Workbook
from openpyxl.styles import Border, Font, Side
from openpyxl.utils import get_column_letter
from decimal import Decimal, ROUND_HALF_UP, getcontext
from docx import Document
from urllib.parse import urlencode
from collections import defaultdict, Counter
from docx.shared import Pt
from docx.oxml.ns import qn
from src.decorators.formatar_datas import formatar_data_extenso, formatar_data_sem_zero
from src.decorators.email_utils import send_reset_password_email, verify_password_reset_token
import re
import plotly.graph_objs as go
import plotly.io as pio
from collections import defaultdict
from src.utils.sa_serialize import sa_to_dict
from sqlalchemy.inspection import inspect as sa_inspect
from src.security.perms import has_perm
from src.authz import is_super_or_perm, can_ferias_bypass_janela, is_super, require_perm
from src.utils.cadastro_status import cadastro_esta_completo
from src.utils.painel import (
    _obter_obm_principal,
    listar_situacoes_atualizacao,
    obter_resumo_atualizacao_cadastral,
    obter_militares_atualizacao_cadastral,
    serializar_militar_atualizacao,
    listar_obms_atualizacao,
    listar_postos_grad_atualizacao,
    obter_detalhes_militar_atualizacao,)
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from src.services.importar_militares import (
    ler_planilha,
    colunas_reconhecidas,
    colunas_nao_reconhecidas,
    analisar_importacao,
    importar_dataframe,
    salvar_historico_importacao,
)
from src.services.militar_situacao_service import (
    parse_date_flex,
    sincronizar_blocos_funcionais,
)
import time
from src.utils.utils import registrar_log_download
from dateutil.relativedelta import relativedelta



@app.route('/gerar-le', methods=['GET', 'POST'])
@login_required
def gerar_le():
    if request.method == 'POST':
        nome = request.form['nome']

        dados = {
            'nota_bg': request.form['nota_bg'],
            'data_do_requerimento': formatar_data_extenso(request.form['data_do_requerimento']),
            # <-- corrigido aqui
            'POSTO/GRADUACAO': request.form['posto_grad'],
            'QUADRO': request.form['quadro'],
            'NOME do MILITAR': nome,
            'OBM do militar': request.form['obm'],
            'tipo_licenca_especial': request.form['tipo_licenca_especial'],
            'data_inicio_licenca_especial a data_fim_licenca_especial': request.form['periodo_licenca'],
            'data_inicio_pedido': formatar_data_sem_zero(request.form['data_inicio_pedido']),
            'data_de_apresentacao': formatar_data_sem_zero(request.form['data_apresentacao']),
            'numero_siged': request.form['numero_siged'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nota_bg', 'POSTO/GRADUACAO', 'QUADRO',
                   'NOME do MILITAR', 'data_inicio_pedido']
        ITALICO = ['numero_siged']

        doc = Document('src/template/nota_tecnica.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                    if chave in ITALICO:
                        novo_run.italic = True
                        if chave == 'numero_siged':
                            novo_run.text = f"({valor})"
                            novo_run.font.size = Pt(10)
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name='nota_le_gerada.docx')

    return render_template('gerar_le.html')


@app.route('/indeferimento-le', methods=['GET', 'POST'])
@login_required
def indeferimento_le():
    if request.method == 'POST':
        nome = request.form['nome']

        dados = {
            'nota_bg': request.form['nota_bg'],
            'POSTO/GRADUACAO': request.form['posto_grad'],
            'QUADRO': request.form['quadro'],
            'NOME do MILITAR': nome,
            'OBM do militar': request.form['obm'],
            'tipo_licenca_especial': request.form['tipo_licenca_especial'],
            'data_inicio_licenca_especial a data_fim_licenca_especial': request.form['periodo_licenca'],
            'numero_siged': request.form['numero_siged'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nota_bg', 'POSTO/GRADUACAO', 'QUADRO',
                   'NOME do MILITAR', 'data_inicio_licenca_especial a data_fim_licenca_especial']
        ITALICO = ['numero_siged']
        doc = Document('src/template/indeferimento_le.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                    if chave in ITALICO:
                        novo_run.italic = True
                        if chave == 'numero_siged':
                            novo_run.text = f"({valor})"
                            novo_run.font.size = Pt(10)
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name='indeferimento_le_gerada.docx')

    return render_template('indeferimento_le.html')


@app.route('/gerar-lp', methods=['GET', 'POST'])
@login_required
def gerar_lp():
    if request.method == 'POST':
        dados = {
            'nota_bg': request.form['nota_bg'],
            'matricula_certidao': request.form['matricula_certidao'],
            'cartorio': request.form['cartorio'],
            'nome_filho': request.form['nome_filho'],
            'cidade_natal': request.form['cidade_natal'],
            'pai': request.form['pai'],
            'mae': request.form['mae'],
            'data_certidao': formatar_data_extenso(request.form['data_certidao']),
            'oficial_responsavel': request.form['oficial_responsavel'],
            'data_inicio_lp': formatar_data_extenso(request.form['data_inicio_lp']),
            'data_apresentacao': formatar_data_extenso(request.form['data_apresentacao']),
            'numero_siged': request.form['numero_siged'],
            'posto_graduacao': request.form['posto_grad'],
            'quadro': request.form['quadro'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }
        NEGRITO = ['nota_bg', 'pai', 'posto_graduacao', 'quadro',
                   'data_inicio_lp', 'data_apresentacao', 'matricula_certidao']

        ITALICO = ['numero_siged']

        doc = Document('src/template/certidao_lp.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                    if chave in ITALICO:
                        novo_run.italic = True
                        if chave == 'numero_siged':
                            novo_run.text = f"({valor})"
                            novo_run.font.size = Pt(10)
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name='certidao_lp_gerada.docx')

    return render_template('gerar_lp.html')


@app.route('/gerar-certidao-casamento', methods=['GET', 'POST'])
@login_required
def gerar_certidao_casamento():
    if request.method == 'POST':

        dados = {
            'nota_bg': request.form['nota_bg'],
            'matricula_certidao': request.form['matricula_certidao'],
            'cartorio': request.form['cartorio'],
            'esposo': request.form['esposo'],
            'posto_grad': request.form['posto_grad'],
            'quadro': request.form['quadro'],
            'esposa': request.form['esposa'],
            'posto_grad_esposa': request.form['posto_grad_esposa'],
            'quadro_esposa': request.form['quadro_esposa'],
            'data_casamento': formatar_data_extenso(request.form['data_casamento']),
            'regime_casamento': request.form['regime_casamento'],
            'escrevente': request.form['escrevente'],
            'data_registro': formatar_data_extenso(request.form['data_registro']),
            'data_inicio_licenca': formatar_data_sem_zero(request.form['data_inicio_licenca']),
            'data_apresentacao': formatar_data_sem_zero(request.form['data_apresentacao']),
            'numero_siged': request.form['numero_siged'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nota_bg', 'matricula_certidao', 'cartorio',
                   'esposo', 'esposa', 'data_casamento', 'regime_casamento', 'data_inicio_licenca', 'data_apresentacao']

        ITALICO = ['numero_siged']

        doc = Document('src/template/certidao_casamento.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                    if chave in ITALICO:
                        novo_run.italic = True
                        if chave == 'numero_siged':
                            novo_run.text = f"({valor})"
                            novo_run.font.size = Pt(10)
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name='certidao_casamento_gerada.docx')

    return render_template('gerar_certidao_casamento.html')


@app.route('/gerar=certidao-obito', methods=['GET', 'POST'])
@login_required
def gerar_certidao_obito():
    if request.method == 'POST':
        dados = {
            'nota_bg': request.form['nota_bg'],
            'numero_certidao': request.form['numero_certidao'],
            'cartorio': request.form['cartorio'],
            'cidade_estado': request.form['cidade_estado'],
            'nome_falecido': request.form['nome_falecido'],
            'cidade_falecido': request.form['cidade_falecido'],
            'data_falecimento': formatar_data_extenso(request.form['data_falecimento']),
            'posto_grad': request.form['posto_grad'],
            'quadro': request.form['quadro'],
            'nome_militar': request.form['nome_militar'],
            'escrevente': request.form['escrevente'],
            'data_inicio_licenca': formatar_data_sem_zero(request.form['data_inicio_licenca']),
            'data_apresentacao': formatar_data_sem_zero(request.form['data_apresentacao']),
            'numero_siged': request.form['numero_siged'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nota_bg', 'data_falecimento', 'posto_grad', 'quadro',
                   'nome_militar', 'data_inicio_licenca', 'data_apresentacao']

        ITALICO = ['numero_siged']

        doc = Document('src/template/certidao_obito.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                    if chave in ITALICO:
                        novo_run.italic = True
                        if chave == 'numero_siged':
                            novo_run.text = f"({valor})"
                            novo_run.font.size = Pt(10)
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name='certidao_obito.docx')

    return render_template('gerar_certidao_obito.html')


@app.route('/gerar-certidao-tempo-servico', methods=['GET', 'POST'])
@login_required
def certidao_tempo_servico():
    if request.method == 'POST':
        nome = request.form['nome_completo'].replace(" ", "_")
        dados = {
            'nome_completo': request.form['nome_completo'],
            'posto_grad': request.form['posto_grad'],
            'cpf': request.form['cpf'],
            'dia_ingresso': request.form['dia_ingresso'],
            'mes_ingresso': request.form['mes_ingresso'],
            'ano_ingresso': request.form['ano_ingresso'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nome_completo', 'posto_grad', 'cpf',
                   'dia_ingresso', 'mes_ingresso', 'ano_ingresso', 'data_atual']

        doc = Document('src/template/declaracao_tempo_de_servico.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name=f'declaracao_tempo_de_servico{nome}.docx')

    return render_template('gerar_certidao_tempo_servico.html')


@app.route('/gerar-certidao-exerc-atp', methods=['GET', 'POST'])
@login_required
def certidao_exercicio_atv_atipica():
    if request.method == 'POST':
        nome = request.form['nome_completo'].replace(" ", "_")
        dados = {
            'nome_completo': request.form['nome_completo'],
            'posto_grad': request.form['posto_grad'],
            'cpf': request.form['cpf'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nome_completo', 'posto_grad', 'cpf']

        ITALICO = ['data_atual']

        doc = Document('src/template/exercicio_atividade_atipica.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True

                    if chave in ITALICO:
                        novo_run.italic = True
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name=f'declaracao_exercicio_atividade_{nome}.docx')

    return render_template('gerar_exc_atv.html')


@app.route('/gerar-declaracao', methods=['GET', 'POST'])
@login_required
def declaracao():
    if request.method == 'POST':
        nome = request.form['nome_militar'].replace(" ", "_")
        dados = {
            'nota_declaracao': request.form['nota_declaracao'],
            'nome_militar': request.form['nome_militar'],
            'orgao': request.form['orgao'],
            'posto_graduacao': request.form['posto_graduacao'],
            'quadro': request.form['quadro'],
            'cpf': request.form['cpf'],
            'rg_cbmam': request.form['rg_cbmam'],
            'matricula': request.form['matricula'],
            'especialidade': request.form['especialidade'],
            'data_concurso': formatar_data_sem_zero(request.form['data_concurso']),
            'numero_bg': request.form['numero_bg'],
            'data_bg': formatar_data_sem_zero(request.form['data_bg']),
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nome_militar', 'posto_graduacao',
                   'quadro', 'cpf', 'nota_declaracao']

        doc = Document('src/template/declaracao1.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name=f'declaracao_{nome}.docx')

    return render_template('declaracao.html')


@app.route('/nota-elogio', methods=['GET', 'POST'])
@login_required
def nota_elogio():
    if request.method == 'POST':
        nome = request.form['nome_militar']

        dados = {
            'nota_bg': request.form['nota_bg'],
            # <-- corrigido aqui
            'posto_graduacao': request.form['posto_graduacao'],
            'quadro': request.form['quadro'],
            'nome_militar': nome,
            'atestador': request.form['atestador'],
            'data_doacao': formatar_data_sem_zero(request.form['data_doacao']),
            'numero_siged': request.form['numero_siged'],
            'numero_coren': request.form['numero_coren'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nota_bg', 'posto_graduacao', 'quadro',
                   'nome_militar', 'data_doacao']
        ITALICO = ['numero_siged']

        doc = Document('src/template/doacao_sangue.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                    if chave in ITALICO:
                        novo_run.italic = True
                        if chave == 'numero_siged':
                            novo_run.text = f"({valor})"
                            novo_run.font.size = Pt(10)
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name=f'elogio_{nome}.docx')

    return render_template('elogio_doacao.html')


@app.route('/portaria-gratificacao', methods=['GET', 'POST'])
@login_required
def portaria_gratificacao():
    if request.method == 'POST':
        # Nome para gerar o arquivo final
        nome_servidor = request.form['nome_servidor'].replace(" ", "_")

        # DADOS PARA PREENCHER O DOCUMENTO
        dados = {
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
            'numero_processo': request.form['numero_processo'],
            'nome_juiz': request.form['nome_juiz'],
            'porcentagem': request.form['porcentagem'],
            'porcentagem_por_extenso': request.form['porcentagem_por_extenso'],
            'nome_servidor': request.form['nome_servidor'],
            'matricula': request.form['matricula'],
            'data_a_contar': formatar_data_sem_zero(request.form['data_a_contar']),
            'memo': request.form['memo']
        }

        # Campos que ficarão em negrito (opcional, pode editar)
        NEGRITO = [
            'nome_servidor', 'matricula', 'porcentagem',
            'numero_processo', 'data_atual'
        ]

        doc = Document('src/template/portaria_gc.docx')

        # PERCORRER TODOS OS PARÁGRAFOS
        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove runs para reconstruir o parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True

                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        # EXPORTA ARQUIVO FINAL
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(
            output,
            as_attachment=True,
            download_name=f'portaria_gratificacao_{nome_servidor}.docx'
        )

    return render_template('portaria_gratificacao.html')
