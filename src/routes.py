import math
from flask_wtf.csrf import validate_csrf
from flask_login import login_required
from flask import abort, request, jsonify, make_response, current_app
from random import shuffle
import os
import zipfile
import qrcode
import pytz
import pandas as pd
import base64
import matplotlib.pyplot as plt
from src.decorators.utils_acumulo import b2_bucket_name, b2_client, b2_delete_all_versions, b2_upload_fileobj
from src.identificacao import buscar_pessoa_por_cpf, normaliza_matricula
from src.formatar_cpf import cadete_restantes, formatar_cpf, get_militar_por_user, is_cadete
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from flask import render_template, redirect, url_for, request, flash, jsonify, session, send_file, make_response, \
    Response, stream_with_context
from flask_login import login_user, logout_user, login_required, current_user
from flask_wtf.csrf import validate_csrf, generate_csrf
from werkzeug.utils import secure_filename
from src import app, database, bcrypt
from src.forms import (AtualizacaoCadastralForm, ControleConvocacaoForm, CriarSenhaForm, DistribuirAtualizacaoForm, FichaAlunosForm, FormFiltroMilitar, FormMilitarInativo,
                       IdentificacaoForm, ImpactoForm, FormLogin, FormMilitar, FormCriarUsuario, FormMotoristas, FormFiltroMotorista, LtsAlunoForm, RecompensaAlunoForm,
                       RestricaoAlunoForm, SancaoAlunoForm, TabelaVencimentoForm, InativarAlunoForm, TokenForm, MatriculaConfirmForm)
from src.models import (ControleConvocacao, Convocacao, DocumentoMilitar, LtsAlunos, Militar, MilitaresInativos, NomeConvocado, PostoGrad, Quadro, Obm, Localidade, Funcao, RecompensaAluno, RestricaoAluno, SancaoAluno, SegundoVinculo, Situacao, SituacaoConvocacao, TarefaAtualizacaoCadete, User, FuncaoUser, PublicacaoBg,
                        EstadoCivil, Especialidade, Destino, Agregacoes, Punicao, Comportamento, MilitarObmFuncao,
                        FuncaoGratificada,
                        MilitaresAgregados, MilitaresADisposicao, LicencaEspecial, LicencaParaTratamentoDeSaude, Paf,
                        Meses, Motoristas, Categoria, TabelaVencimento, ValorDetalhadoPostoGrad, FichaAlunos, AlunoInativo, TokenVerificacao, Viaturas, ViaturaMilitar)
from src.querys import dados_para_mapa, efetivo_oficiais_por_obm, obter_estatisticas_militares, login_usuario
from src.decorators.control import checar_ocupacao
from src.decorators.business_logic import processar_militares_a_disposicao, processar_militares_agregados, \
    processar_militares_le, processar_militares_lts
from datetime import datetime, date, timedelta, timezone
from io import BytesIO
from sqlalchemy.orm import joinedload, selectinload, load_only, aliased
from sqlalchemy import case, func, text, or_, cast, String, and_
from sqlalchemy.exc import IntegrityError
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter
from decimal import Decimal, ROUND_HALF_UP, getcontext
from docx import Document
from urllib.parse import urlencode
from collections import defaultdict, Counter
from docx.shared import Pt
from docx.oxml.ns import qn
from src.decorators.formatar_datas import formatar_data_extenso, formatar_data_sem_zero
import re
import plotly.graph_objs as go
import plotly.io as pio
from src.routes_acumulo import _obms_ativas_do_militar, bp_acumulo
import time
import statistics
from collections import defaultdict
from src.utils.sa_serialize import sa_to_dict
from sqlalchemy.inspection import inspect as sa_inspect


def _pode_pegar_doc(doc: DocumentoMilitar) -> bool:
    # Só o dono (CPF) ou alguém com poder (se quiser permitir admins):
    if current_user.cpf == doc.destinatario_cpf:
        return True
    # Exemplo: permitir DRH baixar em nome do usuário (opcional)
    try:
        # gambizinha pra reutilizar
        return checar_ocupacao('DRH', 'SUPER USER')(lambda: True)() is True
    except Exception:
        return False


@app.route("/db-ping-10")
def db_ping_10():
    times = []
    for _ in range(10):
        t0 = time.perf_counter()
        database.session.execute(text("SELECT 1"))
        times.append((time.perf_counter()-t0)*1000)
    return {
        "avg_ms": round(statistics.mean(times), 1),
        "p50_ms": round(statistics.median(times), 1),
        "min_ms": round(min(times), 1),
        "max_ms": round(max(times), 1),
        "samples": [round(x, 1) for x in times]
    }


@app.route("/db-ping-conn")
def db_ping_conn():
    times = []
    with database.engine.connect() as conn:
        for _ in range(20):
            t0 = time.perf_counter()
            conn.execute(text("SELECT 1")).fetchone()
            times.append((time.perf_counter()-t0)*1000)
    return {
        "avg_ms": round(statistics.mean(times), 1),
        "p50_ms": round(statistics.median(times), 1),
        "min_ms": round(min(times), 1),
        "max_ms": round(max(times), 1),
        "samples": [round(x, 1) for x in times]
    }


@app.route("/db-ping-pool")
def db_ping_pool():
    times = []
    for _ in range(20):
        t0 = time.perf_counter()
        # pega/devolve conexão sempre
        database.session.execute(text("SELECT 1"))
        times.append((time.perf_counter()-t0)*1000)
    return {
        "avg_ms": round(statistics.mean(times), 1),
        "p50_ms": round(statistics.median(times), 1),
        "min_ms": round(min(times), 1),
        "max_ms": round(max(times), 1),
        "samples": [round(x, 1) for x in times]
    }


@app.context_processor
def inject_militar_atual():
    mil = None
    try:
        if current_user.is_authenticated:
            mil = get_militar_por_user(current_user)  # usa teu helper
    except Exception:
        mil = None
    return {
        "militar_atual": mil,
        "militar_id_atual": (mil.id if mil else None),
    }


@app.route('/acesso-negado')
def acesso_negado():
    """Rota para exibir a página de acesso negado."""
    return render_template('acesso_negado.html')


@app.route('/api/estatisticas', methods=['GET'])
def estatisticas():
    """Retorna as estatísticas dos militares em formato JSON."""
    estatisticas = obter_estatisticas_militares()
    return jsonify(estatisticas)


@app.route('/api/login', methods=['POST'])
def api_login():
    data = request.get_json()

    cpf = data.get("cpf")
    senha = data.get("senha")

    if not cpf or not senha:
        return jsonify({"status": "erro", "mensagem": "CPF e senha são obrigatórios"}), 400

    user = login_usuario(cpf, senha)

    if user:
        return jsonify({
            "status": "sucesso",
            "mensagem": "Login realizado com sucesso",
            "usuario": {
                "id": user.id,
                "nome": user.nome,
                "cpf": user.cpf,
                "email": user.email,
                "obm1": user.obm1.sigla if user.obm1 else None,
                "obm2": user.obm2.sigla if user.obm2 else None
            }
        }), 200
    else:
        return jsonify({"status": "erro", "mensagem": "CPF ou senha inválidos"}), 401


@app.route("/")
@login_required
def home():
    if current_user.funcao_user_id == 12:
        return redirect(url_for('home_atualizacao'))

    try:
        # enviando e-mails
        processar_militares_agregados()
        processar_militares_a_disposicao()
        processar_militares_le()
        processar_militares_lts()

        militares_le = LicencaEspecial.query.all()

        militares_lts = LicencaParaTratamentoDeSaude.query.all()

        # estatísticas dos militares para o dashboard
        estatisticas = obter_estatisticas_militares()

        return render_template('home.html', **estatisticas, militares_le=militares_le, militares_lts=militares_lts)
    except Exception as e:
        print(f"Erro: {e}")
        return jsonify({'error': 'Erro ao carregar a página'}), 500


def get_user_ip():
    # Verifica se o cabeçalho X-Forwarded-For está presente
    if request.headers.get('X-Forwarded-For'):
        # Pode conter múltiplos IPs, estou pegando o primeiro
        ip = request.headers.getlist('X-Forwarded-For')[0]
    else:
        # Fallback para o IP remoto
        ip = request.remote_addr
    return ip


@app.route("/login", methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        militar = get_militar_por_user(current_user)
        session['militar_id'] = militar.id
        flash('Você já está logado.', 'alert-info')
        return redirect(url_for('home'))

    form_login = FormLogin()
    if form_login.validate_on_submit() and 'botao_submit_login' in request.form:
        cpf = User.query.filter_by(cpf=form_login.cpf.data).first()

        if cpf and bcrypt.check_password_hash(cpf.senha, form_login.senha.data):
            login_user(cpf, remember=form_login.lembrar_dados.data)
            flash('Login feito com sucesso!', 'alert-success')

            fuso_horario = pytz.timezone('America/Manaus')
            cpf.data_ultimo_acesso = datetime.now(fuso_horario)
            cpf.ip_address = get_user_ip()

            database.session.commit()

            # Agora a verificação de função refinada
            if cpf.funcao_user_id in [1, 2]:  # Diretor ou Chefe
                return redirect(url_for('exibir_ferias_chefe'))
            else:
                return redirect(url_for('home'))
        else:
            flash('Falha no Login, CPF ou senha incorretos.', 'alert-danger')

    return render_template('login.html', form_login=form_login)


@app.route("/criar-conta", methods=['GET', 'POST'])
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'SUPER USER')
def criar_conta():
    form_criar_usuario = FormCriarUsuario()

    choices = [(funcao_user.id, funcao_user.ocupacao)
               for funcao_user in FuncaoUser.query.all()]
    choices.insert(0, ('', '-- Selecione uma opção --'))
    form_criar_usuario.obm_id_1.choices = [('', '-- Selecione uma opção --')] + [(obm.id, obm.sigla) for obm in
                                                                                 Obm.query.all()]
    form_criar_usuario.obm_id_2.choices = [('', '-- Selecione uma opção --')] + [(obm.id, obm.sigla) for obm in
                                                                                 Obm.query.all()]

    form_criar_usuario.localidade_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(localidade.id, localidade.sigla) for localidade in
         Localidade.query.all()]

    form_criar_usuario.funcao_user_id.choices = choices

    if form_criar_usuario.validate_on_submit():
        senha_cript = bcrypt.generate_password_hash(
            form_criar_usuario.senha.data).decode('utf-8')
        usuarios = User(nome=form_criar_usuario.nome.data,
                        email=form_criar_usuario.email.data,
                        funcao_user_id=form_criar_usuario.funcao_user_id.data,
                        cpf=form_criar_usuario.cpf.data,
                        obm_id_1=form_criar_usuario.obm_id_1.data,
                        obm_id_2=form_criar_usuario.obm_id_2.data,
                        localidade_id=form_criar_usuario.localidade_id.data,
                        senha=senha_cript)
        database.session.add(usuarios)
        database.session.commit()
        flash("Usuário cadastrado com sucesso!", "alert-success")
        return redirect(url_for('home'))
    return render_template('criar_conta.html', form_criar_usuario=form_criar_usuario)


@app.route("/adicionar-militar", methods=['GET', 'POST'])
@login_required
@checar_ocupacao('DRH', 'MAPA DA FORÇA', 'SUPER USER', 'DIRETOR DRH')
def adicionar_militar():
    form_militar = FormMilitar()

    form_militar.funcao_gratificada_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(funcao_gratificada.id, funcao_gratificada.gratificacao) for
         funcao_gratificada in FuncaoGratificada.query.all()]

    form_militar.posto_grad_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(posto.id, posto.sigla) for posto in PostoGrad.query.all()]

    form_militar.quadro_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(quadro.id, quadro.quadro) for quadro in Quadro.query.all()]

    form_militar.localidade_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(localidade.id, localidade.sigla) for localidade in
         Localidade.query.all()]

    form_militar.obm_ids_1.choices = [
        ('', '-- Selecione uma opção --')] + [(obm.id, obm.sigla) for obm in
                                              Obm.query.all()]

    form_militar.funcao_ids_1.choices = [
        ('', '-- Selecione uma opção --')] + [(funcao.id, funcao.ocupacao) for
                                              funcao in Funcao.query.all()]

    form_militar.obm_ids_2.choices = [
        ('', '-- Selecione uma opção --')] + [(obm.id, obm.sigla) for obm in
                                              Obm.query.all()]

    form_militar.funcao_ids_2.choices = [
        ('', '-- Selecione uma opção --')] + [(funcao.id, funcao.ocupacao) for
                                              funcao in Funcao.query.all()]

    form_militar.situacao_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(situacao.id, situacao.condicao) for situacao in Situacao.query.all()]

    form_militar.estado_civil.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(estado.id, estado.estado) for estado in EstadoCivil.query.all()]

    form_militar.especialidade_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(especialidade.id, especialidade.ocupacao) for especialidade in
         Especialidade.query.all()]

    form_militar.destino_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(destino.id, destino.local) for destino in Destino.query.all()]

    form_militar.agregacoes_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(agregacoes.id, agregacoes.tipo) for agregacoes in Agregacoes.query.all()]

    form_militar.punicao_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(punicao.id, punicao.sancao) for punicao in Punicao.query.all()]

    form_militar.comportamento_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(comportamento.id, comportamento.conduta) for comportamento in
         Comportamento.query.all()]

    if form_militar.validate_on_submit():

        completa_25_inclusao = datetime.strptime(
            form_militar.completa_25_inclusao.data, '%d/%m/%Y').date()
        completa_30_inclusao = datetime.strptime(
            form_militar.completa_30_inclusao.data, '%d/%m/%Y').date()
        completa_25_anos_sv = datetime.strptime(
            form_militar.completa_25_anos_sv.data, '%d/%m/%Y').date()
        completa_30_anos_sv = datetime.strptime(
            form_militar.completa_30_anos_sv.data, '%d/%m/%Y').date()

        militar = Militar(
            nome_completo=form_militar.nome_completo.data,
            nome_guerra=form_militar.nome_guerra.data,
            cpf=form_militar.cpf.data,
            rg=form_militar.rg.data,
            nome_pai=form_militar.nome_pai.data,
            nome_mae=form_militar.nome_mae.data,
            matricula=form_militar.matricula.data,
            pis_pasep=form_militar.pis_pasep.data,
            num_titulo_eleitor=form_militar.num_titulo_eleitor.data,
            digito_titulo_eleitor=form_militar.digito_titulo_eleitor.data,
            zona=form_militar.zona.data,
            secao=form_militar.secao.data,
            posto_grad_id=form_militar.posto_grad_id.data,
            quadro_id=form_militar.quadro_id.data,
            localidade_id=form_militar.localidade_id.data,
            antiguidade=form_militar.antiguidade.data,
            sexo=form_militar.sexo.data,
            raca=form_militar.raca.data,
            data_nascimento=form_militar.data_nascimento.data,
            inclusao=form_militar.inclusao.data,
            completa_25_inclusao=completa_25_inclusao,
            completa_30_inclusao=completa_30_inclusao,
            punicao_id=form_militar.punicao_id.data,
            comportamento_id=form_militar.comportamento_id.data or None,
            efetivo_servico=form_militar.efetivo_servico.data,
            completa_25_anos_sv=completa_25_anos_sv,
            completa_30_anos_sv=completa_30_anos_sv,
            anos=form_militar.anos.data,
            meses=form_militar.meses.data,
            dias=form_militar.dias.data,
            total_dias=form_militar.total_dias.data or None,
            idade_reserva_grad=0,
            estado_civil=form_militar.estado_civil.data,
            especialidade_id=form_militar.especialidade_id.data,
            pronto=form_militar.pronto.data,
            situacao_id=form_militar.situacao_id.data or None,
            agregacoes_id=form_militar.agregacoes_id.data or None,
            destino_id=form_militar.destino_id.data or None,
            inicio_periodo=form_militar.inicio_periodo.data,
            fim_periodo=form_militar.fim_periodo.data,
            ltip_afastamento_cargo_eletivo=form_militar.ltip_afastamento_cargo_eletivo.data,
            periodo_ltip=form_militar.periodo_ltip.data,
            total_ltip=form_militar.total_ltip.data,
            completa_25_anos_ltip=form_militar.completa_25_anos_ltip.data,
            completa_30_anos_ltip=form_militar.completa_30_anos_ltip.data,
            cursos=form_militar.cursos.data,
            grau_instrucao=form_militar.grau_instrucao.data,
            graduacao=form_militar.graduacao.data,
            pos_graduacao=form_militar.pos_graduacao.data,
            mestrado=form_militar.mestrado.data,
            doutorado=form_militar.doutorado.data,
            cfsd=form_militar.cfsd.data,
            cfc=form_militar.cfc.data,
            cfs=form_militar.cfs.data,
            cas=form_militar.cas.data,
            choa=form_militar.choa.data,
            cfo=form_militar.cfo.data,
            cbo=form_militar.cbo.data,
            cao=form_militar.cao.data,
            csbm=form_militar.csbm.data,
            cursos_civis=form_militar.cursos_civis.data,
            endereco=form_militar.endereco.data,
            complemento=form_militar.complemento.data,
            cidade=form_militar.cidade.data,
            estado=form_militar.estado.data,
            cep=form_militar.cep.data,
            celular=form_militar.celular.data,
            email=form_militar.email.data,
            inclusao_bg=form_militar.inclusao_bg.data,
            soldado_tres=form_militar.soldado_tres.data,
            soldado_dois=form_militar.soldado_dois.data,
            soldado_um=form_militar.soldado_um.data,
            cabo=form_militar.cabo.data,
            terceiro_sgt=form_militar.terceiro_sgt.data,
            segundo_sgt=form_militar.segundo_sgt.data,
            primeiro_sgt=form_militar.primeiro_sgt.data,
            subtenente=form_militar.subtenente.data,
            segundo_tenente=form_militar.segundo_tenente.data,
            primeiro_tenente=form_militar.primeiro_tenente.data,
            cap=form_militar.cap.data,
            maj=form_militar.maj.data,
            tc=form_militar.tc.data,
            cel=form_militar.cel.data,
            funcao_gratificada_id=form_militar.funcao_gratificada_id.data,
            alteracao_nome_guerra=form_militar.alteracao_nome_guerra.data,
            inativo=False,
            ip_address=get_user_ip(),
            usuario_id=current_user.id
        )
        # arquivos = form_militar.arquivo.data

        # # Se for um único arquivo, transforme-o em uma lista para uniformizar o processamento
        # if not isinstance(arquivos, list):
        #     arquivos = [arquivos]

        # for arquivo in arquivos:
        #     if isinstance(arquivo, str):
        #         # Esse erro não deve acontecer se o upload foi feito corretamente, mas informe o usuário se ocorrer
        #         print(f"Erro: Esperado um objeto FileStorage, mas obtido uma string: {arquivo}")
        #     elif arquivo:  # Verifica se o arquivo não está vazio
        #         nome_seguro = secure_filename(arquivo.filename)
        #         caminho = os.path.join(os.path.abspath(os.path.dirname(__file__)), app.config["UPLOAD_FOLDER"],
        #                                nome_seguro)
        #         arquivo.save(caminho)  # Salva o arquivo no caminho especificado
        #         print(f"Arquivo {nome_seguro} salvo com sucesso em {caminho}")
        #     else:
        #         print("Erro: O arquivo está vazio ou é inválido.")

        database.session.add(militar)
        database.session.commit()

        # Adicionando as OBMs e Funções
        obm_funcao_pairs = zip(request.form.getlist(
            'obm_ids_1'), request.form.getlist('funcao_ids_1'))

        # Itera sobre as combinações selecionadas de OBMs e funções
        for obm_id, funcao_id in obm_funcao_pairs:
            # Verifica se obm_id e funcao_id não estão vazios
            if obm_id and funcao_id:
                militar_obm_funcao = MilitarObmFuncao(
                    militar_id=militar.id,
                    # Certifique-se de que o ID é um número inteiro
                    obm_id=int(obm_id),
                    tipo=1,
                    # Certifique-se de que o ID é um número inteiro
                    funcao_id=int(funcao_id)
                )
                database.session.add(militar_obm_funcao)

        # Repetindo o processo para o segundo conjunto de OBMs e funções (caso exista)
        obm_funcao_pairs_2 = zip(request.form.getlist(
            'obm_ids_2'), request.form.getlist('funcao_ids_2'))

        for obm_id_2, funcao_id_2 in obm_funcao_pairs_2:
            if obm_id_2 and funcao_id_2:
                militar_obm_funcao = MilitarObmFuncao(
                    militar_id=militar.id,
                    # Certifique-se de que o ID é um número inteiro
                    obm_id=int(obm_id_2),
                    tipo=2,
                    # Certifique-se de que o ID é um número inteiro
                    funcao_id=int(funcao_id_2)
                )
                database.session.add(militar_obm_funcao)

        database.session.commit()

        # Salvando as publicações de BG
        campos_bg = [
            'transferencia', 'situacao_militar', 'cfsd', 'cfc', 'cfs', 'cas',
            'choa', 'cfo', 'cbo', 'cao', 'csbm', 'soldado_tres',
            'soldado_dois', 'soldado_um', 'cabo', 'terceiro_sgt',
            'segundo_sgt', 'primeiro_sgt', 'subtenente',
            'publicidade_segundo_tenente', 'publicidade_primeiro_tenente',
            'pub_cap', 'pub_maj', 'pub_tc', 'pub_cel', 'pub_alteracao'
        ]

        for campo in campos_bg:
            boletim_geral = getattr(form_militar, campo).data
            if boletim_geral:
                publicacao_bg = PublicacaoBg(
                    militar_id=militar.id,
                    boletim_geral=boletim_geral,
                    tipo_bg=campo
                )
                database.session.add(publicacao_bg)

        # Verifica se a situação selecionada é "AGREGADO"
        situacao_selecionada = Situacao.query.get(
            form_militar.situacao_id.data)
        if situacao_selecionada and situacao_selecionada.condicao == 'AGREGADO':

            # Verifica se há uma publicação BG associada ao militar e à situação
            publicacao_situacao_bg = PublicacaoBg.query.filter_by(militar_id=militar.id,
                                                                  tipo_bg='situacao_militar').first()

            if publicacao_situacao_bg:
                # Criando o registro em 'militares_agregados'
                militar_agregado = MilitaresAgregados(
                    militar_id=militar.id,
                    posto_grad_id=form_militar.posto_grad_id.data,
                    quadro_id=form_militar.quadro_id.data,
                    destino_id=form_militar.destino_id.data,
                    situacao_id=situacao_selecionada.id,
                    inicio_periodo=form_militar.inicio_periodo.data,
                    fim_periodo_agregacao=form_militar.fim_periodo.data,
                    publicacao_bg_id=publicacao_situacao_bg.id
                )

                # Atualizando a posição da agregação (vigente ou término)
                militar_agregado.atualizar_status()

                # Adiciona o registro de agregação e faz o commit no banco
                database.session.add(militar_agregado)
                database.session.commit()
            else:
                # Caso não exista publicação BG associada, você pode decidir o que fazer
                print("Publicação BG não encontrada para o militar agregado.")

        if situacao_selecionada and situacao_selecionada.condicao == 'À DISPOSIÇÃO':

            # Verifica se há uma publicação BG associada ao militar e à situação
            publicacao_situacao_bg = PublicacaoBg.query.filter_by(militar_id=militar.id,
                                                                  tipo_bg='situacao_militar').first()

            if publicacao_situacao_bg:
                # Criando o registro em 'militares_a_disposicao'
                militar_a_disposicao = MilitaresADisposicao(
                    militar_id=militar.id,
                    posto_grad_id=form_militar.posto_grad_id.data,
                    quadro_id=form_militar.quadro_id.data,
                    destino_id=form_militar.destino_id.data,
                    situacao_id=situacao_selecionada.id,
                    inicio_periodo=form_militar.inicio_periodo.data,
                    fim_periodo_disposicao=form_militar.fim_periodo.data,
                    publicacao_bg_id=publicacao_situacao_bg.id
                )

                militar_a_disposicao.atualizar_status()

                # Adiciona o registro de agregação e faz o commit no banco
                database.session.add(militar_a_disposicao)
                database.session.commit()
            else:
                flash(
                    'Publicação BG não encontrada para o militar à disposição.', 'alert-danger')

        if situacao_selecionada and situacao_selecionada.condicao == 'LICENÇA ESPECIAL':
            publicacao_situacao_bg = PublicacaoBg.query.filter_by(militar_id=militar.id,
                                                                  tipo_bg='situacao_militar').first()
            if publicacao_situacao_bg:
                militar_le = LicencaEspecial(
                    militar_id=militar.id,
                    posto_grad_id=form_militar.posto_grad_id.data,
                    quadro_id=form_militar.quadro_id.data,
                    destino_id=form_militar.destino_id.data,
                    situacao_id=situacao_selecionada.id,
                    inicio_periodo_le=form_militar.inicio_periodo.data,
                    fim_periodo_le=form_militar.fim_periodo.data,
                    publicacao_bg_id=publicacao_situacao_bg.id
                )

                militar_le.atualizar_status()

                database.session.add(militar_le)
                database.session.commit()
            else:
                flash(
                    'Publicação BG não encontrada para a Licença Especial.', 'alert-danger')

        if situacao_selecionada and situacao_selecionada.condicao == 'LTS':
            publicacao_situacao_bg = PublicacaoBg.query.filter_by(militar_id=militar.id,
                                                                  tipo_bg='situacao_militar').first()
            if publicacao_situacao_bg:
                militar_lts = LicencaParaTratamentoDeSaude(
                    militar_id=militar.id,
                    posto_grad_id=form_militar.posto_grad_id.data,
                    quadro_id=form_militar.quadro_id.data,
                    destino_id=form_militar.destino_id.data,
                    situacao_id=situacao_selecionada.id,
                    inicio_periodo_lts=form_militar.inicio_periodo.data,
                    fim_periodo_lts=form_militar.fim_periodo.data,
                    publicacao_bg_id=publicacao_situacao_bg.id
                )

                militar_lts.atualizar_status()

                database.session.add(militar_lts)
                database.session.commit()
            else:
                flash('Publicação BG não encontrada para LTS.', 'alert-danger')

        database.session.commit()

        flash('Militar adicionado com sucesso!', 'alert-success')
        return redirect(url_for('home'))

    return render_template('adicionar_militar.html', form_militar=form_militar)


@app.route("/adicionar-militar-inativo", methods=["GET", "POST"])
@login_required
def adicionar_militar_inativo():
    """Rota para adicionar um militar inativo."""
    form_militar = FormMilitarInativo()  # Usando a nova form específica

    form_militar.posto_grad_id.choices = [('', '-- Selecione uma opção --')] + [
        (posto.id, posto.sigla) for posto in PostoGrad.query.all()
    ]
    form_militar.quadro_id.choices = [('', '-- Selecione uma opção --')] + [
        (quadro.id, quadro.quadro) for quadro in Quadro.query.all()
    ]
    form_militar.estado_civil_id.choices = [('', '-- Selecione uma opção --')] + [
        (estado.id, estado.estado) for estado in EstadoCivil.query.all()
    ]

    form_militar.inativo.data = True

    if form_militar.validate_on_submit():
        novo = MilitaresInativos(
            nome_completo=form_militar.nome_completo.data,
            nome_guerra=form_militar.nome_guerra.data,
            estado_civil_id=form_militar.estado_civil_id.data,  # <== aqui a correção
            nome_pai=form_militar.nome_pai.data,
            nome_mae=form_militar.nome_mae.data,
            matricula=form_militar.matricula.data,
            rg=form_militar.rg.data,
            cpf=form_militar.cpf.data,
            pis_pasep=form_militar.pis_pasep.data,
            posto_grad_id=form_militar.posto_grad_id.data,
            quadro_id=form_militar.quadro_id.data,
            sexo=form_militar.sexo.data,
            data_nascimento=form_militar.data_nascimento.data,
            idade_atual=form_militar.idade_atual.data,
            endereco=form_militar.endereco.data,
            complemento=form_militar.complemento.data,
            cidade=form_militar.cidade.data,
            estado=form_militar.estado.data,
            cep=form_militar.cep.data,
            celular=form_militar.celular.data,
            email=form_militar.email.data,
            modalidade=form_militar.modalidade.data,
            doe=form_militar.doe.data,
            usuario_id=current_user.id,
            ip_address=request.remote_addr,
        )

        database.session.add(novo)
        database.session.commit()
        flash("Militar inativo adicionado com sucesso!", "success")
        return redirect(url_for("listar_militares_inativos"))

    return render_template("adicionar_militar_inativo.html", form_militar=form_militar)


@app.route("/militares-inativos")
@login_required
def listar_militares_inativos():
    hierarquia = {
        "CEL": 1,
        "TC": 2,
        "MAJ": 3,
        "CAP": 4,
        "1 TEN": 5,
        "2 TEN": 6,
        "ASP": 7,
        "SUBTENENTE": 8,
        "1 SGT": 9,
        "2 SGT": 10,
        "3 SGT": 11,
        "CB": 12,
        "SD": 13,
    }

    militares = MilitaresInativos.query.join(
        MilitaresInativos.posto_grad).all()

    # Ordenar pela hierarquia definida acima
    militares.sort(key=lambda m: hierarquia.get(
        m.posto_grad.sigla.strip(), 99))

    return render_template("listar_militares_inativos.html", militares=militares)


@app.route("/editar-militar-inativo/<int:id>", methods=["GET", "POST"])
@login_required
def editar_militar_inativo(id):
    militar = MilitaresInativos.query.get_or_404(id)

    form = FormMilitarInativo(obj=militar)

    # ⚠️ Preencha os choices antes de qualquer validação
    form.posto_grad_id.choices = [(p.id, p.sigla)
                                  for p in PostoGrad.query.all()]
    form.quadro_id.choices = [(q.id, q.quadro) for q in Quadro.query.all()]

    form.estado_civil_id.choices = [
        (0, '-- Selecione uma opção --')
    ] + [(estado.id, estado.estado) for estado in EstadoCivil.query.all()]

    if form.validate_on_submit():
        form.populate_obj(militar)
        database.session.commit()
        flash("Dados atualizados com sucesso!", "success")
        return redirect(url_for("listar_militares_inativos"))

    return render_template("adicionar_militar_inativo.html", form_militar=form)


@app.route('/verificar-arquivos', methods=['POST'])
@login_required
def verificar_arquivos():
    data = request.get_json()
    filenames = data.get('filenames', [])

    existing_files = []
    upload_folder = os.path.join(os.path.abspath(
        os.path.dirname(__file__)), app.config["UPLOAD_FOLDER"])

    for filename in filenames:
        file_path = os.path.join(upload_folder, secure_filename(filename))
        if os.path.exists(file_path):
            existing_files.append(filename)

    return jsonify({'exists': existing_files})


@app.route("/exibir-militar/<int:militar_id>", methods=['GET', 'POST'])
@login_required
@checar_ocupacao('DRH', 'MAPA DA FORÇA', 'SUPER USER', 'DIRETOR DRH')
def exibir_militar(militar_id):
    militar = Militar.query.get_or_404(militar_id)
    database.session.expire_all()

    obm_funcao_tipo_1 = MilitarObmFuncao.query.filter_by(militar_id=militar_id, tipo=1) \
        .filter(MilitarObmFuncao.data_fim == None).first()

    obm_funcao_tipo_2 = MilitarObmFuncao.query.filter_by(militar_id=militar_id, tipo=2) \
        .filter(MilitarObmFuncao.data_fim == None).first()

    form_militar = FormMilitar(obj=militar)

    if militar.completa_25_inclusao:
        form_militar.completa_25_inclusao.data = militar.completa_25_inclusao.strftime(
            '%d/%m/%Y')
    if militar.completa_30_inclusao:
        form_militar.completa_30_inclusao.data = militar.completa_30_inclusao.strftime(
            '%d/%m/%Y')
    if militar.completa_25_anos_sv:
        form_militar.completa_25_anos_sv.data = militar.completa_25_anos_sv.strftime(
            '%d/%m/%Y')
    if militar.completa_30_anos_sv:
        form_militar.completa_30_anos_sv.data = militar.completa_30_anos_sv.strftime(
            '%d/%m/%Y')

    form_militar.funcao_gratificada_id.choices = [
        (funcao_gratificada.id, funcao_gratificada.gratificacao) for funcao_gratificada in FuncaoGratificada.query.all()
    ]
    form_militar.posto_grad_id.choices = [
        (posto.id, posto.sigla) for posto in PostoGrad.query.all()
    ]
    form_militar.quadro_id.choices = [
        (quadro.id, quadro.quadro) for quadro in Quadro.query.all()
    ]
    form_militar.localidade_id.choices = [
        (localidade_id.id, localidade_id.sigla) for localidade_id in Localidade.query.all()
    ]

    form_militar.obm_ids_1.choices = ([('', '-- Selecione uma opção --')] +
                                      [(obm.id, obm.sigla) for obm in Obm.query.all()])

    form_militar.funcao_ids_1.choices = ([('', '-- Selecione uma opção --')] +
                                         [(funcao.id, funcao.ocupacao) for funcao in Funcao.query.all()])

    form_militar.obm_ids_2.choices = ([('', '-- Selecione uma opção --')] +
                                      [(obm.id, obm.sigla) for obm in Obm.query.all()])

    form_militar.funcao_ids_2.choices = ([('', '-- Selecione uma opção --')] +
                                         [(funcao.id, funcao.ocupacao) for funcao in Funcao.query.all()])

    form_militar.situacao_id.choices = [
        (situacao.id, situacao.condicao) for situacao in Situacao.query.all()
    ]
    form_militar.estado_civil.choices = [
        (estado.id, estado.estado) for estado in EstadoCivil.query.all()
    ]
    form_militar.especialidade_id.choices = [
        (especialidade.id, especialidade.ocupacao) for especialidade in Especialidade.query.all()
    ]
    form_militar.destino_id.choices = [
        (destino.id, destino.local) for destino in Destino.query.all()
    ]
    form_militar.agregacoes_id.choices = [
        (agregacoes.id, agregacoes.tipo) for agregacoes in Agregacoes.query.all()
    ]
    form_militar.punicao_id.choices = [
        (punicao.id, punicao.sancao) for punicao in Punicao.query.all()
    ]
    form_militar.comportamento_id.choices = ([('', '-- Selecione uma opção --')] +
                                             [(comportamento.id, comportamento.conduta) for comportamento in
                                              Comportamento.query.all()
                                              ])

    if obm_funcao_tipo_1:
        form_militar.obm_ids_1.data = obm_funcao_tipo_1.obm_id
        form_militar.funcao_ids_1.data = obm_funcao_tipo_1.funcao_id

    if obm_funcao_tipo_2:
        form_militar.obm_ids_2.data = obm_funcao_tipo_2.obm_id
        form_militar.funcao_ids_2.data = obm_funcao_tipo_2.funcao_id

    form_militar.sexo.data = militar.sexo if militar.sexo else None
    form_militar.raca.data = militar.raca if militar.raca else None

    hoje = date.today()

    dn = militar.data_nascimento
    # se vier datetime, vira date; se None, fica None
    if isinstance(dn, datetime):
        dn = dn.date()

    if dn:
        idade = hoje.year - dn.year - \
            ((hoje.month, hoje.day) < (dn.month, dn.day))
    else:
        idade = None  # ou 0, se preferir

    form_militar.idade_atual.data = idade
    campos_bg = [
        'transferencia', 'situacao_militar', 'cfsd', 'cfc', 'cfs', 'cas',
        'choa', 'cfo', 'cbo', 'cao', 'csbm', 'soldado_tres',
        'soldado_dois', 'soldado_um', 'cabo', 'terceiro_sgt',
        'segundo_sgt', 'primeiro_sgt', 'subtenente',
        'publicidade_segundo_tenente', 'publicidade_primeiro_tenente',
        'pub_cap', 'pub_maj', 'pub_tc', 'pub_cel', 'pub_alteracao', 'situacao_militar_2',
    ]

    # Recupera as publicações do BG e define campos para publicação do militar
    publicacoes_bg = PublicacaoBg.query.filter_by(militar_id=militar.id).all()

    # Inicializa todos os campos de publicações BG com valor vazio
    for campo in campos_bg:
        if hasattr(form_militar, campo):
            getattr(form_militar, campo).data = ""

    # Itera sobre as publicações do BG e verifica os campos
    for pub in publicacoes_bg:
        if pub.tipo_bg in campos_bg:
            if hasattr(form_militar, pub.tipo_bg):
                # Se o boletim_geral existir, utiliza o valor; caso contrário, mantém como vazio
                getattr(form_militar, pub.tipo_bg).data = pub.boletim_geral or ""

    bg_sit2_ultima = PublicacaoBg.query.filter_by(
        militar_id=militar.id, tipo_bg='situacao_militar_2'
    ).order_by(PublicacaoBg.id.desc()).first()
    bg_sit2_val = bg_sit2_ultima.boletim_geral if bg_sit2_ultima else ""

    if hasattr(form_militar, 'situacao_militar_2'):
        form_militar.situacao_militar_2.data = bg_sit2_val

    def parse_date(d):
        """Aceita date, string 'YYYY-MM-DD' ou 'DD/MM/YYYY'. Retorna date ou None."""
        if not d:
            return None
        if isinstance(d, datetime):
            return d.date()
        if hasattr(d, 'strftime'):  # já é date
            return d
        s = str(d).strip()
        for fmt in ('%Y-%m-%d', '%d/%m/%Y'):
            try:
                return datetime.strptime(s, fmt).date()
            except ValueError:
                pass
        return None  # não converteu

    def safe_bg_id(militar_id):
        """Tenta pegar BG de situacao_militar; se não achar, retorna None (não bloqueia criação)."""
        bg = PublicacaoBg.query.filter_by(
            militar_id=militar_id, tipo_bg='situacao_militar').first()
        return bg.id if bg else None

    def encerrar_status_anteriores(militar, condicao_atual):
        """Fecha registros vigentes de AGREGADO / À DISPOSIÇÃO que não batem com a situação atual."""
        hoje = date.today()
        ontem = hoje - timedelta(days=1)

        # Se a situação atual NÃO é AGREGADO, fecha agregações vigentes
        if condicao_atual != 'AGREGADO':
            agregacoes_vigentes = MilitaresAgregados.query.filter(
                MilitaresAgregados.militar_id == militar.id,
                MilitaresAgregados.inicio_periodo <= hoje,
                or_(
                    MilitaresAgregados.fim_periodo_agregacao == None,
                    MilitaresAgregados.fim_periodo_agregacao >= hoje,
                )
            ).all()

            for ag in agregacoes_vigentes:
                ag.fim_periodo_agregacao = ontem  # <<< aqui muda
                ag.atualizar_status()

        # Se a situação atual NÃO é À DISPOSIÇÃO, fecha disposições vigentes
        if condicao_atual != 'À DISPOSIÇÃO':
            disposicoes_vigentes = MilitaresADisposicao.query.filter(
                MilitaresADisposicao.militar_id == militar.id,
                MilitaresADisposicao.inicio_periodo <= hoje,
                or_(
                    MilitaresADisposicao.fim_periodo_disposicao == None,
                    MilitaresADisposicao.fim_periodo_disposicao >= hoje,
                )
            ).all()

            for disp in disposicoes_vigentes:
                disp.fim_periodo_disposicao = ontem  # <<< e aqui também
                disp.atualizar_status()

    if form_militar.validate_on_submit():
        form_militar.process(request.form)

        militar.nome_completo = form_militar.nome_completo.data
        militar.nome_guerra = form_militar.nome_guerra.data
        militar.cpf = form_militar.cpf.data
        militar.rg = form_militar.rg.data
        militar.nome_pai = form_militar.nome_pai.data
        militar.nome_mae = form_militar.nome_mae.data
        militar.matricula = form_militar.matricula.data
        militar.pis_pasep = form_militar.pis_pasep.data
        militar.num_titulo_eleitor = form_militar.num_titulo_eleitor.data
        militar.digito_titulo_eleitor = form_militar.digito_titulo_eleitor.data
        militar.zona = form_militar.zona.data
        militar.secao = form_militar.secao.data
        militar.posto_grad_id = form_militar.posto_grad_id.data
        militar.quadro_id = form_militar.quadro_id.data
        militar.localidade_id = form_militar.localidade_id.data
        militar.antiguidade = form_militar.antiguidade.data
        militar.sexo = form_militar.sexo.data
        militar.raca = form_militar.raca.data
        militar.data_nascimento = form_militar.data_nascimento.data
        militar.inclusao = form_militar.inclusao.data

        militar.completa_25_inclusao = datetime.strptime(
            str(form_militar.completa_25_inclusao.data), '%d/%m/%Y'
        ).date() if form_militar.completa_25_inclusao.data else None

        militar.completa_30_inclusao = datetime.strptime(
            str(form_militar.completa_30_inclusao.data), '%d/%m/%Y'
        ).date() if form_militar.completa_30_inclusao.data else None

        militar.completa_25_anos_sv = datetime.strptime(
            str(form_militar.completa_25_anos_sv.data), '%d/%m/%Y'
        ).date() if form_militar.completa_25_anos_sv.data else None

        militar.completa_30_anos_sv = datetime.strptime(
            str(form_militar.completa_30_anos_sv.data), '%d/%m/%Y'
        ).date() if form_militar.completa_30_anos_sv.data else None

        militar.inicio_periodo = datetime.strptime(
            str(form_militar.inicio_periodo.data), '%Y-%m-%d'
        ).date() if form_militar.inicio_periodo.data else None

        militar.fim_periodo = datetime.strptime(
            str(form_militar.fim_periodo.data), '%Y-%m-%d'
        ).date() if form_militar.fim_periodo.data else None

        militar.punicao_id = form_militar.punicao_id.data
        militar.comportamento_id = form_militar.comportamento_id.data
        militar.efetivo_servico = form_militar.efetivo_servico.data
        militar.anos = form_militar.anos.data
        militar.meses = form_militar.meses.data
        militar.dias = form_militar.dias.data
        militar.total_dias = form_militar.total_dias.data
        militar.idade_reserva_grad = 0
        militar.estado_civil = form_militar.estado_civil.data
        militar.especialidade_id = form_militar.especialidade_id.data
        militar.pronto = form_militar.pronto.data
        militar.situacao_id = form_militar.situacao_id.data
        militar.agregacoes_id = form_militar.agregacoes_id.data
        militar.destino_id = form_militar.destino_id.data
        militar.inicio_periodo = parse_date(form_militar.inicio_periodo.data)
        militar.fim_periodo = parse_date(form_militar.fim_periodo.data)
        militar.ltip_afastamento_cargo_eletivo = form_militar.ltip_afastamento_cargo_eletivo.data
        militar.periodo_ltip = form_militar.periodo_ltip.data
        militar.total_ltip = form_militar.total_ltip.data
        militar.completa_25_anos_ltip = form_militar.completa_25_anos_ltip.data
        militar.completa_30_anos_ltip = form_militar.completa_30_anos_ltip.data
        militar.cursos = form_militar.cursos.data
        militar.grau_instrucao = form_militar.grau_instrucao.data
        militar.graduacao = form_militar.graduacao.data
        militar.pos_graduacao = form_militar.pos_graduacao.data
        militar.mestrado = form_militar.mestrado.data
        militar.doutorado = form_militar.doutorado.data
        militar.cfsd = form_militar.cfsd.data
        militar.cfc = form_militar.cfc.data
        militar.cfs = form_militar.cfs.data
        militar.cas = form_militar.cas.data
        militar.choa = form_militar.choa.data
        militar.cfo = form_militar.cfo.data
        militar.cbo = form_militar.cbo.data
        militar.cao = form_militar.cao.data
        militar.csbm = form_militar.csbm.data
        militar.cursos_civis = form_militar.cursos_civis.data
        militar.endereco = form_militar.endereco.data
        militar.complemento = form_militar.complemento.data
        militar.cidade = form_militar.cidade.data
        militar.estado = form_militar.estado.data
        militar.cep = form_militar.cep.data
        militar.celular = form_militar.celular.data
        militar.email = form_militar.email.data
        militar.inclusao_bg = form_militar.inclusao_bg.data
        militar.soldado_tres = form_militar.soldado_tres.data
        militar.soldado_dois = form_militar.soldado_dois.data
        militar.soldado_um = form_militar.soldado_um.data
        militar.cabo = form_militar.cabo.data
        militar.terceiro_sgt = form_militar.terceiro_sgt.data
        militar.segundo_sgt = form_militar.segundo_sgt.data
        militar.primeiro_sgt = form_militar.primeiro_sgt.data
        militar.subtenente = form_militar.subtenente.data
        militar.segundo_tenente = form_militar.segundo_tenente.data
        militar.primeiro_tenente = form_militar.primeiro_tenente.data
        militar.cap = form_militar.cap.data
        militar.maj = form_militar.maj.data
        militar.tc = form_militar.tc.data
        militar.cel = form_militar.cel.data
        militar.funcao_gratificada_id = form_militar.funcao_gratificada_id.data
        militar.alteracao_nome_guerra = form_militar.alteracao_nome_guerra.data

        situacao2_id_raw = request.form.get('situacao2_id', '').strip()
        agregacoes2_id_raw = request.form.get('agregacoes2_id', '').strip()
        # aceita 'YYYY-MM-DD' ou 'DD/MM/YYYY'
        inicio_sit2_raw = request.form.get('inicio_situacao2', '').strip()
        fim_sit2_raw = request.form.get('fim_situacao2', '').strip()
        bg_sit2_texto = request.form.get(
            'situacao_militar_2', '').strip()  # texto da publicação

        # Normaliza IDs vazios -> None
        militar.situacao2_id = int(
            situacao2_id_raw) if situacao2_id_raw not in ('', None) else None
        militar.agregacoes2_id = int(
            agregacoes2_id_raw) if agregacoes2_id_raw not in ('', None) else None

        # Datas (usa seu parse_date que aceita 'YYYY-MM-DD' e 'DD/MM/YYYY')
        militar.inicio_situacao2 = parse_date(inicio_sit2_raw)
        militar.fim_situacao2 = parse_date(fim_sit2_raw)

        # Publicação (segue a lógica do primeiro bloco: só cria nova se mudou)
        if bg_sit2_texto:
            pub_existente = PublicacaoBg.query.filter_by(
                militar_id=militar.id, tipo_bg='situacao_militar_2'
            ).order_by(PublicacaoBg.id.desc()).first()

            if not pub_existente or pub_existente.boletim_geral != bg_sit2_texto:
                nova_pub2 = PublicacaoBg(
                    militar_id=militar.id,
                    tipo_bg='situacao_militar_2',
                    boletim_geral=bg_sit2_texto
                )
                database.session.add(nova_pub2)

        # Verifica se OBM 1 e Função 1 foram modificadas ou removidas (selecionada opção vazia)
        if form_militar.obm_ids_1.data and form_militar.funcao_ids_1.data and form_militar.obm_ids_1.data != '' and form_militar.funcao_ids_1.data != '':
            # Somente faz a alteração se houver diferença entre os dados do banco e os do formulário
            obm_id_1 = int(form_militar.obm_ids_1.data)
            funcao_id_1 = int(form_militar.funcao_ids_1.data)

            if not obm_funcao_tipo_1 or (
                    obm_id_1 != obm_funcao_tipo_1.obm_id or funcao_id_1 != obm_funcao_tipo_1.funcao_id):
                # Feche o registro antigo se houver um
                if obm_funcao_tipo_1:
                    obm_funcao_tipo_1.data_fim = datetime.now()

                nova_relacao_1 = MilitarObmFuncao(
                    militar_id=militar_id,
                    obm_id=obm_id_1,
                    funcao_id=funcao_id_1,
                    tipo=1,
                    data_criacao=datetime.now(),
                    data_fim=None
                )
                database.session.add(nova_relacao_1)
        else:
            # Se a opção 'Selecione uma opção' foi escolhida, finalize a relação atual
            if obm_funcao_tipo_1:
                obm_funcao_tipo_1.data_fim = datetime.now()

        # Verifica se OBM 2 e Função 2 foram modificadas ou removidas
        if form_militar.obm_ids_2.data and form_militar.funcao_ids_2.data and form_militar.obm_ids_2.data != '' and form_militar.funcao_ids_2.data != '':
            # Somente faz a alteração se houver diferença entre os dados do banco e os do formulário
            obm_id_2 = int(form_militar.obm_ids_2.data)
            funcao_id_2 = int(form_militar.funcao_ids_2.data)

            if not obm_funcao_tipo_2 or (
                    obm_id_2 != obm_funcao_tipo_2.obm_id or funcao_id_2 != obm_funcao_tipo_2.funcao_id):
                # Feche o registro antigo se houver um
                if obm_funcao_tipo_2:
                    obm_funcao_tipo_2.data_fim = datetime.now()

                nova_relacao_2 = MilitarObmFuncao(
                    militar_id=militar_id,
                    obm_id=obm_id_2,
                    funcao_id=funcao_id_2,
                    tipo=2,
                    data_criacao=datetime.now(),
                    data_fim=None
                )
                database.session.add(nova_relacao_2)
        else:
            # Se a opção 'Selecione uma opção' foi escolhida, finalize a relação atual
            if obm_funcao_tipo_2:
                obm_funcao_tipo_2.data_fim = datetime.now()

        for campo_bg in campos_bg:
            if hasattr(form_militar, campo_bg):
                boletim_geral_data = getattr(form_militar, campo_bg).data
                if boletim_geral_data:
                    publicacao_existente = PublicacaoBg.query.filter_by(
                        militar_id=militar.id, tipo_bg=campo_bg).first()
                    if publicacao_existente:
                        if publicacao_existente.boletim_geral != boletim_geral_data:
                            nova_publicacao = PublicacaoBg(
                                militar_id=militar.id,
                                tipo_bg=campo_bg,
                                boletim_geral=boletim_geral_data
                            )
                            database.session.add(nova_publicacao)
                    else:
                        nova_publicacao = PublicacaoBg(
                            militar_id=militar.id,
                            tipo_bg=campo_bg,
                            boletim_geral=boletim_geral_data
                        )
                        database.session.add(nova_publicacao)

                # Situação escolhida no formulário
        situacao_selecionada = Situacao.query.get(
            form_militar.situacao_id.data)
        condicao_atual = situacao_selecionada.condicao if situacao_selecionada else None

        # Garante que não fiquem registros vigentes "errados"
        encerrar_status_anteriores(militar, condicao_atual)

        if situacao_selecionada and situacao_selecionada.condicao == 'AGREGADO':
            bg_id = safe_bg_id(militar.id)  # pode ser None
            militar_agregado = MilitaresAgregados.query.filter_by(
                militar_id=militar.id).first()
            if not militar_agregado:
                militar_agregado = MilitaresAgregados(militar_id=militar.id)
                database.session.add(militar_agregado)

            militar_agregado.posto_grad_id = form_militar.posto_grad_id.data
            militar_agregado.quadro_id = form_militar.quadro_id.data
            militar_agregado.destino_id = form_militar.destino_id.data
            militar_agregado.situacao_id = situacao_selecionada.id
            militar_agregado.inicio_periodo = parse_date(
                form_militar.inicio_periodo.data)
            militar_agregado.fim_periodo_agregacao = parse_date(
                form_militar.fim_periodo.data)
            militar_agregado.publicacao_bg_id = bg_id
            militar_agregado.atualizar_status()

        # À DISPOSIÇÃO
        if situacao_selecionada and situacao_selecionada.condicao == 'À DISPOSIÇÃO':
            bg_id = safe_bg_id(militar.id)
            militar_a_disposicao = MilitaresADisposicao.query.filter_by(
                militar_id=militar.id).first()
            if not militar_a_disposicao:
                militar_a_disposicao = MilitaresADisposicao(
                    militar_id=militar.id)
                database.session.add(militar_a_disposicao)

            militar_a_disposicao.posto_grad_id = form_militar.posto_grad_id.data
            militar_a_disposicao.quadro_id = form_militar.quadro_id.data
            militar_a_disposicao.destino_id = form_militar.destino_id.data
            militar_a_disposicao.situacao_id = situacao_selecionada.id
            militar_a_disposicao.inicio_periodo = parse_date(
                form_militar.inicio_periodo.data)
            militar_a_disposicao.fim_periodo_disposicao = parse_date(
                form_militar.fim_periodo.data)
            militar_a_disposicao.publicacao_bg_id = bg_id
            militar_a_disposicao.atualizar_status()

        # LICENÇA ESPECIAL
        if situacao_selecionada and situacao_selecionada.condicao == 'LICENÇA ESPECIAL':
            bg_id = safe_bg_id(militar.id)
            militar_le = LicencaEspecial.query.filter_by(
                militar_id=militar.id).first()
            if not militar_le:
                militar_le = LicencaEspecial(militar_id=militar.id)
                database.session.add(militar_le)

            militar_le.posto_grad_id = form_militar.posto_grad_id.data
            militar_le.quadro_id = form_militar.quadro_id.data
            militar_le.destino_id = form_militar.destino_id.data
            militar_le.situacao_id = situacao_selecionada.id
            militar_le.inicio_periodo_le = parse_date(
                form_militar.inicio_periodo.data)
            militar_le.fim_periodo_le = parse_date(
                form_militar.fim_periodo.data)
            militar_le.publicacao_bg_id = bg_id
            militar_le.atualizar_status()

        # LTS
        if situacao_selecionada and situacao_selecionada.condicao == 'LTS':
            bg_id = safe_bg_id(militar.id)
            militar_lts = LicencaParaTratamentoDeSaude.query.filter_by(
                militar_id=militar.id).first()
            if not militar_lts:
                militar_lts = LicencaParaTratamentoDeSaude(
                    militar_id=militar.id)
                database.session.add(militar_lts)

            militar_lts.posto_grad_id = form_militar.posto_grad_id.data
            militar_lts.quadro_id = form_militar.quadro_id.data
            militar_lts.destino_id = form_militar.destino_id.data
            militar_lts.situacao_id = situacao_selecionada.id
            militar_lts.inicio_periodo_lts = parse_date(
                form_militar.inicio_periodo.data)
            militar_lts.fim_periodo_lts = parse_date(
                form_militar.fim_periodo.data)
            militar_lts.publicacao_bg_id = bg_id
            militar_lts.atualizar_status()

        try:
            database.session.commit()
            flash('Militar atualizado com sucesso!', 'alert-success')
            return redirect(url_for('militares', militar_id=militar_id))
        except Exception as e:
            database.session.rollback()
            flash(
                f'Erro ao atualizar o Militar. Tente novamente. {e}', 'alert-danger')
    documentos_militar = DocumentoMilitar.query.filter_by(
        militar_id=militar.id).order_by(DocumentoMilitar.criado_em.desc()).all()
    return render_template(
        'exibir_militar.html',
        form_militar=form_militar,
        militar=militar,
        documentos_militar=documentos_militar,
        bg_sit2_val=bg_sit2_val,
    )


@app.route('/inativar-militar/<int:militar_id>', methods=['POST'])
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'SUPER USER', 'DRH')
def inativar_militar(militar_id):
    militar = Militar.query.get_or_404(militar_id)

    if militar.inativo:
        flash('Este militar já está inativo.', 'alert-warning')
        return redirect(url_for('exibir_militar', militar_id=militar_id))

    militar.inativo = True

    # (Opcional, mas recomendado) Encerrar vínculos ativos de OBM/Função
    ativos = MilitarObmFuncao.query.filter_by(
        militar_id=militar_id, data_fim=None).all()
    for rel in ativos:
        rel.data_fim = datetime.now()

    # (Opcional) guardar trilha de auditoria, se tiver colunas (ver seção 3)
    militar.inativado_em = datetime.utcnow()
    militar.inativado_por_id = current_user.id
    militar.motivo_inativacao = request.form.get('motivo_inativacao') or None

    try:
        database.session.commit()
        flash('Militar inativado com sucesso.', 'alert-success')
    except Exception as e:
        database.session.rollback()
        flash(f'Erro ao inativar: {e}', 'alert-danger')

    return redirect(url_for('militares'))


@app.post("/exibir-militar/<int:militar_id>/enviar-doc")
@login_required
@checar_ocupacao('DRH', 'MAPA DA FORÇA', 'SUPER USER', 'DIRETOR DRH')
def enviar_documento_militar(militar_id):
    current_app.logger.info(
        ">>> POST enviar_documento_militar para id=%s", militar_id)

    militar = Militar.query.get_or_404(militar_id)
    file = request.files.get("doc_para_militar")

    if not file or not (file.filename or "").strip():
        flash("Selecione um arquivo.", "alert-warning")
        return redirect(url_for('exibir_militar', militar_id=militar_id))

    # calcula tamanho em bytes sem consumir o stream
    try:
        pos = file.stream.tell()
    except Exception:
        pos = 0
    try:
        file.stream.seek(0, 2)  # fim
        tamanho_bytes = file.stream.tell()
    finally:
        file.stream.seek(pos, 0)  # volta

    # sobe para o B2 (guarda só a key)
    try:
        # fica dentro de 'acumulo/'
        prefix = f"acumulo/{datetime.utcnow().year}/{militar.id}/docs"
        key = b2_upload_fileobj(file, key_prefix=prefix)
    except Exception as e:
        current_app.logger.exception("Falha ao subir doc no B2")
        flash(f"Erro ao enviar arquivo: {e}", "alert-danger")
        return redirect(url_for('exibir_militar', militar_id=militar_id))

    obs = (request.form.get("obs_para_militar") or "").strip() or None

    doc = DocumentoMilitar(
        militar_id=militar.id,
        destinatario_cpf=militar.cpf,
        nome_original=file.filename,
        content_type=file.mimetype or "application/octet-stream",
        tamanho_bytes=tamanho_bytes,
        object_key=key,
        criado_por_user_id=current_user.id,
        observacao=obs
    )

    try:
        database.session.add(doc)
        database.session.commit()
    except Exception as e:
        database.session.rollback()
        current_app.logger.exception("Falha ao gravar DocumentoMilitar")
        flash(f"Erro ao salvar no banco: {e}", "alert-danger")
        return redirect(url_for('exibir_militar', militar_id=militar_id))

    flash("Documento disponibilizado para o militar.", "alert-success")
    return redirect(url_for('exibir_militar', militar_id=militar_id))


@app.post("/documentos/<int:doc_id>/revogar")
@login_required
@checar_ocupacao('DRH', 'MAPA DA FORÇA', 'SUPER USER', 'DIRETOR DRH')
def revogar_documento_militar(doc_id):
    doc = DocumentoMilitar.query.get_or_404(doc_id)
    if doc.baixado_em:
        flash("Documento já foi baixado; não é possível revogar.", "alert-warning")
        return redirect(url_for('exibir_militar', militar_id=doc.militar_id))

    try:
        b2_delete_all_versions(doc.object_key)
    except Exception:
        current_app.logger.exception("Falha ao remover do B2 (revogar)")

    database.session.delete(doc)
    database.session.commit()
    flash("Documento revogado e removido do Backblaze.", "alert-success")
    return redirect(url_for('exibir_militar', militar_id=doc.militar_id))


TZ_MANAUS = pytz.timezone('America/Manaus')


def _to_manaus(dt):
    """Converte dt (naive/aware/None) para aware em America/Manaus."""
    if dt is None:
        return None
    if dt.tzinfo is None:
        # tratamos como horário local salvo sem tz
        return TZ_MANAUS.localize(dt)
    return dt.astimezone(TZ_MANAUS)            # converte para Manaus


@app.get("/meus-documentos")
@login_required
def meus_documentos():
    pendentes = (DocumentoMilitar.query
                 .filter_by(destinatario_cpf=current_user.cpf)
                 .filter(DocumentoMilitar.baixado_em.is_(None))
                 .order_by(DocumentoMilitar.criado_em.desc())
                 .all())

    baixados = (DocumentoMilitar.query
                .filter_by(destinatario_cpf=current_user.cpf)
                .filter(DocumentoMilitar.baixado_em.isnot(None))
                .order_by(DocumentoMilitar.baixado_em.desc())
                .limit(50)
                .all())

    # primeira visita
    show_intro = (request.cookies.get("meus_docs_intro_seen") != "1")

    NOVO_LIMITE_DIAS = 3
    now_mao = datetime.now(TZ_MANAUS)
    novo_limite = now_mao - timedelta(days=NOVO_LIMITE_DIAS)

    # marca cada doc como "novo" (conversão robusta para Manaus)
    for d in [*pendentes, *baixados]:
        criado_local = _to_manaus(d.criado_em)
        d.is_new = bool(criado_local and criado_local >= novo_limite)

    resp = make_response(render_template(
        "meus_documentos.html",
        pendentes=pendentes,
        baixados=baixados,
        show_intro=show_intro,
        novo_limite_dias=NOVO_LIMITE_DIAS,  # só para tooltip/texto, se quiser
    ))
    if show_intro:
        resp.set_cookie("meus_docs_intro_seen", "1",
                        max_age=60*60*24*365, httponly=False, samesite="Lax")
    return resp


@app.get("/documentos/<int:doc_id>/download")
@login_required
def download_documento(doc_id):
    doc = DocumentoMilitar.query.get_or_404(doc_id)
    if not _pode_pegar_doc(doc):
        abort(403)
    if doc.baixado_em:
        flash("Este documento já foi baixado e não está mais disponível.",
              "alert-warning")
        return redirect(url_for('meus_documentos'))

    s3 = b2_client()
    try:
        obj = s3.get_object(Bucket=b2_bucket_name(), Key=doc.object_key)
        body = obj["Body"]
    except Exception:
        current_app.logger.exception("Falha ao abrir objeto no B2")
        abort(404)

    def stream():
        try:
            for chunk in iter(lambda: body.read(8192), b""):
                yield chunk
        finally:
            try:
                body.close()
            except Exception:
                pass

    resp = Response(stream_with_context(stream()),
                    mimetype=doc.content_type or "application/octet-stream")
    resp.headers["Content-Disposition"] = f'attachment; filename="{doc.nome_original}"'

    # >>> pegue o app e dados que você precisa ANTES de fechar o contexto
    app = current_app._get_current_object()
    key = doc.object_key
    doc_id = doc.id

    @resp.call_on_close
    def _cleanup(_app=app, _key=key, _doc_id=doc_id):
        # reabre contexto da aplicação
        with _app.app_context():
            try:
                b2_delete_all_versions(_key)  # precisa permissão de delete
            except Exception:
                _app.logger.exception(
                    "Falha ao deletar objeto no B2 após download")
            try:
                # atualiza sem manter a instância anexada
                DocumentoMilitar.query.filter_by(id=_doc_id)\
                    .update({"baixado_em": datetime.utcnow()})
                database.session.commit()
            except Exception:
                database.session.rollback()
                _app.logger.exception("Falha ao marcar doc como baixado")

    return resp


@app.get("/documentos/<int:doc_id>/status")
@login_required
def status_documento(doc_id):
    doc = DocumentoMilitar.query.get_or_404(doc_id)
    if not _pode_pegar_doc(doc):
        abort(403)
    return {"baixado": bool(doc.baixado_em)}


@app.route("/militares", methods=['GET'])
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'MAPA DA FORÇA', 'DRH', 'SUPER USER', 'DIRETOR DRH')
def militares():
    f = FormFiltroMilitar()

    # choices
    f.obm_id_1.choices = [(o.id, o.sigla)
                          for o in Obm.query.order_by(Obm.sigla).all()]
    f.funcao_id.choices = [(x.id, x.ocupacao)
                           for x in Funcao.query.order_by(Funcao.ocupacao).all()]
    f.posto_grad_id.choices = [
        (p.id, p.sigla) for p in PostoGrad.query.order_by(PostoGrad.sigla.asc()).all()]
    f.quadro_id.choices = [(q.id, q.quadro)
                           for q in Quadro.query.order_by(Quadro.quadro).all()]
    f.especialidade_id.choices = [
        (e.id, e.ocupacao) for e in Especialidade.query.order_by(Especialidade.ocupacao).all()]
    f.localidade_id.choices = [
        (l.id, l.sigla) for l in Localidade.query.order_by(Localidade.sigla).all()]
    f.situacao_id.choices = [(s.id, s.condicao)
                             for s in Situacao.query.order_by(Situacao.condicao).all()]

    # paginação & busca
    page = request.args.get('page', 1, type=int)
    search = (request.args.get('search') or '').strip()

    # MULTI filtros (listas)
    obm_ids = request.args.getlist('obm_ids', type=int)
    funcao_ids = request.args.getlist('funcao_ids', type=int)
    posto_grad_ids = request.args.getlist('posto_grad_ids', type=int)
    quadro_ids = request.args.getlist('quadro_ids', type=int)
    especialidade_ids = request.args.getlist('especialidade_ids', type=int)
    localidade_ids = request.args.getlist('localidade_ids', type=int)
    situacao_ids = request.args.getlist('situacao_ids', type=int)

    sexo_filtro = (request.args.get('sexo') or '').strip().upper()

    # base
    query = (Militar.query
             .options(
                 selectinload(Militar.obm_funcoes).selectinload(
                     MilitarObmFuncao.obm),
                 selectinload(Militar.obm_funcoes).selectinload(
                     MilitarObmFuncao.funcao),
                 selectinload(Militar.posto_grad),
                 selectinload(Militar.quadro)
             )
             .filter(Militar.inativo.is_(False))
             )

    if search:
        like = f"%{search}%"
        query = query.filter(or_(
            Militar.nome_completo.ilike(like),
            Militar.nome_guerra.ilike(like),
            Militar.cpf.ilike(like),
            Militar.rg.ilike(like),
            Militar.matricula.ilike(like),
        ))

    # aplica filtros simples (FK direta)
    if posto_grad_ids:
        query = query.filter(Militar.posto_grad_id.in_(posto_grad_ids))
    if quadro_ids:
        query = query.filter(Militar.quadro_id.in_(quadro_ids))
    if especialidade_ids:
        query = query.filter(Militar.especialidade_id.in_(especialidade_ids))
    if localidade_ids:
        query = query.filter(Militar.localidade_id.in_(localidade_ids))
    if situacao_ids:
        query = query.filter(Militar.situacao_id.in_(situacao_ids))

    # filtros por OBM/Função ativos via join
    if obm_ids or funcao_ids:
        query = (query
                 .join(MilitarObmFuncao, MilitarObmFuncao.militar_id == Militar.id)
                 .filter(MilitarObmFuncao.data_fim.is_(None)))
        if obm_ids:
            query = query.filter(MilitarObmFuncao.obm_id.in_(obm_ids))
        if funcao_ids:
            query = query.filter(MilitarObmFuncao.funcao_id.in_(funcao_ids))

        query = query.distinct()

    sexo_norm = func.lower(func.trim(Militar.sexo))

    if sexo_filtro == 'M':
        query = query.filter(
            Militar.sexo.isnot(None),
            sexo_norm.like('m%')
        )
    elif sexo_filtro == 'F':
        query = query.filter(
            Militar.sexo.isnot(None),
            sexo_norm.like('f%')
        )

    # paginação
    per_page = 100
    query = query.order_by(Militar.nome_completo.asc())
    militares_paginados = query.paginate(page=page, per_page=per_page)

    def fmt_cpf(cpf):
        d = re.sub(r'\D', '', cpf or '')
        return f"{d[:3]}.{d[3:6]}.{d[6:9]}-{d[9:11]}" if len(d) == 11 else (cpf or '')

    # Preparar dados para renderização
    militares = []
    for militar in militares_paginados.items:
        # Selecionar apenas OBMs e funções ativas (data_fim is None)
        obm_funcoes_ativas = [
            of for of in militar.obm_funcoes if of.data_fim is None
        ]

        # Ordenar OBMs e funções ativas por data_criacao (opcional)
        obm_funcoes_ativas = sorted(
            obm_funcoes_ativas,
            key=lambda of: of.data_criacao,
            reverse=True
        )

        # Extrair OBMs e Funções ativas
        obms_recentes = [
            of.obm.sigla if of.obm else 'OBM não encontrada'
            for of in obm_funcoes_ativas
        ]
        funcoes_recentes = [
            of.funcao.ocupacao if of.funcao else 'Função não encontrada'
            for of in obm_funcoes_ativas
        ]

        militares.append({
            'id': militar.id,
            'nome_completo': militar.nome_completo,
            'nome_guerra': militar.nome_guerra,
            'cpf': militar.cpf,
            'cpf_fmt': fmt_cpf(militar.cpf),
            'rg': militar.rg,
            'obms': obms_recentes,  # Lista com as OBMs ativas
            'funcoes': funcoes_recentes,  # Lista com as funções ativas
            'posto_grad': militar.posto_grad.sigla if militar.posto_grad else '',
            'quadro': militar.quadro.quadro if militar.quadro else '',
            'matricula': militar.matricula,
        })

    return render_template(
        'militares.html',
        militares=militares,
        form_militar=f,              # substitui aqui
        page=page,
        has_next=militares_paginados.has_next,
        has_prev=militares_paginados.has_prev,
        next_page=militares_paginados.next_num,
        prev_page=militares_paginados.prev_num,
        pages=militares_paginados.pages,
        total=militares_paginados.total,
        start=(page-1)*per_page+1 if militares_paginados.total else 0,
        end=min(page*per_page, militares_paginados.total),
        has_novo_militar=('novo_militar' in current_app.view_functions),
    )


@app.route("/militares-inativos", methods=['GET'])
@login_required
def militares_inativos():
    try:
        page = request.args.get('page', 1, type=int)
        search = request.args.get('search', '', type=str)

        query = Militar.query.options(
            joinedload(Militar.posto_grad),
            joinedload(Militar.quadro),
            joinedload(Militar.especialidade),
            joinedload(Militar.localidade),
            joinedload(Militar.situacao),
            joinedload(Militar.obm_funcoes)
        ).filter(Militar.situacao.has(Situacao.condicao.in_(['RESERVA', 'INATIVO'])))

        if search:
            query = query.filter(Militar.nome_completo.ilike(f"%{search}%"))

        militares_inativos = query.order_by(
            Militar.nome_completo.asc()).paginate(page=page, per_page=100)

        return render_template(
            'militares_inativos.html',
            militares=militares_inativos.items,
            page=page,
            has_next=militares_inativos.has_next,
            has_prev=militares_inativos.has_prev,
            next_page=militares_inativos.next_num,
            prev_page=militares_inativos.prev_num
        )

    except Exception as e:
        app.logger.error(f"Erro ao processar a requisição: {str(e)}")
        return jsonify({'error': 'Ocorreu um erro ao processar a requisição.', 'details': str(e)}), 500


@app.route('/tabela-militares', methods=['GET', 'POST'])
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'MAPA DA FORÇA', 'SUPER USER', 'DRH', 'DIRETOR DRH')
def tabela_militares():
    today = date.today()
    try:
        page = request.args.get('page', 1, type=int)
        search = request.args.get('search', '', type=str)
        query = Militar.query.options(
            joinedload(Militar.posto_grad),
            joinedload(Militar.quadro),
            joinedload(Militar.especialidade),
            joinedload(Militar.localidade),
            joinedload(Militar.situacao),
            joinedload(Militar.obm_funcoes),
            joinedload(Militar.destino)
        ).filter(Militar.inativo.is_(False))

        total_militares = query.count()

        # Filtro de busca por nome
        if search:
            query = query.filter(Militar.nome_completo.ilike(f"%{search}%"))

        sexo_filtro = request.args.get('sexo', '', type=str).strip().upper()

        vals = request.values  # une args + form
        obm_ids = vals.getlist('obm_ids', type=int)
        funcao_ids = vals.getlist('funcao_ids', type=int)
        posto_grad_ids = vals.getlist('posto_grad_ids', type=int)
        quadro_ids = vals.getlist('quadro_ids', type=int)
        especialidade_ids = vals.getlist('especialidade_ids', type=int)
        localidade_ids = vals.getlist('localidade_ids', type=int)
        situacao_ids = vals.getlist('situacao_ids', type=int)
        sexo_norm = func.lower(func.trim(Militar.sexo))
        # FK diretas
        if posto_grad_ids:
            query = query.filter(Militar.posto_grad_id.in_(posto_grad_ids))
        if quadro_ids:
            query = query.filter(Militar.quadro_id.in_(quadro_ids))
        if especialidade_ids:
            query = query.filter(
                Militar.especialidade_id.in_(especialidade_ids))
        if localidade_ids:
            query = query.filter(Militar.localidade_id.in_(localidade_ids))
        if situacao_ids:
            query = query.filter(Militar.situacao_id.in_(situacao_ids))

        # OBM/Função ativas
        if obm_ids or funcao_ids:
            mo = (database.session.query(MilitarObmFuncao.militar_id)
                  .filter(MilitarObmFuncao.data_fim.is_(None)))
            if obm_ids:
                mo = mo.filter(MilitarObmFuncao.obm_id.in_(obm_ids))
            if funcao_ids:
                mo = mo.filter(MilitarObmFuncao.funcao_id.in_(funcao_ids))

            mo = mo.distinct()  # DISTINCT "puro" (sem ON)

            # restringe os militares aos que têm OBM/Função ativa
            query = query.filter(Militar.id.in_(mo))

        if sexo_filtro == 'M':
            query = query.filter(
                Militar.sexo.isnot(None),
                sexo_norm.like('m%')   # "m", "masculino", "MASCULINO", etc
            )
        elif sexo_filtro == 'F':
            query = query.filter(
                Militar.sexo.isnot(None),
                sexo_norm.like('f%')   # "f", "feminino", etc
            )

        base_filtrada = query.order_by(None)

        filtrados_sq = (
            base_filtrada
            .with_entities(Militar.id)
            .distinct()        # em PG vira DISTINCT ON (militar.id)
            # <-- essencial pra não “vazar” ORDER BY nome_completo
            .order_by(None)
            .subquery()
        )

        # contagem de agregados dentre os filtrados
        agregados_ids = [
            x[0]
            for x in (
                database.session.query(MilitaresAgregados.militar_id)
                .join(filtrados_sq, MilitaresAgregados.militar_id == filtrados_sq.c.id)
                .filter(
                    MilitaresAgregados.inicio_periodo <= today,
                    or_(
                        MilitaresAgregados.fim_periodo_agregacao == None,
                        MilitaresAgregados.fim_periodo_agregacao >= today,
                    )
                )
                .distinct()
                .all()
            )
        ]
        agregados_count = len(agregados_ids)

        # contagem de "à disposição" dentre os filtrados
        adisposicao_ids = [
            x[0]
            for x in (
                database.session.query(MilitaresADisposicao.militar_id)
                .join(filtrados_sq, MilitaresADisposicao.militar_id == filtrados_sq.c.id)
                .filter(
                    MilitaresADisposicao.inicio_periodo <= today,
                    or_(
                        MilitaresADisposicao.fim_periodo_disposicao == None,
                        MilitaresADisposicao.fim_periodo_disposicao >= today,
                    )
                )
                .distinct()
                .all()
            )
        ]
        adisposicao_count = len(adisposicao_ids)

        militares_filtrados = (
            base_filtrada               # mesma base sem ORDER BY no subquery
            .order_by(Militar.nome_completo.asc())
            .all()
        )

        query = query.order_by(Militar.nome_completo.asc())

        # Retorna todos os resultados
        militares_filtrados = query.all()

        # Preparar dados para o template
        militares_filtrados_data = []
        for militar in militares_filtrados:
            obm_funcoes_ativas = [
                {
                    'obm': of.obm.sigla if of.obm else 'OBM não encontrada',
                    'funcao': of.funcao.ocupacao if of.funcao else 'Função não encontrada'
                }
                for of in militar.obm_funcoes if of.data_fim is None
            ]

            destino_txt = 'N/A'
            try:
                if getattr(militar, 'destino', None):
                    destino_txt = getattr(
                        militar.destino, 'local', None) or 'N/A'
                elif getattr(militar, 'destino_id', None):
                    d = Destino.query.get(militar.destino_id)
                    destino_txt = getattr(d, 'local', None) or str(
                        militar.destino_id)
            except Exception:
                pass

            inclusao_fmt = militar.inclusao.strftime(
                '%d/%m/%Y') if militar.inclusao else 'N/A'

            if militar.id in adisposicao_ids:
                situacao_exibe = 'À DISPOSIÇÃO'
            elif militar.id in agregados_ids:
                situacao_exibe = 'AGREGADO'
            else:
                situacao_exibe = militar.situacao.condicao if militar.situacao else 'N/A'

            sexo_raw = (militar.sexo or '').strip()
            s = sexo_raw.lower()

            if s.startswith('m'):
                sexo_exibe = 'Masculino'
            elif s.startswith('f'):
                sexo_exibe = 'Feminino'
            elif sexo_raw:
                sexo_exibe = sexo_raw
            else:
                sexo_exibe = 'N/A'

            militares_filtrados_data.append({
                'id': militar.id,
                'nome_completo': militar.nome_completo,
                'nome_guerra': militar.nome_guerra,
                'sexo': sexo_exibe,  # <-- NOVO
                'cpf': militar.cpf,
                'rg': militar.rg,
                'matricula': militar.matricula,
                'posto_grad': militar.posto_grad.sigla if militar.posto_grad else 'N/A',
                'quadro': militar.quadro.quadro if militar.quadro else 'N/A',
                'especialidade': militar.especialidade.ocupacao if militar.especialidade else 'N/A',
                'localidade': militar.localidade.sigla if militar.localidade else 'N/A',
                'situacao': situacao_exibe,
                'destino': destino_txt,
                'inclusao': inclusao_fmt,
                'obms': [item['obm'] for item in obm_funcoes_ativas],
                'funcoes': [item['funcao'] for item in obm_funcoes_ativas],
            })

        return render_template(
            'relacao_militares.html',
            militares=militares_filtrados_data,
            total_militares=total_militares,
            militares_filtrados_count=len(militares_filtrados_data),
            agregados_count=agregados_count,
            adisposicao_count=adisposicao_count,
            adisposicao_ids=adisposicao_ids,   # opcional
            agregados_ids=agregados_ids,       # opcional
        )

    except Exception as e:
        app.logger.error(f"Erro ao processar a requisição: {str(e)}")
        return jsonify({'error': 'Ocorreu um erro ao processar a requisição.', 'details': str(e)}), 500


@app.route("/export-excel", methods=["GET"])
@login_required
def export_excel():
    try:
        if 'militares_filtrados' not in session:
            return jsonify({'error': 'Nenhum dado filtrado disponível para download.'}), 400

        dados = session['militares_filtrados']
        df = pd.DataFrame(dados)

        # Definindo a ordem das colunas (ajuste conforme necessário)
        colunas_em_ordem = [
            'Nome Completo',
            'Nome de Guerra',
            'CPF',
            'RG',
            'Matrícula',
            'Posto/Graduação',
            'Quadro',
            'Especialidade',
            'Localidade',
            'Situação',
            'OBM 1',
            'Função 1',
            'OBM 2',
            'Função 2'
        ]

        # Reordenando o DataFrame de acordo com a lista definida
        df = df[colunas_em_ordem]

        # Salva para um arquivo Excel na memória
        output = BytesIO()
        writer = pd.ExcelWriter(output, engine='xlsxwriter')

        # Converter o DataFrame para um Excel
        df.to_excel(writer, index=False, sheet_name='Militares')

        # Acessar o workbook e a planilha
        workbook = writer.book
        worksheet = writer.sheets['Militares']

        # Definindo a largura das colunas
        column_widths = {
            'A': 25,  # Nome Completo
            'B': 25,  # Nome de Guerra
            'C': 15,  # CPF
            'D': 15,  # RG
            'E': 30,  # Matrícula
            'F': 20,  # Posto/Graduação
            'G': 30,  # Quadro
            'H': 30,  # 'Especialidade'
            'I': 20,  # Localidade
            'J': 20,  # Situação
            'K': 20,  # OBM 1
            'L': 20,  # Função 1
            'M': 20,  # OBM 2
            'N': 20  # Função 2
        }

        for col, width in column_widths.items():
            worksheet.set_column(f'{col}:{col}', width)

        # Estilizando o cabeçalho
        header_format = workbook.add_format({
            'bg_color': '#00008b',  # Cor de fundo azul
            'font_color': '#FFFFFF',  # Cor da fonte branca
            'bold': True,  # Texto em negrito
            'border': 1  # Bordas
        })

        # Aplicar o formato ao cabeçalho
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)

        writer.close()
        output.seek(0)

        return send_file(output, download_name="militares_filtrados.xlsx", as_attachment=True)

    except Exception as e:
        print(f"Erro ao processar a requisição: {e}")
        return jsonify({'error': 'Ocorreu um erro ao processar a requisição.'}), 500


@app.route("/militares-a-disposicao")
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'MAPA DA FORÇA', 'DRH', 'SUPER USER', 'DIRETOR DRH')
def militares_a_disposicao():
    militares_a_disposicao = (
        MilitaresADisposicao.query
        .join(Militar, MilitaresADisposicao.militar_id == Militar.id)
        .filter(Militar.inativo.is_(False))
        .all()
    )
    return render_template('militares_a_disposicao.html', militares=militares_a_disposicao)


@app.route("/militares-agregados")
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'MAPA DA FORÇA', 'DRH', 'SUPER USER', 'DIRETOR DRH')
def militares_agregados():
    militares_agregados = (
        MilitaresAgregados.query
        .join(Militar, MilitaresAgregados.militar_id == Militar.id)
        .filter(Militar.inativo.is_(False))
        .all()
    )
    return render_template('militares_agregados.html', militares=militares_agregados)


@app.route("/licenca-especial")
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'MAPA DA FORÇA', 'DRH', 'SUPER USER', 'DIRETOR DRH')
def licenca_especial():
    militares_le = (
        LicencaEspecial.query
        .join(Militar, LicencaEspecial.militar_id == Militar.id)
        .filter(Militar.inativo.is_(False))
        .all()
    )
    return render_template('licenca_especial.html', militares_le=militares_le)


@app.route("/licenca-para-tratamento-de-saude")
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'MAPA DA FORÇA', 'DRH', 'SUPER USER', 'DIRETOR DRH')
def lts():
    militares_lts = (
        LicencaParaTratamentoDeSaude.query
        .join(Militar, LicencaParaTratamentoDeSaude.militar_id == Militar.id)
        .filter(Militar.inativo.is_(False))
        .all()
    )
    return render_template('licenca_para_tratamento_de_saude.html', militares_lts=militares_lts)


@app.route("/exportar-excel/<string:tabela>")
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'MAPA DA FORÇA', 'DRH', 'SUPER USER', 'DIRETOR DRH')
def exportar_excel(tabela):
    # Mapeamento das tabelas para consultas
    tabela_mapping = {
        'militares_agregados': {
            'model': MilitaresAgregados,
            'columns': [
                'posto_grad.sigla', 'quadro.quadro', 'militar.nome_completo',
                'destino.local', 'situacao.condicao', 'inicio_periodo',
                'fim_periodo_agregacao', 'status', 'publicacao_bg.boletim_geral'
            ]
        },
        'militares_a_disposicao': {
            'model': MilitaresADisposicao,
            'columns': [
                'posto_grad.sigla', 'quadro.quadro', 'militar.nome_completo',
                'destino.local', 'situacao.condicao', 'inicio_periodo',
                'fim_periodo_disposicao', 'status', 'publicacao_bg.boletim_geral'
            ]
        }
    }

    # Verifica se a tabela especificada está no mapeamento
    if tabela not in tabela_mapping:
        return jsonify({'error': 'Tabela não encontrada'}), 404

    # Obter modelo e colunas da tabela especificada
    tabela_info = tabela_mapping[tabela]
    modelo = tabela_info['model']
    colunas = tabela_info['columns']

    # Consultar os dados da tabela
    dados = modelo.query.all()

    # Construir dados para o DataFrame do Pandas
    export_data = []
    for item in dados:
        export_row = {
            'Posto/Graduação': getattr(item.posto_grad, 'sigla', 'N/A'),
            'Quadro': getattr(item.quadro, 'quadro', 'N/A'),
            'Nome Completo': getattr(item.militar, 'nome_completo', 'N/A'),
            'Destino': getattr(item.destino, 'local', 'N/A'),
            'Situação': getattr(item.situacao, 'condicao', 'N/A'),
            'A contar de': item.inicio_periodo.strftime('%d/%m/%Y') if item.inicio_periodo else 'N/A',
            'Término': (
                item.fim_periodo_agregacao.strftime('%d/%m/%Y') if tabela == 'militares_agregados' else
                item.fim_periodo_disposicao.strftime(
                    '%d/%m/%Y') if item.fim_periodo_disposicao else 'N/A'
            ),
            'Status': item.status,
            'Documento Autorizador': getattr(item.publicacao_bg, 'boletim_geral', 'N/A')
        }
        export_data.append(export_row)

    # Criar DataFrame
    df = pd.DataFrame(export_data)

    # Gerar arquivo Excel em memória
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Dados')
        workbook = writer.book
        worksheet = writer.sheets['Dados']

        # Definir formatos personalizados
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#1F4E78',  # Azul escuro
            'font_color': 'white',
            'align': 'center',
            'valign': 'vcenter'
        })

        cell_centered_format = workbook.add_format(
            {'align': 'center', 'valign': 'vcenter'})
        cell_left_format = workbook.add_format(
            {'align': 'left', 'valign': 'vcenter'})

        # Ajustar largura das colunas
        column_widths = {
            'Posto/Graduação': 15,
            'Quadro': 15,
            'Nome Completo': 30,
            'Destino': 20,
            'Situação': 15,
            'A contar de': 12,
            'Término': 12,
            'Status': 15,
            'Documento Autorizador': 30
        }

        # Aplicar largura das colunas e formatação
        for col_num, column_title in enumerate(df.columns):
            width = column_widths.get(column_title, 15)  # Valor padrão 15
            worksheet.set_column(col_num, col_num, width)

            # Aplicar estilo ao cabeçalho
            worksheet.write(0, col_num, column_title, header_format)

            # Aplicar centralização condicional para colunas específicas
            if column_title in ['Nome Completo', 'Documento Autorizador']:
                worksheet.set_column(col_num, col_num, width, cell_left_format)
            else:
                worksheet.set_column(
                    col_num, col_num, width, cell_centered_format)

    output.seek(0)

    # Enviar arquivo Excel para download
    nome_arquivo = f'{tabela}.xlsx'
    return send_file(output, as_attachment=True, download_name=nome_arquivo,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route("/api/usuarios", methods=["GET"])
@login_required
@checar_ocupacao('DIRETOR', 'SUPER USER')
def api_usuarios():
    # Parâmetros vindos do DataTables
    draw = int(request.args.get('draw', 1))
    start = int(request.args.get('start', 0))
    length = int(request.args.get('length', 10))
    search_value = request.args.get('search[value]', '')

    query = User.query.join(FuncaoUser, User.funcao_user_id == FuncaoUser.id)

    if search_value:
        query = query.filter(
            database.or_(
                User.nome.ilike(f"%{search_value}%"),
                User.cpf.ilike(f"%{search_value}%"),
                FuncaoUser.ocupacao.ilike(f"%{search_value}%"),
            )
        )

    total_records = User.query.count()
    filtered_records = query.count()

    usuarios = query.add_columns(
        User.id, User.nome, User.cpf,
        FuncaoUser.ocupacao.label('funcao_ocupacao')
    ).offset(start).limit(length).all()

    data = []
    for usuario in usuarios:
        data.append([
            usuario.nome,
            usuario.cpf,
            usuario.funcao_ocupacao,
            f"""
            <a href="{url_for('exibir_usuario', id_usuario=usuario.id)}" class="btn btn-sm btn-primary"><i class="bi bi-eye-fill"></i></a>
            <a href="{url_for('exibir_usuario', id_usuario=usuario.id)}" class="btn btn-sm btn-warning"><i class="bi bi-pencil-fill"></i></a>
            <a href="{url_for('excluir_usuario', usuario_id=usuario.id)}" class="btn btn-sm btn-danger"
            onclick="return confirm('Tem certeza que deseja excluir este usuário?')">
            <i class="bi bi-trash-fill"></i></a>
            """
        ])

    return {
        "draw": draw,
        "recordsTotal": total_records,
        "recordsFiltered": filtered_records,
        "data": data
    }


@app.route("/usuarios", methods=['GET'])
@login_required
@checar_ocupacao('DIRETOR', 'SUPER USER')
def usuarios():
    return render_template('usuarios.html')


@app.route('/usuario/<int:id_usuario>', methods=['GET', 'POST'])
@login_required
@checar_ocupacao('DIRETOR', 'SUPER USER')
def exibir_usuario(id_usuario):
    # Carregar o usuário diretamente
    usuario = User.query.get_or_404(id_usuario)

    # Carregar informações extras sobre a função com join para exibição na tabela
    usuario_info = User.query \
        .join(FuncaoUser, User.funcao_user_id == FuncaoUser.id) \
        .add_columns(User.nome, User.cpf, User.id, User.email,
                     FuncaoUser.ocupacao.label('funcao_ocupacao')) \
        .filter(User.id == id_usuario) \
        .first_or_404()

    form = FormCriarUsuario(obj=usuario)
    form.current_user_id = id_usuario
    form.funcao_user_id.choices = [
        (funcao.id, funcao.ocupacao) for funcao in FuncaoUser.query.all()]
    form.obm_id_1.choices = [('', '-- Selecione uma opção --')] + [(obm.id, obm.sigla) for obm in
                                                                   Obm.query.all()]
    form.obm_id_2.choices = [('', '-- Selecione uma opção --')] + [(obm.id, obm.sigla) for obm in
                                                                   Obm.query.all()]
    form.localidade_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(localidade.id, localidade.sigla) for localidade in
         Localidade.query.all()]

    if form.validate_on_submit():
        usuario.nome = form.nome.data
        usuario.email = form.email.data
        usuario.cpf = form.cpf.data
        usuario.funcao_user_id = form.funcao_user_id.data
        usuario.obm_id_1 = form.obm_id_1.data
        usuario.obm_id_2 = form.obm_id_2.data
        usuario.localidade_id = form.localidade_id.data

        if form.senha.data:
            usuario.senha = bcrypt.generate_password_hash(
                form.senha.data).decode('utf-8')

        try:
            database.session.commit()
            flash('Usuário atualizado com sucesso!', 'alert-success')
            return redirect(url_for('usuarios', id_usuario=id_usuario))
        except Exception as e:
            database.session.rollback()
            flash(
                f'Erro ao atualizar o usuário. Tente novamente. {e}', 'alert-danger')

    return render_template('usuario_detalhes.html', usuario=usuario_info, form=form)


@app.route('/perfil/<int:id_usuario>', methods=['GET', 'POST'])
@login_required
def perfil(id_usuario):
    # Verificar se o usuário logado está acessando seu próprio perfil
    if current_user.id != id_usuario:
        flash('Você não tem permissão para acessar este perfil.', 'alert-danger')
        return redirect(url_for('home'))

    usuario = User.query.get_or_404(id_usuario)

    usuario_info = User.query \
        .join(FuncaoUser, User.funcao_user_id == FuncaoUser.id) \
        .add_columns(User.nome, User.cpf, User.id, User.email,
                     FuncaoUser.ocupacao.label('funcao_ocupacao')) \
        .filter(User.id == id_usuario) \
        .first_or_404()

    form = FormCriarUsuario(obj=usuario)
    form.current_user_id = id_usuario
    form.funcao_user_id.choices = [
        (funcao.id, funcao.ocupacao) for funcao in FuncaoUser.query.all()]
    form.obm_id_1.choices = [('', '-- Selecione uma opção --')] + [(obm.id, obm.sigla) for obm in
                                                                   Obm.query.all()]
    form.obm_id_2.choices = [('', '-- Selecione uma opção --')] + [(obm.id, obm.sigla) for obm in
                                                                   Obm.query.all()]
    form.localidade_id.choices = [
        ('', '-- Selecione uma opção --')
    ] + [(localidade.id, localidade.sigla) for localidade in
         Localidade.query.all()]

    if form.validate_on_submit():
        usuario.nome = form.nome.data
        usuario.email = form.email.data
        usuario.cpf = form.cpf.data
        usuario.funcao_user_id = form.funcao_user_id.data
        usuario.obm_id_1 = form.obm_id_1.data
        usuario.obm_id_2 = form.obm_id_2.data
        usuario.localidade_id = form.localidade_id.data

        if form.senha.data:
            usuario.senha = bcrypt.generate_password_hash(
                form.senha.data).decode('utf-8')

        try:
            database.session.commit()
            flash('Usuário atualizado com sucesso!', 'alert-success')
            return redirect(url_for('home', id_usuario=id_usuario))
        except Exception as e:
            database.session.rollback()
            flash(
                f'Erro ao atualizar o usuário. Tente novamente. {e}', 'alert-danger')

    return render_template('perfil.html', usuario=usuario_info, form=form)


@app.route("/exportar-pafs/<string:tabela>")
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'MAPA DA FORÇA', 'DRH', 'SUPER USER', 'DIRETOR DRH')
def exportar_pafs(tabela):
    if tabela != "pafs":
        return "Tabela inválida", 400

        # Consulta os dados
    militares_pafs = (
        database.session.query(Militar, Paf)
        .outerjoin(Paf, Paf.militar_id == Militar.id)
        .all()
    )

    # Criação do arquivo Excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Pafs e Militares"

    # Cabeçalhos com filtros
    colunas = [
        "Posto/Grad", "Nome", "Matrícula", "Quadro", "Mês Usufruto",
        "Qtd. Dias 1º Período", "1º Período de Férias", "Fim 1º Período",
        "Qtd. Dias 2º Período", "2º Período de Férias", "Fim 2º Período",
        "Qtd. Dias 3º Período", "3º Período de Férias", "Fim 3º Período"
    ]
    ws.append(colunas)

    # Estilo do cabeçalho
    for col_num, col_name in enumerate(colunas, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = col_name
        cell.font = Font(bold=True)
    ws.auto_filter.ref = f"A1:{get_column_letter(len(colunas))}1"

    # Adiciona os dados
    for militar, paf in militares_pafs:
        ws.append([
            militar.posto_grad.sigla if militar.posto_grad else "",
            militar.nome_completo,
            militar.matricula,
            militar.quadro.quadro if militar.quadro else "",
            paf.mes_usufruto if paf else "",
            paf.qtd_dias_primeiro_periodo if paf else "",
            paf.primeiro_periodo_ferias.strftime(
                "%d/%m/%Y") if paf and paf.primeiro_periodo_ferias else "",
            paf.fim_primeiro_periodo.strftime(
                "%d/%m/%Y") if paf and paf.fim_primeiro_periodo else "",
            paf.qtd_dias_segundo_periodo if paf else "",
            paf.segundo_periodo_ferias.strftime(
                "%d/%m/%Y") if paf and paf.segundo_periodo_ferias else "",
            paf.fim_segundo_periodo.strftime(
                "%d/%m/%Y") if paf and paf.fim_segundo_periodo else "",
            paf.qtd_dias_terceiro_periodo if paf else "",
            paf.terceiro_periodo_ferias.strftime(
                "%d/%m/%Y") if paf and paf.terceiro_periodo_ferias else "",
            paf.fim_terceiro_periodo.strftime(
                "%d/%m/%Y") if paf and paf.fim_terceiro_periodo else "",
        ])

    # Ajusta largura das colunas
    for col_num, _ in enumerate(colunas, 1):
        ws.column_dimensions[get_column_letter(col_num)].auto_size = True

    # Salva em um arquivo na memória
    output = BytesIO()
    wb.save(output)
    output.seek(0)

    # Prepara resposta HTTP
    response = make_response(output.read())
    response.headers["Content-Disposition"] = "attachment; filename=pafs_militares.xlsx"
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    return response


def validate_vacation_period(start_date, days):
    next_year = date.today().year + 1
    end_date = start_date + timedelta(days=days - 1)
    # if start_date.year != next_year or end_date.year != next_year:
    #     raise ValueError("As férias devem ocorrer apenas no próximo ano.")
    if end_date > date(next_year, 12, 31):
        raise ValueError("As férias não podem ultrapassar 31 de dezembro.")


@app.route('/grafico-todos-militares', methods=['GET'])
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'SUPER USER', 'DIRETOR DRH')
def grafico_todos_militares():
    # Seleciona todos os militares
    militares = (
        database.session.query(Militar, Paf)
        .outerjoin(Paf, Paf.militar_id == Militar.id)
        .all()
    )

    # Contar número de militares de férias por mês
    ferias_por_mes = {mes.id: 0 for mes in Meses.query.all()}
    for militar, paf in militares:
        if paf:
            if paf.primeiro_periodo_ferias:
                mes = paf.primeiro_periodo_ferias.month
                ferias_por_mes[mes] += 1
            if paf.segundo_periodo_ferias:
                mes = paf.segundo_periodo_ferias.month
                ferias_por_mes[mes] += 1
            if paf.terceiro_periodo_ferias:
                mes = paf.terceiro_periodo_ferias.month
                ferias_por_mes[mes] += 1

    # Gerar gráfico
    labels = [mes.mes for mes in Meses.query.all()]
    values = [ferias_por_mes[mes.id] for mes in Meses.query.all()]

    plt.figure(figsize=(10, 6))
    plt.bar(labels, values, color='skyblue')
    plt.xlabel('Mês')
    plt.ylabel('Número de Militares de Férias')
    plt.title('Militares de Férias por Mês')
    plt.xticks(rotation=25)

    # Salvar gráfico em um buffer de memória
    buf = BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    image_base64 = base64.b64encode(buf.getvalue()).decode('utf-8')
    buf.close()

    return Response(response=image_base64, status=200, mimetype='text/plain')


def paf_ano_vigente():
    return current_app.config.get('PAF_ANO_VIGENTE', datetime.now().year)


@app.route('/ferias_dados', methods=['GET', 'POST'])
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'SUPER USER', 'DIRETOR DRH')
def ferias_dados():
    draw = request.form.get('draw', type=int)
    start = request.form.get('start', type=int)
    length = request.form.get('length', type=int)
    search_value = request.form.get('search[value]', type=str)

    # ✅ ano vem do frontend; fallback 2026
    ano = request.form.get('ano', type=int) or 2026

    query = (
        database.session.query(Militar, Paf)
        .outerjoin(Paf, and_(
            Militar.id == Paf.militar_id,
            Paf.ano_referencia == ano
        ))
    )

    if search_value:
        query = query.filter(
            (Militar.nome_completo.ilike(f'%{search_value}%')) |
            (Militar.matricula.ilike(f'%{search_value}%')) |
            (Militar.quadro.has(quadro=search_value))
        )

    total_records = query.count()

    militares_pafs = query.offset(start).limit(length).all()

    data = []
    for militar, paf in militares_pafs:
        data.append({
            "posto_grad": militar.posto_grad.sigla if militar.posto_grad else "",
            "nome_completo": militar.nome_completo,
            "matricula": militar.matricula,
            "quadro": militar.quadro.quadro if militar.quadro else "",
            "mes_usufruto": paf.mes_usufruto if paf else "",
            "qtd_dias_1": paf.qtd_dias_primeiro_periodo if paf else "",
            "inicio_1": str(paf.primeiro_periodo_ferias) if paf and paf.primeiro_periodo_ferias else "",
            "fim_1": str(paf.fim_primeiro_periodo) if paf and paf.fim_primeiro_periodo else "",
            "qtd_dias_2": paf.qtd_dias_segundo_periodo if paf else "",
            "inicio_2": str(paf.segundo_periodo_ferias) if paf and paf.segundo_periodo_ferias else "",
            "fim_2": str(paf.fim_segundo_periodo) if paf and paf.fim_segundo_periodo else "",
            "qtd_dias_3": paf.qtd_dias_terceiro_periodo if paf else "",
            "inicio_3": str(paf.terceiro_periodo_ferias) if paf and paf.terceiro_periodo_ferias else "",
            "fim_3": str(paf.fim_terceiro_periodo) if paf and paf.fim_terceiro_periodo else "",
            "id": militar.id
        })

    return jsonify({
        "draw": draw,
        "recordsTotal": total_records,
        "recordsFiltered": total_records,
        "data": data
    })


@app.route('/ferias', methods=['GET'])
@login_required
@checar_ocupacao('SUPER USER')
def exibir_ferias():
    ano_vigente = 2026  # depois tu puxa de config
    return render_template('ferias.html', ano_atual=datetime.now().year, ano_vigente=ano_vigente)


@app.route('/pafs/nao_preenchidos')
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'SUPER USER')
def pafs_nao_preenchidos():
    subquery_pafs = database.session.query(Paf.militar_id).subquery()

    prioridade_obm = case(
        (Obm.sigla == 'GAB SUBCMT-GERAL', 1),
        else_=2
    )

    ordem_posto = case(
        (PostoGrad.sigla == 'CEL', 1),
        (PostoGrad.sigla == 'TC', 2),
        (PostoGrad.sigla == 'MAJ', 3),
        (PostoGrad.sigla == 'CAP', 4),
        (PostoGrad.sigla == '1 TEN', 5),
        (PostoGrad.sigla == '2 TEN', 6),
        (PostoGrad.sigla == 'AL OF', 7),
        (PostoGrad.sigla == 'ALUNO OFICIAL', 8),
        (PostoGrad.sigla == 'SUBTENENTE', 9),
        (PostoGrad.sigla == '1 SGT', 10),
        (PostoGrad.sigla == '2 SGT', 11),
        (PostoGrad.sigla == '3 SGT', 12),
        (PostoGrad.sigla == 'AL SGT', 13),
        (PostoGrad.sigla == 'CB', 14),
        (PostoGrad.sigla == 'SD', 15),
        (PostoGrad.sigla == 'AL SD', 16),
        else_=99
    )

    sub_militares = (
        database.session.query(
            Militar.id.label("militar_id"),
            Militar.nome_completo,
            PostoGrad.sigla.label("posto_grad"),
            Quadro.quadro.label("quadro"),
            Obm.sigla.label("obm"),
            ordem_posto.label("ordem"),
            func.row_number().over(
                partition_by=Militar.id,
                order_by=[prioridade_obm.asc(), MilitarObmFuncao.id.desc()]
            ).label("linha")
        )
        .join(MilitarObmFuncao, Militar.id == MilitarObmFuncao.militar_id)
        .join(Obm, Obm.id == MilitarObmFuncao.obm_id)
        .join(PostoGrad, PostoGrad.id == Militar.posto_grad_id)
        .join(Quadro, Quadro.id == Militar.quadro_id)
        .filter(MilitarObmFuncao.data_fim.is_(None))  # OBMs ativas
        .filter(~Militar.id.in_(subquery_pafs))  # sem PAF
        .subquery()
    )

    militares_sem_paf = (
        database.session.query(
            sub_militares.c.nome_completo,
            sub_militares.c.posto_grad,
            sub_militares.c.quadro,
            sub_militares.c.obm
        )
        .filter(sub_militares.c.linha == 1)
        .order_by(sub_militares.c.ordem, sub_militares.c.obm)
        .all()
    )

    return render_template("pafs_nao_preenchidos.html", militares=militares_sem_paf)


@app.route('/api-sesuite', methods=['GET'])
def api_sesuite():
    # Realiza a consulta no banco de dados
    militares = (
        database.session.query(Militar, Paf)
        .outerjoin(Paf, Paf.militar_id == Militar.id)
        .all()
    )

    # Lista para armazenar os dados que serão retornados no JSON
    resultado = []

    # Itera sobre os resultados da consulta
    for militar, paf in militares:
        # Cria um dicionário com os dados do militar e do Paf
        militar_data = {
            'posto_grad': militar.posto_grad.sigla if militar.posto_grad else None,
            'nome_completo': militar.nome_completo,
            'matricula': militar.matricula,
            'quadro': militar.quadro.quadro if militar.quadro else None,
            'ferias_usufruto': {
                'mes_usufruto': paf.mes_usufruto if paf else None,
                'qtd_dias_primeiro_periodo': paf.qtd_dias_primeiro_periodo if paf else None,
                'primeiro_periodo_ferias': paf.primeiro_periodo_ferias.strftime(
                    '%Y-%m-%d') if paf and paf.primeiro_periodo_ferias else None,
                'fim_primeiro_periodo': paf.fim_primeiro_periodo.strftime(
                    '%Y-%m-%d') if paf and paf.fim_primeiro_periodo else None,
                'qtd_dias_segundo_periodo': paf.qtd_dias_segundo_periodo if paf else None,
                'segundo_periodo_ferias': paf.segundo_periodo_ferias.strftime(
                    '%Y-%m-%d') if paf and paf.segundo_periodo_ferias else None,
                'fim_segundo_periodo': paf.fim_segundo_periodo.strftime(
                    '%Y-%m-%d') if paf and paf.fim_segundo_periodo else None,
                'qtd_dias_terceiro_periodo': paf.qtd_dias_terceiro_periodo if paf else None,
                'terceiro_periodo_ferias': paf.terceiro_periodo_ferias.strftime(
                    '%Y-%m-%d') if paf and paf.terceiro_periodo_ferias else None,
                'fim_terceiro_periodo': paf.fim_terceiro_periodo.strftime(
                    '%Y-%m-%d') if paf and paf.fim_terceiro_periodo else None
            }
        }
        # Adiciona o dicionário à lista de resultados
        resultado.append(militar_data)

    # Retorna a lista de resultados como JSON
    return jsonify(resultado)


@app.route("/debug/militar/<int:militar_id>/full")
@login_required
def debug_militar_full(militar_id):
    # profundidade padrão (pode diminuir se ainda ficar pesado)
    depth = request.args.get("depth", default=4, type=int)

    # monta options de eager load para TODAS as relationships do Militar
    mapper = sa_inspect(Militar)
    rel_options = [
        selectinload(getattr(Militar, rel.key))
        for rel in mapper.relationships
    ]

    militar = (
        database.session.query(Militar)
        .options(*rel_options)
        .get(militar_id)
    )

    if not militar:
        return jsonify({"error": "Militar não encontrado"}), 404

    payload = sa_to_dict(militar, depth=depth, root_class=Militar)

    return jsonify(payload)


@app.route('/ferias-chefe', methods=['GET'])
@login_required
@checar_ocupacao('DIRETOR DRH', 'DIRETOR', 'CHEFE', 'SUPER USER')
def exibir_ferias_chefe():
    meses = {
        "Janeiro": 1, "Fevereiro": 2, "Março": 3, "Abril": 4, "Maio": 5, "Junho": 6,
        "Julho": 7, "Agosto": 8, "Setembro": 9, "Outubro": 10, "Novembro": 11, "Dezembro": 12
    }

    dia_atual = datetime.now().day
    obms_adicionais = [16, 17, 18, 19, 20, 21, 22, 23, 24, 25]
    lista_obms = []

    obm1 = Obm.query.get(current_user.obm_id_1)
    if obm1:
        lista_obms.append(obm1)

    if current_user.obm_id_2:
        obm2 = Obm.query.get(current_user.obm_id_2)
        if obm2:
            lista_obms.append(obm2)

    if current_user.obm_id_1 == 16:
        adicionais = Obm.query.filter(Obm.id.in_(obms_adicionais)).all()
        lista_obms.extend(adicionais)

    ano_vigente = paf_ano_vigente()
    return render_template(
        'ferias_chefe2.html',
        lista_obms=lista_obms,
        ano_atual=datetime.now().year,
        dia_atual=dia_atual,
        ano_vigente=ano_vigente,
    )


@app.route('/pafs/tabela/<int:obm_id>', methods=['GET'])
@login_required
@checar_ocupacao('DIRETOR DRH', 'DIRETOR', 'CHEFE', 'SUPER USER')
def carregar_tabela_obm(obm_id):
    ano = int(request.args.get("ano") or datetime.now().year)

    meses = {
        "Janeiro": 1, "Fevereiro": 2, "Março": 3, "Abril": 4, "Maio": 5, "Junho": 6,
        "Julho": 7, "Agosto": 8, "Setembro": 9, "Outubro": 10, "Novembro": 11, "Dezembro": 12
    }
    current_month = datetime.now().month
    current_date = datetime.now().date()

    obm = Obm.query.get(obm_id)
    if not obm:
        return "<div class='alert alert-danger'>OBM não encontrada</div>", 404

    militares_pafs = (
        database.session.query(Militar, Paf)
        .outerjoin(Paf, database.and_(
            Paf.militar_id == Militar.id,
            Paf.ano_referencia == ano
        ))
        .options(joinedload(Militar.obm_funcoes))
        .join(MilitarObmFuncao, Militar.id == MilitarObmFuncao.militar_id)
        .filter(
            MilitarObmFuncao.obm_id == obm_id,
            MilitarObmFuncao.data_fim.is_(None)
        )
        .all()
    )

    return render_template(
        'partial_tabela_obm.html',
        obm=obm,
        militares_pafs=militares_pafs,
        meses=meses,
        current_month=current_month,
        current_date=current_date,
        ano=ano,
    )


@app.route('/grafico-ferias/<int:obm_id>', methods=['GET'])
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'SUPER USER')
def grafico_ferias(obm_id):
    # Lista de OBMs adicionais para obm_id_1 == 16
    obms_adicionais = [16, 17, 18, 19, 20, 21, 22, 23, 24, 25]

    # Seleciona militares da OBM específica
    militares = (
        database.session.query(Militar, Paf)
        .outerjoin(Paf, Paf.militar_id == Militar.id)
        .join(MilitarObmFuncao, Militar.id == MilitarObmFuncao.militar_id)
        .filter(
            (MilitarObmFuncao.obm_id == obm_id) |
            (MilitarObmFuncao.obm_id.in_(obms_adicionais) if obm_id == 16 else False)
        )
        .all()
    )

    # Contar número de militares de férias por mês
    ferias_por_mes = {mes.id: 0 for mes in Meses.query.all()}
    for militar, paf in militares:
        if paf:
            if paf.primeiro_periodo_ferias:
                mes = paf.primeiro_periodo_ferias.month
                ferias_por_mes[mes] += 1
            if paf.segundo_periodo_ferias:
                mes = paf.segundo_periodo_ferias.month
                ferias_por_mes[mes] += 1
            if paf.terceiro_periodo_ferias:
                mes = paf.terceiro_periodo_ferias.month
                ferias_por_mes[mes] += 1

    # Gerar gráfico
    labels = [mes.mes for mes in Meses.query.all()]
    values = [ferias_por_mes[mes.id] for mes in Meses.query.all()]

    plt.figure(figsize=(10, 6))
    plt.bar(labels, values, color='skyblue')
    plt.xlabel('Mês')
    plt.ylabel('Número de Militares de Férias')
    plt.title('Militares de Férias por Mês')
    plt.xticks(rotation=25)

    # Salvar gráfico em um buffer de memória
    buf = BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    image_base64 = base64.b64encode(buf.getvalue()).decode('utf-8')
    buf.close()

    return Response(response=image_base64, status=200, mimetype='text/plain')


def parse_date(date_string):
    try:
        return datetime.strptime(date_string, '%Y-%m-%d').date()
    except (ValueError, TypeError):
        return None


@app.route('/pafs/update', methods=['POST'])
@login_required
def update_paf():
    hoje = datetime.now().day
    if (hoje < 10 or hoje > 20) and getattr(current_user, 'funcao_user_id', None) != 6:
        return jsonify({"message": "Alterações só são permitidas de 10 a 20 de cada mês."}), 403

    data = request.form
    militar_id = int(data.get('militar_id') or 0)

    # ✅ novo: ano do form, se não vier usa ano atual do servidor
    ano = int(data.get('ano_referencia') or datetime.now().year)

    # ✅ novo: busca por militar + ano
    paf = Paf.query.filter_by(militar_id=militar_id,
                              ano_referencia=ano).first()
    if not paf:
        paf = Paf(militar_id=militar_id, ano_referencia=ano)
        database.session.add(paf)

    mes_usufruto = data.get('mes_usufruto')

    qtd_dias_primeiro_periodo = int(data.get('qtd_dias_1') or 0)
    primeiro_periodo_inicio = parse_date(data.get('inicio_1'))
    primeiro_periodo_fim = parse_date(data.get('fim_1'))

    qtd_dias_segundo_periodo = int(data.get('qtd_dias_2') or 0)
    segundo_periodo_inicio = parse_date(data.get('inicio_2'))
    segundo_periodo_fim = parse_date(data.get('fim_2'))

    qtd_dias_terceiro_periodo = int(data.get('qtd_dias_3') or 0)
    terceiro_periodo_inicio = parse_date(data.get('inicio_3'))
    terceiro_periodo_fim = parse_date(data.get('fim_3'))

    # validação (mantém tua lógica)
    try:
        if primeiro_periodo_inicio:
            validate_vacation_period(
                primeiro_periodo_inicio, qtd_dias_primeiro_periodo)
        if segundo_periodo_inicio:
            validate_vacation_period(
                segundo_periodo_inicio, qtd_dias_segundo_periodo)
        if terceiro_periodo_inicio:
            validate_vacation_period(
                terceiro_periodo_inicio, qtd_dias_terceiro_periodo)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    paf.mes_usufruto = mes_usufruto
    paf.qtd_dias_primeiro_periodo = qtd_dias_primeiro_periodo
    paf.primeiro_periodo_ferias = primeiro_periodo_inicio
    paf.fim_primeiro_periodo = primeiro_periodo_fim

    paf.qtd_dias_segundo_periodo = qtd_dias_segundo_periodo
    paf.segundo_periodo_ferias = segundo_periodo_inicio
    paf.fim_segundo_periodo = segundo_periodo_fim

    paf.qtd_dias_terceiro_periodo = qtd_dias_terceiro_periodo
    paf.terceiro_periodo_ferias = terceiro_periodo_inicio
    paf.fim_terceiro_periodo = terceiro_periodo_fim

    paf.usuario_id = current_user.id
    paf.data_alteracao = datetime.now()

    database.session.commit()
    return jsonify({"message": "Dados salvos com sucesso!"})


@app.route('/adicionar-motorista', methods=['GET', 'POST'])
@login_required
def adicionar_motorista():
    form_motorista = FormMotoristas()

    # Carregar todos os militares de uma vez com os dados necessários
    militares_query = (
        database.session.query(
            Militar.id,
            Militar.nome_completo,
            Militar.matricula,
            PostoGrad.sigla.label("posto_grad_sigla"),
            Obm.sigla.label("obm_sigla")
        )
        .outerjoin(PostoGrad, Militar.posto_grad_id == PostoGrad.id)
        .outerjoin(MilitarObmFuncao, (MilitarObmFuncao.militar_id == Militar.id) & (MilitarObmFuncao.data_fim == None))
        .outerjoin(Obm, MilitarObmFuncao.obm_id == Obm.id)
        .order_by(Militar.nome_completo)
        .all()
    )

    form_motorista.nome_completo.choices = [
        (militar.id, militar.nome_completo) for militar in militares_query if militar.id is not None
    ]

    militares = {
        militar.id: {
            'matricula': militar.matricula,
            'obm_id_1': militar.obm_sigla,
            'posto_grad_id': militar.posto_grad_sigla
        }
        for militar in militares_query
    }

    form_motorista.categoria_id.choices = [
        ('', '-- Selecione uma categoria --')
    ] + [(categoria.id, categoria.sigla) for categoria in Categoria.query.all()]

    if form_motorista.validate_on_submit():
        try:
            novo_motorista = Motoristas(
                militar_id=form_motorista.nome_completo.data,
                categoria_id=form_motorista.categoria_id.data,
                boletim_geral=form_motorista.boletim_geral.data,
                siged=form_motorista.siged.data,
                usuario_id=current_user.id,
                desclassificar="NÃO",
                created=datetime.utcnow()
            )

            if form_motorista.cnh_imagem.data and form_motorista.cnh_imagem.data.filename != '':
                file = form_motorista.cnh_imagem.data
                ext = file.filename.split('.')[-1]
                timestamp = datetime.utcnow().strftime('%Y%m%d%H%M%S')

                nome_militar = next(
                    (m.nome_completo for m in militares_query if m.id ==
                     form_motorista.nome_completo.data),
                    'motorista'
                ).replace(" ", "_")

                nome_arquivo = secure_filename(
                    f"{nome_militar}_cnh_{timestamp}.{ext}")
                file_bytes = file.read()

                # Upload para o root do bucket (sem suappasta)
                app.supabase.storage.from_('motoristas').upload(
                    path=nome_arquivo,
                    file=file_bytes,
                    file_options={"content-type": file.mimetype}
                )

                # Salva apenas o nome do arquivo no banco
                novo_motorista.cnh_imagem = nome_arquivo

            database.session.add(novo_motorista)
            database.session.commit()
            flash('Motorista cadastrado com sucesso!', 'success')
            return redirect(url_for('adicionar_motorista'))

        except Exception as e:
            database.session.rollback()
            flash(f'Erro ao cadastrar motorista: {str(e)}', 'danger')

    return render_template('adicionar_motorista.html', form_motorista=form_motorista, militares=militares)


@app.route('/motoristas', methods=['GET', 'POST'])
@login_required
def motoristas():
    form_filtro = FormFiltroMotorista()

    form_filtro.obm_id.choices = [
        ('', '-- Selecione OBM --')] + [(obm.id, obm.sigla) for obm in Obm.query.all()]
    form_filtro.posto_grad_id.choices = [('', '-- Selecione Posto/Grad --')] + [
        (posto.id, posto.sigla) for posto in PostoGrad.query.all()]
    form_filtro.categoria_id.choices = [('', '-- Selecione uma categoria --')] + [(
        categoria.id, categoria.sigla) for categoria in Categoria.query.all()]

    page = request.args.get('page', 1, type=int)
    per_page = 10
    search = request.args.get('search', '', type=str)
    obm_id = request.args.get('obm_id', '', type=str)
    posto_grad_id = request.args.get('posto_grad_id', '', type=str)
    categoria_id = request.args.get('categoria_id', '', type=str)

    # Query base — junta Militar e Motoristas
    # IMPORTANTE: exclui motoristas desclassificados (desclassificar == 'SIM')
    query = Motoristas.query.join(Militar).filter(
        Motoristas.desclassificar != 'SIM')

    # Filtro por OBM
    if obm_id:
        subquery = MilitarObmFuncao.query.filter_by(
            obm_id=obm_id).with_entities(MilitarObmFuncao.militar_id)
        query = query.filter(Motoristas.militar_id.in_(subquery))

    # Filtro por Posto/Graduação
    if posto_grad_id:
        query = query.filter(Militar.posto_grad_id == posto_grad_id)

    # Filtro por Categoria
    if categoria_id:
        query = query.filter(Motoristas.categoria_id == categoria_id)

    # Filtro por Nome
    if search:
        query = query.filter(Militar.nome_completo.ilike(f'%{search}%'))

    # Apenas registros ativos (se você usa modified para histórico)
    query = query.filter(Motoristas.modified.is_(None))

    # Paginação
    motoristas_paginados = query.order_by(
        Militar.nome_completo.asc()).paginate(page=page, per_page=per_page)

    # Contagem de militares (total geral) e motoristas válidos (exclui desclassificados)
    total_militares = Militar.query.count()
    total_motoristas = Motoristas.query.filter(Motoristas.modified.is_(
        None), Motoristas.desclassificar != 'SIM').count()

    # Gráfico: Percentual de militares que são motoristas (exclui desclassificados)
    labels_motoristas = ['Motoristas', 'Não são motoristas']
    values_motoristas = [total_motoristas, total_militares - total_motoristas]
    fig_motoristas = go.Figure(
        data=[go.Pie(labels=labels_motoristas, values=values_motoristas, hole=0.4)])
    grafico_motoristas = pio.to_json(fig_motoristas)

    # Gráfico: Motoristas por categoria (exclui desclassificados)
    categorias = database.session.query(
        Categoria.sigla,
        database.func.count(Motoristas.id)
    ).join(Motoristas).filter(Motoristas.modified.is_(None), Motoristas.desclassificar != 'SIM').group_by(Categoria.sigla).all()
    labels_categorias = [c[0] for c in categorias]
    values_categorias = [c[1] for c in categorias]
    fig_categorias = go.Figure(
        data=[go.Pie(labels=labels_categorias, values=values_categorias, hole=0.4)])
    grafico_categorias = pio.to_json(fig_categorias)

    # Gráfico: Motoristas por OBM (exclui desclassificados)
    obms = database.session.query(
        Obm.sigla,
        database.func.count(Motoristas.id)
    ).join(MilitarObmFuncao, Obm.id == MilitarObmFuncao.obm_id).join(
        Motoristas, MilitarObmFuncao.militar_id == Motoristas.militar_id
    ).filter(Motoristas.modified.is_(None), Motoristas.desclassificar != 'SIM').group_by(Obm.sigla).all()
    labels_obms = [obm[0] for obm in obms]
    values_obms = [obm[1] for obm in obms]
    fig_obms = go.Figure(
        data=[go.Pie(labels=labels_obms, values=values_obms, hole=0.4)])
    grafico_obms = pio.to_json(fig_obms)

    return render_template(
        'motoristas.html',
        motoristas=motoristas_paginados,
        search=search,
        form_filtro=form_filtro,
        grafico_motoristas=grafico_motoristas,
        grafico_categorias=grafico_categorias,
        grafico_obms=grafico_obms
    )


@app.route('/atualizar-motorista/<int:motorista_id>', methods=['GET', 'POST'])
@login_required
def atualizar_motorista(motorista_id):
    motorista = Motoristas.query.get_or_404(motorista_id)

    form_motorista = FormMotoristas(obj=motorista)

    # Define a opção única para o militar atual
    militar_atual = (motorista.militar.id, motorista.militar.nome_completo)
    form_motorista.nome_completo.choices = [militar_atual]
    form_motorista.nome_completo.data = motorista.militar.id

    # Carrega opções de categoria
    form_motorista.categoria_id.choices = [('', '-- Selecione uma categoria --')] + [
        (categoria.id, categoria.sigla) for categoria in Categoria.query.all()
    ]

    # Preenche dados exibidos
    form_motorista.matricula.data = motorista.militar.matricula
    form_motorista.posto_grad_id.data = motorista.militar.posto_grad.sigla if motorista.militar.posto_grad else None
    form_motorista.obm_id_1.data = motorista.militar.obm_funcoes[
        0].obm.sigla if motorista.militar.obm_funcoes else None

    if request.method == 'POST':
        # Ação específica: desclassificar (vindo do modal)
        if request.form.get('action') == 'desclassificar':
            motivo = request.form.get('motivo_desclassificacao', None)
            try:
                # Marca o antigo como modificado (histórico)
                motorista.modified = datetime.utcnow()
                database.session.commit()

                # Novo registro marcando desclassificado
                novo_motorista = Motoristas(
                    militar_id=motorista.militar_id,
                    categoria_id=motorista.categoria_id,  # mantém a categoria atual
                    boletim_geral=motorista.boletim_geral,
                    siged=motorista.siged,
                    usuario_id=current_user.id,
                    vencimento_cnh=motorista.vencimento_cnh,
                    created=datetime.utcnow(),
                    desclassificar='SIM'
                )

                # opcional: salvar motivo concatenado (ou use outra coluna se preferir)
                if motivo:
                    # limite para evitar overflow da coluna (sua coluna é String(30) — cuidado!)
                    # nesse exemplo eu salvo apenas uma flag e concateno motivo em outra tabela/coluna seria ideal.
                    novo_motorista.desclassificar = 'SIM'
                    # se quiser armazenar motivo curto (<=30): usar outra coluna; aqui guardo prefixo do motivo se couber
                    # novo_motorista.desclassificar = (motivo[:30])  # *substitui* a flag se preferir
                    # Para não perder a info, recomendo criar uma coluna `desclassificar_motivo` (ver migração abaixo)
                # Salva novo registro
                database.session.add(novo_motorista)
                database.session.commit()

                flash('Motorista desclassificado com sucesso!', 'warning')
                return redirect(url_for('motoristas'))

            except Exception as e:
                database.session.rollback()
                flash(f'Erro ao desclassificar motorista: {str(e)}', 'danger')
                print("Erro ao desclassificar motorista:", e)
                # cai para re-render do template com flash

        # Caso seja submissão normal (salvar/editar)
        if form_motorista.validate_on_submit():
            try:
                # Marca antigo como modificado (histórico)
                motorista.modified = datetime.utcnow()
                database.session.commit()

                novo_motorista = Motoristas(
                    militar_id=motorista.militar_id,
                    categoria_id=motorista.categoria_id,
                    boletim_geral=motorista.boletim_geral,
                    siged=motorista.siged,
                    usuario_id=current_user.id,
                    vencimento_cnh=motorista.vencimento_cnh,
                    created=datetime.utcnow(),
                    desclassificar='SIM',
                    desclassificar_por=current_user.id,
                    desclassificar_em=datetime.utcnow()
                )

                # Verifica se imagem foi enviada
                if form_motorista.cnh_imagem.data and form_motorista.cnh_imagem.data.filename != '':
                    file = form_motorista.cnh_imagem.data
                    ext = file.filename.split('.')[-1]
                    timestamp = datetime.utcnow().strftime('%Y%m%d%H%M%S')
                    nome_formatado = motorista.militar.nome_completo.replace(
                        " ", "_")
                    nome_arquivo = secure_filename(
                        f"{nome_formatado}_cnh_{timestamp}.{ext}")
                    file_bytes = file.read()

                    # Upload para o root do bucket
                    app.supabase.storage.from_('motoristas').upload(
                        path=nome_arquivo,
                        file=file_bytes,
                        file_options={"content-type": file.mimetype}
                    )

                    # Salva apenas o nome no banco
                    novo_motorista.cnh_imagem = nome_arquivo

                # Salva novo registro no banco
                database.session.add(novo_motorista)
                database.session.commit()

                flash('Motorista atualizado com sucesso!', 'success')
                return redirect(url_for('motoristas'))

            except Exception as e:
                database.session.rollback()
                flash(f'Erro ao atualizar motorista: {str(e)}', 'danger')
                print("Erro ao atualizar motorista:", e)

    return render_template(
        'atualizar_motorista.html',
        form_motorista=form_motorista,
        motorista=motorista
    )


@app.route('/motoristas-desclassificados', methods=['GET', 'POST'])
@login_required
def motoristas_desclassificados():
    form_filtro = FormFiltroMotorista()

    form_filtro.obm_id.choices = [
        ('', '-- Selecione OBM --')] + [(obm.id, obm.sigla) for obm in Obm.query.all()]
    form_filtro.posto_grad_id.choices = [('', '-- Selecione Posto/Grad --')] + [
        (posto.id, posto.sigla) for posto in PostoGrad.query.all()]
    form_filtro.categoria_id.choices = [('', '-- Selecione uma categoria --')] + [(
        categoria.id, categoria.sigla) for categoria in Categoria.query.all()]

    page = request.args.get('page', 1, type=int)
    per_page = 10
    search = request.args.get('search', '', type=str)
    obm_id = request.args.get('obm_id', '', type=str)
    posto_grad_id = request.args.get('posto_grad_id', '', type=str)
    categoria_id = request.args.get('categoria_id', '', type=str)

    # Query base: apenas desclassificados atuais (registros não modificados e desclassificar == 'SIM')
    query = Motoristas.query.join(Militar).filter(
        Motoristas.desclassificar == 'SIM', Motoristas.modified.is_(None))

    # Filtro por OBM
    if obm_id:
        subquery = MilitarObmFuncao.query.filter_by(
            obm_id=obm_id).with_entities(MilitarObmFuncao.militar_id)
        query = query.filter(Motoristas.militar_id.in_(subquery))

    # Filtro por Posto/Graduação
    if posto_grad_id:
        query = query.filter(Militar.posto_grad_id == posto_grad_id)

    # Filtro por Categoria
    if categoria_id:
        query = query.filter(Motoristas.categoria_id == categoria_id)

    # Filtro por Nome
    if search:
        query = query.filter(Militar.nome_completo.ilike(f'%{search}%'))

    # Paginação
    motoristas_paginados = query.order_by(Motoristas.desclassificar_em.desc(
    ).nullslast(), Militar.nome_completo.asc()).paginate(page=page, per_page=per_page)

    # Contagem total de desclassificados (para resumo)
    total_desclassificados = Motoristas.query.filter(
        Motoristas.desclassificar == 'SIM', Motoristas.modified.is_(None)).count()

    # Gráfico 1: Desclassificados por categoria
    categorias = database.session.query(
        Categoria.sigla,
        database.func.count(Motoristas.id)
    ).join(Motoristas).filter(Motoristas.desclassificar == 'SIM', Motoristas.modified.is_(None)).group_by(Categoria.sigla).all()
    labels_categorias = [c[0] for c in categorias]
    values_categorias = [c[1] for c in categorias]
    fig_categorias = go.Figure(
        data=[go.Pie(labels=labels_categorias, values=values_categorias, hole=0.4)])
    grafico_categorias = pio.to_json(fig_categorias)

    # Gráfico 2: Desclassificados por OBM
    obms = database.session.query(
        Obm.sigla,
        database.func.count(Motoristas.id)
    ).join(MilitarObmFuncao, Obm.id == MilitarObmFuncao.obm_id).join(
        Motoristas, MilitarObmFuncao.militar_id == Motoristas.militar_id
    ).filter(Motoristas.desclassificar == 'SIM', Motoristas.modified.is_(None)).group_by(Obm.sigla).all()
    labels_obms = [o[0] for o in obms]
    values_obms = [o[1] for o in obms]
    fig_obms = go.Figure(
        data=[go.Pie(labels=labels_obms, values=values_obms, hole=0.4)])
    grafico_obms = pio.to_json(fig_obms)

    # Gráfico 3: Evolução mensal de desclassificados (últimos 12 meses)
    # Usa date_trunc para agrupar por mês; funciona em PostgreSQL (Supabase)
    mensal = database.session.query(
        database.func.date_trunc(
            'month', Motoristas.desclassificar_em).label('mes'),
        database.func.count(Motoristas.id)
    ).filter(
        Motoristas.desclassificar == 'SIM',
        Motoristas.modified.is_(None),
        Motoristas.desclassificar_em.isnot(None)
    ).group_by('mes').order_by('mes').all()

    meses = [row[0].strftime('%Y-%m') for row in mensal] if mensal else []
    valores_mensais = [row[1] for row in mensal] if mensal else []
    fig_mensal = go.Figure(data=[go.Bar(x=meses, y=valores_mensais)])
    fig_mensal.update_layout(xaxis_title='Mês', yaxis_title='Desclassificados')
    grafico_mensal = pio.to_json(fig_mensal)

    return render_template(
        'motoristas_desclassificados.html',
        motoristas=motoristas_paginados,
        search=search,
        form_filtro=form_filtro,
        total_desclassificados=total_desclassificados,
        grafico_categorias=grafico_categorias,
        grafico_obms=grafico_obms,
        grafico_mensal=grafico_mensal
    )


@app.route("/viaturas", methods=["GET"])
def escolher_obm():
    obms = Obm.query.order_by(Obm.sigla.asc()).all()
    return render_template("viaturas_escolher_obm.html", obms=obms)


@app.route("/<int:obm_id>/viaturas", methods=["GET"])
def gerenciar_viaturas(obm_id):
    obm = Obm.query.get_or_404(obm_id)

    # Viaturas já dessa OBM
    viaturas_da_obm = (Viaturas.query
                       .filter(Viaturas.obm_id == obm_id)
                       .order_by(Viaturas.prefixo.asc(), Viaturas.placa.asc())
                       .all())

    # Viaturas sem OBM (ou de outra OBM) – mostramos as sem OBM para facilitar atribuição
    viaturas_sem_obm = (Viaturas.query
                        .filter(Viaturas.obm_id.is_(None))
                        .order_by(Viaturas.prefixo.asc(), Viaturas.placa.asc())
                        .all())

    # Motoristas preferencialmente desta OBM (se houver vínculo Militar↔OBM via MilitarObmFuncao)
    # Se não houver essa tabela/ligação, troque por Motoristas.query.all()
    motoristas = (database.session.query(Motoristas)
                  .join(Militar, Motoristas.militar_id == Militar.id)
                  .join(MilitarObmFuncao, MilitarObmFuncao.militar_id == Militar.id)
                  .filter(MilitarObmFuncao.obm_id == obm_id)
                  .order_by(Militar.nome_completo.asc())
                  .all())

    # Mapa de motoristas atuais por viatura (para preencher selects)
    motoristas_por_viatura = {}
    for v in viaturas_da_obm:
        vms = (ViaturaMilitar.query
               .filter_by(viatura_id=v.id)
               .all())
        motoristas_por_viatura[v.id] = [vm.militar_id for vm in vms]

    return render_template(
        "viaturas_gerenciar.html",
        obm=obm,
        viaturas_da_obm=viaturas_da_obm,
        viaturas_sem_obm=viaturas_sem_obm,
        motoristas=motoristas,
        motoristas_por_viatura=motoristas_por_viatura,
    )


@app.route("/<int:obm_id>/viaturas/atribuir", methods=["POST"])
def atribuir_viaturas_obm(obm_id):
    """
    Recebe a lista 'assigned_ids[]' (viaturas que devem ficar na OBM).
    Vamos:
      - Setar obm_id = obm_id para as IDs enviadas
      - Remover desta OBM (setar NULL) todas as que estavam e não vieram no POST
    """
    obm = Obm.query.get_or_404(obm_id)
    ids_enviados = request.form.getlist("assigned_ids[]")
    ids_enviados = [int(x) for x in ids_enviados]

    # Viaturas que hoje estão nessa OBM
    atuais = Viaturas.query.filter_by(obm_id=obm_id).all()
    ids_atuais = {v.id for v in atuais}

    # 1) adicionar/mover para esta OBM os enviados que não estão
    if ids_enviados:
        (Viaturas.query
         .filter(Viaturas.id.in_(ids_enviados))
         .update({Viaturas.obm_id: obm_id}, synchronize_session=False))

    # 2) remover desta OBM os que estavam e não estão mais na lista
    ids_remover = list(ids_atuais - set(ids_enviados))
    if ids_remover:
        (Viaturas.query
         .filter(Viaturas.id.in_(ids_remover))
         .update({Viaturas.obm_id: None}, synchronize_session=False))

    database.session.commit()
    flash("Atribuições de viaturas atualizadas para a OBM {}.".format(
        obm.sigla), "success")
    return redirect(url_for("gerenciar_viaturas", obm_id=obm_id))


@app.route("/viaturas/<int:viatura_id>/motoristas", methods=["POST"])
def salvar_motoristas_viatura(viatura_id):
    """
    Salva até 5 motoristas (militar_id) para a viatura.
    O gatilho no banco impede >5, mas também checamos aqui para UX.
    """
    v = Viaturas.query.get_or_404(viatura_id)
    selecionados = request.form.getlist("motoristas[]")
    selecionados = [int(x) for x in selecionados if x]

    if len(selecionados) > 5:
        flash("Selecione no máximo 5 motoristas para a viatura.", "warning")
        return redirect(url_for("viaturas_admin.gerenciar_viaturas", obm_id=v.obm_id or 0))

    # Atualiza o conjunto: remove os que saíram e adiciona os novos
    atuais = ViaturaMilitar.query.filter_by(viatura_id=viatura_id).all()
    ids_atuais = {vm.militar_id for vm in atuais}
    novos = set(selecionados) - ids_atuais
    remover = ids_atuais - set(selecionados)

    if remover:
        (ViaturaMilitar.query
            .filter(ViaturaMilitar.viatura_id == viatura_id,
                    ViaturaMilitar.militar_id.in_(remover))
            .delete(synchronize_session=False))

    for mid in novos:
        database.session.add(ViaturaMilitar(
            viatura_id=viatura_id, militar_id=mid))

    try:
        database.session.commit()
        flash("Motoristas atualizados para a viatura {}.".format(
            v.prefixo or v.placa), "success")
    except IntegrityError:
        database.session.rollback()
        flash("Não foi possível salvar: limite atingido ou motorista duplicado.", "danger")

    return redirect(url_for("viaturas_admin.gerenciar_viaturas", obm_id=v.obm_id or 0))


@app.route('/usuario/<usuario_id>/excluir', methods=['GET', 'POST'])
@login_required
@checar_ocupacao('DIRETOR', 'SUPER USER')
def excluir_usuario(usuario_id):
    usuario = User.query.get(usuario_id)
    if not usuario:
        flash('Usuário não encontrado', 'alert-warning')
        return redirect(url_for('usuarios'))

    # desvincula militares que apontam para esse usuário
    militares = Militar.query.filter_by(usuario_id=usuario.id).all()
    for m in militares:
        m.usuario_id = None
        database.session.add(m)

    database.session.delete(usuario)
    database.session.commit()
    flash('Usuário excluído permanentemente', 'alert-danger')
    return redirect(url_for('usuarios'))


@app.route('/militar/<int:militar_id>/excluir', methods=['GET', 'POST'])
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'MAPA DA FORÇA', 'SUPER USER')
def excluir_militar(militar_id):
    militar = Militar.query.get(militar_id)
    database.session.delete(militar)
    database.session.commit()
    flash('Militar e registros vinculados excluídos permanentemente', 'alert-danger')
    return redirect(url_for('militares'))


@app.route('/sair')
@login_required
def sair():
    logout_user()
    flash('Faça o Login para continuar', 'alert-success')
    return redirect(url_for('login'))


@app.route('/listar-cnhs')
@login_required
def listar_cnhs():
    # Lista do root do bucket motoristas (com path vazio!)
    arquivos = app.supabase.storage.from_('motoristas').list('')

    print("🟡 Arquivos retornados do Supabase:")
    if arquivos:
        for item in arquivos:
            print(" -", item['name'])
    else:
        print("⚠️ Nenhum arquivo retornado!")

    # Garante que estamos pegando só os arquivos (não pastas)
    nomes_arquivos = [item['name']
                      for item in arquivos if item['name'] and not item['name'].endswith('/')]

    return render_template('listar_cnhs.html', arquivos=nomes_arquivos)


@app.route('/vencimentos/novo', methods=['GET', 'POST'])
@login_required
def novo_vencimento():
    form = TabelaVencimentoForm()

    if 'tabela_id' in session:
        form.nome.validators = []
        form.lei.validators = []
        form.data_inicio.validators = []
        form.data_fim.validators = []

    postos = PostoGrad.query.all()
    form.posto_grad.choices = [(p.id, p.sigla) for p in postos]

    if request.method == 'POST' and form.validate_on_submit():
        # Finalizar a tabela e ir para o cálculo
        if 'finalizar' in request.form:
            session.pop('tabela_id', None)
            flash("Tabela finalizada com sucesso!", "success")
            # Altere para seu endpoint real
            return redirect(url_for('home'))

        # Criar nova tabela se ainda não houver na sessão
        if 'tabela_id' not in session:
            tabela = TabelaVencimento(
                nome=form.nome.data,
                lei=form.lei.data,
                data_inicio=form.data_inicio.data,
                data_fim=form.data_fim.data
            )
            database.session.add(tabela)
            database.session.flush()
            session['tabela_id'] = tabela.id
        else:
            tabela = TabelaVencimento.query.get(session['tabela_id'])
            if tabela is None:
                session.pop('tabela_id', None)
                flash(
                    "A tabela anterior foi removida ou expirou. Por favor, inicie novamente.", "warning")
                return redirect(url_for('novo_vencimento'))

        valor = ValorDetalhadoPostoGrad(
            tabela_id=tabela.id,
            posto_grad_id=form.posto_grad.data,
            soldo=form.soldo.data,
            grat_tropa=form.grat_tropa.data,
            gams=form.gams.data,
            valor_bruto=form.valor_bruto.data,
            curso_25=form.curso_25.data,
            curso_30=form.curso_30.data,
            curso_35=form.curso_35.data,
            bruto_esp=form.bruto_esp.data,
            bruto_mestre=form.bruto_mestre.data,
            bruto_dout=form.bruto_dout.data,
            fg_1=form.fg_1.data,
            fg_2=form.fg_2.data,
            fg_3=form.fg_3.data,
            fg_4=form.fg_4.data,
            aux_moradia=form.aux_moradia.data,
            etapas_capital=form.etapas_capital.data,
            etapas_interior=form.etapas_interior.data,
            seg_hora=form.seg_hora.data,
            motorista_a=form.motorista_a.data,
            motorista_b=form.motorista_b.data,
            motorista_ab=form.motorista_ab.data,
            motorista_cde=form.motorista_cde.data,
            tecnico_raiox=form.tecnico_raiox.data,
            tecnico_lab=form.tecnico_lab.data,
            mecanico=form.mecanico.data,
            fluvial=form.fluvial.data,
            explosivista=form.explosivista.data,
            coe=form.coe.data,
            tripulante=form.tripulante.data,
            piloto=form.piloto.data,
            aviacao=form.aviacao.data,
            mergulhador=form.mergulhador.data
        )

        database.session.add(valor)
        database.session.commit()

        flash("Posto adicionado à tabela com sucesso!", "success")
        return redirect(url_for('novo_vencimento', step=2))

    return render_template("form_tabela_vencimento.html", form=form, step=request.args.get('step'))


getcontext().prec = 10


def arred(valor):
    return Decimal(valor).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)


# Função para cálculo de dias no padrão 30/360 Europeu
def dias360_europeu(inicio: datetime, fim: datetime) -> int:
    d1, m1, a1 = inicio.day, inicio.month, inicio.year
    d2, m2, a2 = fim.day, fim.month, fim.year

    if d1 == 31:
        d1 = 30
    if d2 == 31:
        if d1 < 30:
            d2 = 1
            m2 += 1
            if m2 > 12:
                m2 = 1
                a2 += 1
        else:
            d2 = 30

    return (a2 - a1) * 360 + (m2 - m1) * 30 + (d2 - d1)


@app.route('/impacto/calcular', methods=['GET', 'POST'])
@login_required
def calcular_impacto():
    form = ImpactoForm()
    postos = PostoGrad.query.order_by(PostoGrad.id).all()
    form.posto_origem.choices = [(p.id, p.sigla) for p in postos]
    form.posto_destino.choices = [(p.id, p.sigla) for p in postos]

    resultado = None
    tabelas_usadas = []

    if request.method == 'GET' and request.args.get('show_modal') == '1':
        if "resultado" in session and "tabelas_usadas" in session:
            resultado = session.pop("resultado")
            tabelas_usadas = session.pop("tabelas_usadas")

    if request.method == 'POST' and form.validate_on_submit():
        data_inicio = form.data_inicio.data
        data_fim = form.data_fim.data
        efetivo = form.efetivo.data
        posto_origem_id = form.posto_origem.data
        posto_destino_id = form.posto_destino.data

        tabelas = TabelaVencimento.query.filter(
            TabelaVencimento.data_fim >= data_inicio,
            TabelaVencimento.data_inicio <= data_fim
        ).order_by(TabelaVencimento.data_inicio).all()

        if not tabelas:
            flash("Nenhuma tabela de vencimento encontrada para esse período.", "danger")
            return render_template("impacto_calculo.html", form=form)

        resultado_final = {
            "detalhes": [],
            "total": Decimal('0.00')
        }
        tabelas_usadas = []

        for tabela in tabelas:
            inicio_periodo = max(data_inicio, tabela.data_inicio)
            fim_periodo = min(data_fim, tabela.data_fim)

            dias = dias360_europeu(
                inicio_periodo, fim_periodo + timedelta(days=1))
            meses = Decimal(dias) / Decimal(30)
            coef = meses / Decimal(12)

            valor_origem = ValorDetalhadoPostoGrad.query.filter_by(
                tabela_id=tabela.id, posto_grad_id=posto_origem_id
            ).first()
            valor_destino = ValorDetalhadoPostoGrad.query.filter_by(
                tabela_id=tabela.id, posto_grad_id=posto_destino_id
            ).first()

            if not valor_origem or not valor_destino:
                flash(
                    f"Valores de postos não encontrados na tabela {tabela.nome}.", "danger")
                return render_template("impacto_calculo.html", form=form)

            diferenca = Decimal(valor_destino.valor_bruto) - \
                Decimal(valor_origem.valor_bruto)
            impacto_mensal = arred(diferenca * efetivo)
            retroativo = arred((impacto_mensal / Decimal(30)) * dias)
            ferias = arred((impacto_mensal / Decimal(3)) * coef)
            decimo = arred(impacto_mensal * coef)
            subtotal = retroativo + ferias + decimo

            resultado_final["detalhes"].append({
                "nome": tabela.nome,
                "inicio": tabela.data_inicio.strftime("%d/%m/%Y"),
                "fim": tabela.data_fim.strftime("%d/%m/%Y"),
                "dias": dias,
                "meses": str(meses),
                "coef": str(coef),
                "diferenca": str(diferenca),
                "impacto_mensal": str(impacto_mensal),
                "retroativo": str(retroativo),
                "ferias": str(ferias),
                "decimo": str(decimo),
                "total": str(subtotal)
            })

            resultado_final["total"] += subtotal

            tabelas_usadas.append({
                "nome": tabela.nome,
                "inicio": tabela.data_inicio.strftime("%d/%m/%Y"),
                "fim": tabela.data_fim.strftime("%d/%m/%Y")
            })

        # Cálculo do impacto atual fixo
        data_inicio_atual = date.today()
        data_fim_atual = data_fim
        dias_atual = dias360_europeu(
            data_inicio_atual, data_fim_atual + timedelta(days=1))
        meses_coef = Decimal(dias_atual) / Decimal(30)
        coef_proporcional = meses_coef / Decimal(12)

        tabela_atual = next((t for t in tabelas if t.data_inicio <=
                            data_inicio_atual and t.data_fim >= data_fim_atual), None)
        if tabela_atual:
            valor_origem_atual = ValorDetalhadoPostoGrad.query.filter_by(
                tabela_id=tabela_atual.id, posto_grad_id=posto_origem_id).first()
            valor_destino_atual = ValorDetalhadoPostoGrad.query.filter_by(
                tabela_id=tabela_atual.id, posto_grad_id=posto_destino_id).first()

            if valor_origem_atual and valor_destino_atual:
                diferenca_atual = Decimal(
                    valor_destino_atual.valor_bruto) - Decimal(valor_origem_atual.valor_bruto)
                impacto_mensal_atual = arred(diferenca_atual * efetivo)
                subtotal_atual = arred(impacto_mensal_atual * meses_coef)
                ferias_atual = arred(
                    (impacto_mensal_atual / Decimal(3)) * coef_proporcional)
                decimo_atual = arred(impacto_mensal_atual * coef_proporcional)
                total_sem_retroativo = subtotal_atual + ferias_atual + decimo_atual
                impacto_mensal_estimado = arred(
                    total_sem_retroativo / meses_coef)

                resultado_final["atual"] = {
                    "dias": dias_atual,
                    "meses_coef": str(meses_coef),
                    "coef": str(coef_proporcional),
                    "diferenca": str(diferenca_atual),
                    "impacto_mensal": str(impacto_mensal_atual),
                    "subtotal": str(subtotal_atual),
                    "ferias": str(ferias_atual),
                    "decimo": str(decimo_atual),
                    "total_sem_retroativo": str(total_sem_retroativo),
                    "impacto_mensal_estimado": str(impacto_mensal_estimado)
                }

        # REGISTRA IMPACTOS NA SESSÃO
        impactos_registrados = session.get("impactos_registrados", [])
        impactos_registrados.append(resultado_final)
        session["impactos_registrados"] = impactos_registrados

        session["resultado"] = resultado_final
        session["tabelas_usadas"] = tabelas_usadas

        params = urlencode({"show_modal": "1"})
        return redirect(f"{url_for('calcular_impacto')}?{params}")

    return render_template("impacto_calculo.html", form=form, resultado=resultado, tabelas_usadas=tabelas_usadas)


def calcular_semana(data_convocacao, data_base=None):
    if not data_base:
        # data base inicial da primeira semana
        data_base = datetime(2025, 5, 5)
    dias_passados = (data_convocacao - data_base).days
    numero_semana = dias_passados // 7 + 1
    return f"Semana {numero_semana}"


@app.route('/dashboard', methods=['GET', 'POST'])
@login_required
@checar_ocupacao('SUPER USER')
def dashboard():
    data = None
    if request.method == 'POST':
        convocados = int(request.form['convocados'])
        faltaram = int(request.form['faltaram'])
        desistiram = int(request.form['desistiram'])
        data_input = request.form['data']
        data_dt = datetime.strptime(data_input, '%Y-%m-%d')

        vagas_abertas = faltaram + desistiram
        semana = calcular_semana(data_dt)

        data_dict = {
            'Situação': ['Presentes', 'Faltaram', 'Desistiram', 'Vagas Abertas'],
            'Quantidade': [convocados - vagas_abertas, faltaram, desistiram, vagas_abertas]
        }

        # Salvar no banco
        registro = Convocacao(
            data=datetime.strptime(data_input, '%Y-%m-%d'),
            convocados=convocados,
            faltaram=faltaram,
            desistiram=desistiram,
            vagas_abertas=vagas_abertas,
            semana=semana
        )
        database.session.add(registro)
        database.session.commit()

        return render_template('dashboard.html', data=data_dict)

    return render_template('dashboard.html', data=None)


@app.route('/export-dashboard', methods=['POST'])
def export_dashboard():
    data = request.json
    df = pd.DataFrame(data)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Dados')
    output.seek(0)
    return send_file(output, download_name='dashboard.xlsx', as_attachment=True)


@app.route('/relatorio-convocacao', methods=['GET'])
@login_required
def relatorio_convocacao():
    registros = Convocacao.query.order_by(Convocacao.data).all()

    dados_por_semana = defaultdict(list)
    totais_semanais = {}
    somatorios_geral = {
        "convocados": 0,
        "presentes": 0,
        "faltaram": 0,
        "desistiram": 0,
        "vagas": 0
    }

    for r in registros:
        presentes = r.convocados - r.faltaram - r.desistiram
        vagas = r.faltaram + r.desistiram

        item = {
            "data": r.data.strftime('%d/%m/%Y'),
            "convocados": r.convocados,
            "faltaram": r.faltaram,
            "desistiram": r.desistiram,
            "presentes": presentes,
            "vagas": vagas
        }

        semana = r.semana
        dados_por_semana[semana].append(item)

        somatorios_geral["convocados"] += r.convocados
        somatorios_geral["presentes"] += presentes
        somatorios_geral["faltaram"] += r.faltaram
        somatorios_geral["desistiram"] += r.desistiram
        somatorios_geral["vagas"] += vagas

        if semana not in totais_semanais:
            totais_semanais[semana] = {
                "convocados": 0,
                "presentes": 0,
                "faltaram": 0,
                "desistiram": 0,
                "vagas": 0
            }

        totais_semanais[semana]["convocados"] += r.convocados
        totais_semanais[semana]["presentes"] += presentes
        totais_semanais[semana]["faltaram"] += r.faltaram
        totais_semanais[semana]["desistiram"] += r.desistiram
        totais_semanais[semana]["vagas"] += vagas

    dados_ordenados = dict(
        sorted(dados_por_semana.items(), key=lambda x: int(x[0].split()[-1])))
    totais_ordenados = dict(
        sorted(totais_semanais.items(), key=lambda x: int(x[0].split()[-1])))

    return render_template('relatorio.html', dados=dados_ordenados, totais_semanais=totais_ordenados, somatorios=somatorios_geral)


@app.route('/relatorio-convocacao/excel', methods=['GET'])
def relatorio_convocacao_excel():
    registros = Convocacao.query.order_by(Convocacao.data).all()
    dados = []

    for r in registros:
        presentes = r.convocados - r.faltaram - r.desistiram
        vagas = r.faltaram + r.desistiram
        dados.append({
            "Data": r.data.strftime('%Y-%m-%d'),
            "Convocados": r.convocados,
            "Faltaram": r.faltaram,
            "Desistiram": r.desistiram,
            "Presentes": presentes,
            "Vagas Abertas": vagas
        })

    df = pd.DataFrame(dados)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Convocacoes')
    output.seek(0)

    return send_file(output, download_name="convocacoes_admin.xlsx", as_attachment=True)


@app.route('/adicionar-convocacao', methods=['GET', 'POST'])
@login_required
def adicionar_convocacao():
    form = ControleConvocacaoForm()

    # Situações
    form.situacao_convocacao_id.choices = [
        (s.id, s.situacao) for s in SituacaoConvocacao.query.all()
    ]

    # Nomes (apenas o texto do nome na label)
    nomes = NomeConvocado.query.all()
    form.nome.choices = [(n.id, n.nome) for n in nomes]

    # ➊  ———  dicionário p/ preencher via JS
    nomes_data = {
        n.id: {
            "inscricao": n.inscricao or "",
            "classificacao": n.classificacao or "",
            "nota_final": n.nota_final or "",
        } for n in nomes
    }

    if form.validate_on_submit():
        selected_nome = NomeConvocado.query.get(form.nome.data)
        novo = ControleConvocacao(
            classificacao=form.classificacao.data,
            inscricao=form.inscricao.data,
            nome=selected_nome.nome,
            nota_final=form.nota_final.data,
            ordem_de_convocacao=form.ordem_de_convocacao.data,
            apresentou=form.apresentou.data,
            situacao_convocacao_id=form.situacao_convocacao_id.data,
            matricula=form.matricula.data,
            numero_da_matricula_doe=form.numero_da_matricula_doe.data,
            bg_matricula_doe=form.bg_matricula_doe.data,
            portaria_convocacao=form.portaria_convocacao.data,
            bg_portaria_convocacao=form.bg_portaria_convocacao.data,
            doe_portaria_convocacao=form.doe_portaria_convocacao.data,
            notificacao_pessoal=form.notificacao_pessoal.data,
            termo_desistencia=form.termo_desistencia.data,
            siged_desistencia=form.siged_desistencia.data
        )
        database.session.add(novo)
        database.session.delete(selected_nome)  # remove da fila
        database.session.commit()

        flash('Registro salvo com sucesso!', 'success')
        # ajuste a rota conforme seu sistema
        return redirect(url_for('adicionar_convocacao'))

    return render_template('form_convocacao.html', form=form, nomes_data=nomes_data)


@app.route('/controle-convocacao', methods=['GET'])
@login_required
def controle_convocacao():
    page = request.args.get('page', 1, type=int)
    per_page = 100
    search = request.args.get('search', '').strip()

    # coleta todos os filtros
    filtros = {
        'classificacao': request.args.get('classificacao', '').strip(),
        'inscricao': request.args.get('inscricao', '').strip(),
        'nota_final': request.args.get('nota_final', '').strip(),
        'ordem_de_convocacao': request.args.get('ordem_de_convocacao', '').strip(),
        # 'sim' | 'nao' | None
        'apresentou': request.args.get('apresentou'),
        'situacao_convocacao_id': request.args.get('situacao_convocacao_id', type=int),
        'matricula': request.args.get('matricula'),
        'numero_da_matricula_doe': request.args.get('numero_da_matricula_doe', '').strip(),
        'bg_matricula_doe': request.args.get('bg_matricula_doe', '').strip(),
        'portaria_convocacao': request.args.get('portaria_convocacao', '').strip(),
        'bg_portaria_convocacao': request.args.get('bg_portaria_convocacao', '').strip(),
        'doe_portaria_convocacao': request.args.get('doe_portaria_convocacao', '').strip(),
        'notificacao_pessoal': request.args.get('notificacao_pessoal'),
        'termo_desistencia': request.args.get('termo_desistencia'),
        'siged_desistencia': request.args.get('siged_desistencia', '').strip(),
    }

    query = ControleConvocacao.query

    # busca rápida por nome
    if search:
        query = query.filter(ControleConvocacao.nome.ilike(f'%{search}%'))

    # aplica filtros textuais (LIKE)
    like_map = {
        'classificacao': ControleConvocacao.classificacao,
        'inscricao': ControleConvocacao.inscricao,
        'nota_final': ControleConvocacao.nota_final,
        'ordem_de_convocacao': ControleConvocacao.ordem_de_convocacao,
        'numero_da_matricula_doe': ControleConvocacao.numero_da_matricula_doe,
        'bg_matricula_doe': ControleConvocacao.bg_matricula_doe,
        'portaria_convocacao': ControleConvocacao.portaria_convocacao,
        'bg_portaria_convocacao': ControleConvocacao.bg_portaria_convocacao,
        'doe_portaria_convocacao': ControleConvocacao.doe_portaria_convocacao,
        'siged_desistencia': ControleConvocacao.siged_desistencia,
    }
    for campo, coluna in like_map.items():
        if filtros[campo]:
            query = query.filter(coluna.ilike(f"% {filtros[campo]} %"))

    # filtros exatos / booleanos
    if filtros['situacao_convocacao_id']:
        query = query.filter(
            ControleConvocacao.situacao_convocacao_id == filtros['situacao_convocacao_id'])
    bool_map = {
        'apresentou': ControleConvocacao.apresentou,
        'matricula': ControleConvocacao.matricula,
        'notificacao_pessoal': ControleConvocacao.notificacao_pessoal,
        'termo_desistencia': ControleConvocacao.termo_desistencia,
    }
    for campo, coluna in bool_map.items():
        if filtros[campo] in ('sim', 'nao'):
            query = query.filter(coluna.is_(filtros[campo] == 'sim'))

    convocacoes_paginadas = query.order_by(
        ControleConvocacao.id.asc()).paginate(page=page, per_page=per_page)

    # dados para o gráfico
    situacoes_list = [
        c.situacao.situacao if c.situacao else 'Indefinido'
        for c in convocacoes_paginadas.items
    ]
    contagem_situacoes = dict(Counter(situacoes_list))

    csrf_token = generate_csrf()
    return render_template(
        'controle_convocacao.html',
        convocacoes=convocacoes_paginadas,
        contagem_situacoes=contagem_situacoes,
        csrf_token=csrf_token          # <- aqui
    )


@app.route('/importar-convocados', methods=['GET', 'POST'])
@login_required
def importar_convocados():
    if request.method == 'POST':
        arquivo = request.files['arquivo']
        if arquivo.filename.endswith('.xlsx'):
            filename = secure_filename(arquivo.filename)

            # Garante que a pasta 'uploads' existe
            os.makedirs('uploads', exist_ok=True)

            caminho = os.path.join('uploads', filename)
            arquivo.save(caminho)

            df = pd.read_excel(caminho)

            for _, row in df.iterrows():
                nome = NomeConvocado(
                    nome=row['nome'],
                    inscricao=row.get('inscricao', ''),
                    classificacao=row.get('classificacao', ''),
                    nota_final=row.get('nota_final', '')
                )
                database.session.add(nome)
            database.session.commit()

            flash('Nomes importados com sucesso!', 'success')
            return redirect(url_for('adicionar_convocacao'))
        else:
            flash('Formato inválido. Envie um arquivo .xlsx', 'danger')

    return render_template('importar_convocados.html')


@app.route('/gerar-qrcodes', methods=['GET', 'POST'])
@login_required
def gerar_qrcodes():
    if request.method == 'POST':
        arquivo = request.files.get('arquivo')
        if not arquivo or not arquivo.filename.endswith('.xlsx'):
            flash('Envie um arquivo .xlsx válido', 'danger')
            return redirect(request.url)

        # lê planilha em memória
        df = pd.read_excel(arquivo)

        if not {'nome_completo', 'qrcode_link'} <= set(df.columns.str.lower()):
            flash('Planilha deve conter colunas nome_completo e qrcode_link', 'danger')
            return redirect(request.url)

        # cria zip em memória
        buffer_zip = BytesIO()
        with zipfile.ZipFile(buffer_zip, 'w', zipfile.ZIP_DEFLATED) as zf:
            for _, row in df.iterrows():
                nome = str(row['nome_completo']).strip()
                link = str(row['qrcode_link']).strip()
                if not nome or not link:
                    continue

                # gera QR
                qr_img = qrcode.make(link)
                img_bytes = BytesIO()
                qr_img.save(img_bytes, format='PNG')
                img_bytes.seek(0)

                # filename seguro (sem espaços / acentos)
                fname = secure_filename(f'{nome}.png')
                zf.writestr(fname, img_bytes.read())

        buffer_zip.seek(0)
        return send_file(
            buffer_zip,
            mimetype='application/zip',
            as_attachment=True,
            download_name='qrcodes.zip'
        )

    # GET → mostra formulário
    return render_template('gerar_qrcodes.html')


@app.route('/controle-convocacao/exportar', methods=['GET'])
@login_required
def exportar_convocacoes():
    filtros = {
        'classificacao': request.args.get('classificacao', '').strip(),
        'inscricao': request.args.get('inscricao', '').strip(),
        'nota_final': request.args.get('nota_final', '').strip(),
        'ordem_de_convocacao': request.args.get('ordem_de_convocacao', '').strip(),
        'apresentou': request.args.get('apresentou'),
        'situacao_convocacao_id': request.args.get('situacao_convocacao_id', type=int),
        'matricula': request.args.get('matricula'),
        'numero_da_matricula_doe': request.args.get('numero_da_matricula_doe', '').strip(),
        'bg_matricula_doe': request.args.get('bg_matricula_doe', '').strip(),
        'portaria_convocacao': request.args.get('portaria_convocacao', '').strip(),
        'bg_portaria_convocacao': request.args.get('bg_portaria_convocacao', '').strip(),
        'doe_portaria_convocacao': request.args.get('doe_portaria_convocacao', '').strip(),
        'notificacao_pessoal': request.args.get('notificacao_pessoal'),
        'termo_desistencia': request.args.get('termo_desistencia'),
        'siged_desistencia': request.args.get('siged_desistencia', '').strip(),
    }

    query = ControleConvocacao.query

    # filtros LIKE
    like_map = {
        'classificacao': ControleConvocacao.classificacao,
        'inscricao': ControleConvocacao.inscricao,
        'nota_final': ControleConvocacao.nota_final,
        'ordem_de_convocacao': ControleConvocacao.ordem_de_convocacao,
        'numero_da_matricula_doe': ControleConvocacao.numero_da_matricula_doe,
        'bg_matricula_doe': ControleConvocacao.bg_matricula_doe,
        'portaria_convocacao': ControleConvocacao.portaria_convocacao,
        'bg_portaria_convocacao': ControleConvocacao.bg_portaria_convocacao,
        'doe_portaria_convocacao': ControleConvocacao.doe_portaria_convocacao,
        'siged_desistencia': ControleConvocacao.siged_desistencia,
    }
    for campo, coluna in like_map.items():
        if filtros[campo]:
            query = query.filter(coluna.ilike(f"%{filtros[campo]}%"))

    if filtros['situacao_convocacao_id']:
        query = query.filter(
            ControleConvocacao.situacao_convocacao_id == filtros['situacao_convocacao_id'])

    bool_map = {
        'apresentou': ControleConvocacao.apresentou,
        'matricula': ControleConvocacao.matricula,
        'notificacao_pessoal': ControleConvocacao.notificacao_pessoal,
        'termo_desistencia': ControleConvocacao.termo_desistencia,
    }
    for campo, coluna in bool_map.items():
        if filtros[campo] in ('sim', 'nao'):
            query = query.filter(coluna.is_(filtros[campo] == 'sim'))

    registros = query.order_by(ControleConvocacao.id.asc()).all()

    data = []
    for c in registros:
        data.append({
            'Classificação': c.classificacao,
            'Inscrição': c.inscricao,
            'Nome': c.nome,
            'Nota Final': c.nota_final,
            'Ordem Convocação': c.ordem_de_convocacao,
            'Apresentou': 'Sim' if c.apresentou else 'Não',
            'Situação': c.situacao.situacao if c.situacao else '-',
            'Matrícula': 'Sim' if c.matricula else 'Não',
            'Nº Mat. DOE': c.numero_da_matricula_doe,
            'BG Mat. DOE': c.bg_matricula_doe,
            'Portaria Conv.': c.portaria_convocacao,
            'BG Portaria': c.bg_portaria_convocacao,
            'DOE Portaria': c.doe_portaria_convocacao,
            'Notif. Pessoal': 'Sim' if c.notificacao_pessoal else 'Não',
            'Termo Desist.': 'Sim' if c.termo_desistencia else 'Não',
            'SIGED Desist.': c.siged_desistencia,
            'Criado em': c.data_criacao.strftime('%d/%m/%Y') if c.data_criacao else '-'
        })

    df = pd.DataFrame(data)

    # Criar buffer na memória para o arquivo Excel
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)

    response = make_response(output.read())
    response.headers["Content-Disposition"] = "attachment; filename=convocacoes_filtradas.xlsx"
    response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return response


@app.route('/atualizar-campo-convocacao', methods=['POST'])
@login_required
def atualizar_campo_convocacao():
    dados = request.get_json()

    token = request.headers.get("X-CSRFToken", "")
    try:
        validate_csrf(token)
    except Exception:
        return jsonify({'sucesso': False, 'erro': 'CSRF inválido'}), 400

    convoc_id = dados.get('id')
    campo = dados.get('campo')
    valor = dados.get('valor', '')

    # campos que podem ser editados inline
    campos_permitidos = {
        'ordem_de_convocacao': str,
        'apresentou': bool,
        'situacao_convocacao_id': int,
        'matricula': bool,
        'numero_da_matricula_doe': str,
        'bg_matricula_doe': str,
        'portaria_convocacao': str,
        'bg_portaria_convocacao': str,
        'doe_portaria_convocacao': str,
        'notificacao_pessoal': bool,
        'termo_desistencia': bool,
        'siged_desistencia': str
    }

    if campo not in campos_permitidos:
        return jsonify({'sucesso': False, 'erro': 'Campo não permitido'}), 400

    conv = ControleConvocacao.query.get_or_404(convoc_id)

    tipo = campos_permitidos[campo]

    try:
        if tipo is bool:
            valor = str(valor).strip().lower() in ('1', 'true', 'sim', 'on')
        elif tipo is int:
            valor = int(valor)
        else:
            valor = valor.strip()
    except Exception:
        return jsonify({'sucesso': False, 'erro': 'Erro ao converter valor'}), 400

    setattr(conv, campo, valor)
    database.session.commit()

    return jsonify({'sucesso': True})


@app.route('/ficha-alunos-soldados', methods=['GET', 'POST'])
@login_required
def ficha_aluno():
    form = FichaAlunosForm()
    ''''
    Pelotões:
        1° Pelotão: Rio Javari
        2° Pelotão: Rio Juruá
        3° Pelotão: Rio Japurá 
        4° Pelotão: Rio Purus
    '''
    # Preenchendo choices se necessário
    form.pelotao.choices = [('', '— Selecionar —'),
                            ('Rio Javari', 'Rio Javari'), ('Rio Juruá', 'Rio Juruá'),
                            ('Rio Japurá', 'Rio Japurá'), ('Rio Purus', 'Rio Purus')]

    form.estado_civil.choices = [('', '— Selecionar —'),
                                 ('Solteiro', 'Solteiro'), ('Casado', 'Casado'),
                                 ('Divorciado', 'Divorciado'), ('Viúvo', 'Viúvo')]

    form.estado.choices = [('', '— Selecionar —'),
                           ('AM', 'Amazonas'), ('AC', 'Acre')]

    form.categoria_cnh.choices = [('', '— Selecionar —'),
                                  ('A', 'A'), ('B', 'B'), ('AB',
                                                           'AB'), ('C', 'C'), ('D', 'D'),
                                  ('E', 'E'), ('AC', 'AC'), ('AD', 'AD'), ('AE', 'AE')]

    foto_url = None

    if form.validate_on_submit():
        foto_filename = None
        if form.foto.data:
            filename = secure_filename(form.foto.data.filename)
            foto_path = os.path.join('uploads/fotos', filename)
            form.foto.data.save(foto_path)
            foto_filename = foto_path

        novo_aluno = FichaAlunos(
            nome_completo=form.nome_completo.data or 'NÃO INFORMADO',
            nome_guerra=form.nome_guerra.data or None,
            idade_atual=form.idade_atual.data,
            cpf=form.cpf.data or None,
            rg=form.rg.data or None,
            estado_civil=form.estado_civil.data or None,
            nome_pai=form.nome_pai.data or 'NÃO INFORMADO',
            nome_mae=form.nome_mae.data or 'NÃO INFORMADO',
            pelotao=form.pelotao.data or None,
            email=form.email.data or None,
            telefone=form.telefone.data or None,
            telefone_emergencia=form.telefone_emergencia.data or None,
            rua=form.rua.data or None,
            bairro=form.bairro.data or None,
            complemento=form.complemento.data or None,
            estado=form.estado.data or None,
            formacao_academica=form.formacao_academica.data or 'NÃO INFORMADO',
            tipo_sanguineo=form.tipo_sanguineo.data or None,
            categoria_cnh=form.categoria_cnh.data or None,
            comportamento=(form.comportamento.data or 'Bom'),
            nota_comportamento=(
                form.nota_comportamento.data if form.nota_comportamento.data is not None else 5.0),
            caso_aluno_nao_resida_em_manaus=form.hospedagem_aluno_de_fora.data or None,
            foto=foto_filename or None,
            matricula=form.matricula.data or None
        )
        database.session.add(novo_aluno)
        database.session.commit()
        flash('Ficha do aluno salva com sucesso!', 'success')
        return redirect(url_for('ficha_aluno'))
    else:
        # Mostra erros explícitos (agora você vai ver por causa do padding/top/flash)
        if form.errors:
            flash('Corrija os campos destacados para salvar a ficha.', 'danger')

    return render_template('ficha_alunos.html', form=form, foto_url=foto_url, ano_atual=datetime.now().year,
                           aluno=None,          # <- importante
                           is_edicao=False      # <- flag
                           )


@app.route('/fichas')
@login_required
def listar_fichas():
    search = request.args.get('search', '').strip()

    query = FichaAlunos.query.filter(FichaAlunos.ativo == True)

    if search:
        query = query.filter(FichaAlunos.nome_completo.ilike(f"%{search}%"))

    alunos = query.order_by(FichaAlunos.nome_completo.asc()).all()

    idade_chart = Counter([a.idade_atual for a in alunos if a.idade_atual])
    cnh_chart = Counter([a.categoria_cnh for a in alunos if a.categoria_cnh])
    comportamento_raw = [a.comportamento.strip().capitalize()
                         for a in alunos if a.comportamento]
    comportamento_chart = Counter(comportamento_raw)

    return render_template(
        'fichas.html',
        alunos=alunos,
        search=search,
        idade_chart=idade_chart,
        cnh_chart=cnh_chart,
        comportamento_chart=comportamento_chart,
        ano_atual=datetime.now().year
    )


@app.route('/fichas/<int:aluno_id>')
def ficha_detalhada(aluno_id):
    aluno = FichaAlunos.query.get_or_404(aluno_id)
    return render_template('ficha_detalhada.html', aluno=aluno, ano_atual=datetime.now().year)


def calcular_comportamento(nota):
    if nota < 4:
        return "Mau"
    elif nota < 5:
        return "Insuficiente"
    elif nota < 8:
        return "Bom"
    elif nota < 9:
        return "Ótimo"
    else:
        return "Excepcional"


@app.route('/fichas/<int:aluno_id>/editar', methods=['GET', 'POST'])
@login_required
def editar_ficha(aluno_id):
    aluno = FichaAlunos.query.get_or_404(aluno_id)
    form = FichaAlunosForm(obj=aluno)

    form.pelotao.choices = [('Rio Javari', 'Rio Javari'), ('Rio Juruá', 'Rio Juruá'),
                            ('Rio Japurá', 'Rio Japurá'), ('Rio Purus', 'Rio Purus')]
    form.estado_civil.choices = [('Solteiro', 'Solteiro'), ('Casado', 'Casado'),
                                 ('Divorciado', 'Divorciado'), ('Viúvo', 'Viúvo')]
    form.estado.choices = [('AM', 'Amazonas'), ('AC', 'Acre')]
    form.categoria_cnh.choices = [
        ('A', 'A (Moto)'), ('B', 'B (Carro)'), ('AB', 'AB (Moto + Carro)'),
        ('C', 'C (Caminhão)'), ('D', 'D (Ônibus)'), ('E', 'E (Carreta)'),
        ('AC', 'AC (Moto + Caminhão)'), ('AD', 'AD (Moto + Ônibus)'), ('AE', 'AE (Moto + Carreta)')]
    form.comportamento.choices = [
        ('Excepcional', 'Excepcional'), ('Ótimo', 'Ótimo'), ('Bom', 'Bom'), ('Insuficiente', 'Insuficiente'), ('Mau', 'Mau')]

    foto_url = url_for('static', filename=aluno.foto) if aluno.foto else url_for(
        'static', filename='img/avatar-default.png')

    if form.validate_on_submit():
        # Evita erro com FileStorage
        foto_antiga = aluno.foto  # salva foto atual
        form_data = {k: v for k, v in form.data.items() if k != 'foto'}
        for key, value in form_data.items():
            setattr(aluno, key, value)

        # Processa nova imagem se foi enviada
        if form.foto.data and form.foto.data.filename:
            from werkzeug.utils import secure_filename
            import os

            upload_folder = os.path.join('static', 'uploads', 'fotos')
            os.makedirs(upload_folder, exist_ok=True)

            filename = secure_filename(form.foto.data.filename)
            foto_path = os.path.join(upload_folder, filename)
            form.foto.data.save(foto_path)
            aluno.foto = foto_path
        else:
            aluno.foto = foto_antiga  # mantém a antiga se nenhuma nova enviada

        database.session.commit()
        flash("Ficha atualizada com sucesso!", "success")
        return redirect(url_for('listar_fichas', aluno_id=aluno.id))

    return render_template('ficha_alunos.html', form=form, foto_url=foto_url, aluno=aluno, ano_atual=datetime.now().year,
                           is_edicao=True)


@app.route('/fichas/<int:aluno_id>/inativar', methods=['GET', 'POST'])
@login_required
def inativar_aluno(aluno_id):
    aluno = FichaAlunos.query.get_or_404(aluno_id)

    if aluno.inativo:
        flash('Este aluno já está marcado como inativo.', 'warning')
        return redirect(url_for('editar_ficha', aluno_id=aluno.id))

    form = InativarAlunoForm()

    if form.validate_on_submit():
        novo_inativo = AlunoInativo(
            ficha_aluno_id=aluno.id,
            motivo_saida=form.motivo_saida.data,
            data_saida=form.data_saida.data
        )
        aluno.ativo = False
        database.session.add(novo_inativo)
        database.session.commit()
        flash('Aluno marcado como inativo com sucesso.', 'success')
        return redirect(url_for('listar_fichas'))

    return render_template('inativar_aluno.html', form=form, aluno=aluno, ano_atual=datetime.now().year)


@app.route('/alunos-inativos')
@login_required
def listar_alunos_inativos():
    nome = request.args.get('nome', '')
    motivo = request.args.get('motivo', '')

    query = AlunoInativo.query.join(FichaAlunos)

    if nome:
        query = query.filter(FichaAlunos.nome_completo.ilike(f'%{nome}%'))
    if motivo:
        query = query.filter(AlunoInativo.motivo_saida == motivo)

    alunos = query.order_by(AlunoInativo.data_saida.desc()).all()

    return render_template('listar_alunos_inativos.html', alunos=alunos, ano_atual=datetime.now().year)


@app.route('/pelotao/<slug>', methods=['GET'])
@login_required
def listar_por_pelotao(slug):
    mapa_pelotoes = {
        'rio-javari': 'Rio Javari',
        'rio-jurua': 'Rio Juruá',
        'rio-japura': 'Rio Japurá',
        'rio-purus': 'Rio Purus'
    }

    nome_pelotao = mapa_pelotoes.get(slug)
    if not nome_pelotao:
        abort(404)

    termo = request.args.get('termo', '').strip()
    query = FichaAlunos.query.filter(FichaAlunos.pelotao == nome_pelotao)

    if termo:
        query = query.filter(FichaAlunos.nome_completo.ilike(f'%{termo}%'))

    alunos = query.order_by(FichaAlunos.nome_completo.asc()).all()

    # GERAÇÃO DOS DADOS PARA OS GRÁFICOS
    idade_chart = Counter([a.idade_atual for a in alunos if a.idade_atual])
    cnh_chart = Counter([a.categoria_cnh for a in alunos if a.categoria_cnh])
    comportamento_raw = [a.comportamento.strip().capitalize()
                         for a in alunos if a.comportamento]
    comportamento_chart = Counter(comportamento_raw)

    return render_template('fichas.html',
                           alunos=alunos,
                           termo_busca=termo,
                           titulo=f'Alunos do {nome_pelotao}',
                           idade_chart=idade_chart,
                           cnh_chart=cnh_chart,
                           comportamento_chart=comportamento_chart,
                           ano_atual=datetime.now().year
                           )


@app.route('/fichas/<int:aluno_id>/lts', methods=['GET', 'POST'])
@login_required
def registrar_lts(aluno_id):
    aluno = FichaAlunos.query.get_or_404(aluno_id)
    form = LtsAlunoForm()

    if form.validate_on_submit():
        nova_lts = LtsAlunos(
            ficha_aluno_id=aluno.id,
            boletim_interno=form.boletim_interno.data,
            data_inicio=form.data_inicio.data,
            data_fim=form.data_fim.data,
            usuario_id=current_user.id
        )

        database.session.add(nova_lts)
        database.session.commit()
        flash('LTS registrada com sucesso!', 'success')
        return redirect(url_for('editar_ficha', aluno_id=aluno.id))

    return render_template('registrar_lts_aluno.html', form=form, aluno=aluno, ano_atual=datetime.now().year)


@app.route('/alunos-em-lts')
@login_required
def listar_alunos_em_lts():
    hoje = datetime.utcnow().date()

    licencas_ativas = LtsAlunos.query.join(FichaAlunos).filter(
        LtsAlunos.data_inicio <= hoje,
        LtsAlunos.data_fim >= hoje
    ).order_by(LtsAlunos.data_inicio.asc()).all()

    return render_template('alunos_em_lts.html', licencas=licencas_ativas, ano_atual=datetime.now().year)


@app.route('/fichas/<int:aluno_id>/restricao', methods=['GET', 'POST'])
@login_required
def registrar_restricao(aluno_id):
    aluno = FichaAlunos.query.get_or_404(aluno_id)
    form = RestricaoAlunoForm()

    if form.validate_on_submit():
        existe_igual = RestricaoAluno.query.filter_by(
            ficha_aluno_id=aluno.id,
            descricao=form.descricao.data,
            data_inicio=form.data_inicio.data,
            data_fim=form.data_fim.data
        ).first()

        if existe_igual:
            flash('Restrição já registrada para esse período.', 'warning')
            return redirect(url_for('editar_ficha', aluno_id=aluno.id))

        nova_restricao = RestricaoAluno(
            ficha_aluno_id=aluno.id,
            descricao=form.descricao.data,
            data_inicio=form.data_inicio.data,
            data_fim=form.data_fim.data,
            usuario_id=current_user.id
        )
        database.session.add(nova_restricao)
        database.session.commit()
        flash('Restrição registrada com sucesso!', 'success')
        return redirect(url_for('editar_ficha', aluno_id=aluno.id))

    return render_template('registrar_restricao.html', form=form, aluno=aluno, ano_atual=datetime.now().year)


@app.route('/restricoes-ativas')
@login_required
def restricoes_ativas():
    hoje = date.today()

    restricoes = RestricaoAluno.query.join(FichaAlunos).filter(
        RestricaoAluno.data_inicio <= hoje,
        RestricaoAluno.data_fim >= hoje
    ).order_by(RestricaoAluno.data_inicio.asc()).all()

    return render_template('restricoes_ativas.html', restricoes=restricoes, ano_atual=datetime.now().year)


@app.route('/restricoes-ativas/excel')
@login_required
def exportar_restricoes_excel():
    hoje = date.today()
    restricoes = RestricaoAluno.query.join(FichaAlunos).filter(
        RestricaoAluno.data_inicio <= hoje,
        RestricaoAluno.data_fim >= hoje
    ).all()

    dados = [{
        'Nome do Aluno': r.ficha_aluno.nome_completo,
        'Pelotão': r.ficha_aluno.pelotao,
        'Motivo': r.descricao,
        'Data Início': r.data_inicio.strftime('%d/%m/%Y'),
        'Data Fim': r.data_fim.strftime('%d/%m/%Y'),
        'Registrado por': r.usuario.nome,
        'Data Registro': r.data_criacao.strftime('%d/%m/%Y %H:%M'),
    } for r in restricoes]

    df = pd.DataFrame(dados)

    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Restrições Ativas')

    output.seek(0)
    return send_file(output, download_name='restricoes_ativas.xlsx', as_attachment=True)


@app.route('/restricoes-ativas/print')
@login_required
def imprimir_restricoes_ativas():
    hoje = date.today()
    restricoes = RestricaoAluno.query.join(FichaAlunos).filter(
        RestricaoAluno.data_inicio <= hoje,
        RestricaoAluno.data_fim >= hoje
    ).all()
    return render_template('restricoes_print.html', restricoes=restricoes)


@app.route('/fichas/<int:aluno_id>/imprimir')
@login_required
def imprimir_ficha_aluno(aluno_id):
    aluno = FichaAlunos.query.get_or_404(aluno_id)
    return render_template('ficha_detalhada_print.html', aluno=aluno)


@app.route('/dashboard-obms')
@login_required
def dashboard_obms():
    dados = dados_para_mapa()
    return render_template('dashboard_obms.html', dados=dados)


@app.route('/atualizacao-cadastral', methods=['GET', 'POST'])
def atualizacao_cadastral():
    form = IdentificacaoForm()

    if form.validate_on_submit():
        print("VALIDOU ✅")
        cpf_raw = form.cpf.data
        email_digitado = form.email.data.strip().lower()

        # mantém seu formato padrão com máscara
        cpf_formatado = formatar_cpf(cpf_raw)

        # 👉 NOVO: procurar em Militar OU FichaAlunos
        pessoa = buscar_pessoa_por_cpf(cpf_formatado)
        if not pessoa:
            flash("⚠️ CPF não encontrado no sistema (Militar/Aluno). Verifique e tente novamente ou contate a DRH.", "danger")
            return render_template("atualizacao/identificacao.html", form=form)

        session['email_atualizacao'] = email_digitado

        # Já existe User com esse CPF?
        user = User.query.filter_by(cpf=cpf_formatado).first()
        if user:
            flash(
                "⚠️ Já existe uma conta vinculada a este CPF. Faça login para continuar.", "warning")
            return redirect(url_for('login_atualizacao'))

        # 👉 Guarda no fluxo de validação de identidade
        session['cpf_em_validacao'] = cpf_formatado
        session['pessoa_tipo'] = pessoa['tipo']           # 'militar' | 'aluno'
        session['pessoa_id'] = pessoa['obj'].id           # id correspondente

        return redirect(url_for('confirmar_matricula'))

    return render_template("atualizacao/identificacao.html", form=form)


@app.route('/confirmar-matricula', methods=['GET', 'POST'])
def confirmar_matricula():
    cpf = session.get('cpf_em_validacao')
    pessoa_tipo = session.get('pessoa_tipo')  # 'militar' | 'aluno'
    pessoa_id = session.get('pessoa_id')

    if not cpf or not pessoa_tipo or not pessoa_id:
        flash("Sessão expirada ou inválida. Refaça a identificação.", "warning")
        return redirect(url_for('atualizacao_cadastral'))

    form = MatriculaConfirmForm()

    # Carrega a pessoa do tipo correto
    if pessoa_tipo == 'militar':
        pessoa = Militar.query.get(pessoa_id)
        if not pessoa:
            flash("Registro militar não encontrado para o CPF em validação.", "danger")
            _limpa_sessao_validacao()
            return redirect(url_for('atualizacao_cadastral'))
        nome_pessoa = getattr(pessoa, 'nome_completo',
                              getattr(pessoa, 'nome', ''))
        matricula_oficial = pessoa.matricula

    else:  # 'aluno'
        pessoa = FichaAlunos.query.get(pessoa_id)
        if not pessoa:
            flash("Registro de aluno não encontrado para o CPF em validação.", "danger")
            _limpa_sessao_validacao()
            return redirect(url_for('atualizacao_cadastral'))
        nome_pessoa = getattr(pessoa, 'nome_completo', '')
        matricula_oficial = pessoa.matricula

    if form.validate_on_submit():
        matricula_informada = (form.matricula_completa.data or "").strip()

        if normaliza_matricula(matricula_informada) != normaliza_matricula(matricula_oficial or ""):
            flash("❌ Matrícula não confere com nossos registros para este CPF.", "danger")
            return render_template('atualizacao/confirmar_matricula.html',
                                   form=form, cpf=cpf, militar_nome=nome_pessoa)

        session['matricula_validada'] = True
        # mantém pessoa_tipo/pessoa_id já na sessão

        flash("✅ Identidade confirmada com sucesso. Crie sua senha.", "success")
        return redirect(url_for('criar_senha', cpf=cpf))

    return render_template('atualizacao/confirmar_matricula.html',
                           form=form, cpf=cpf, militar_nome=nome_pessoa,
                           matricula=matricula_oficial)


@app.route('/criar-senha/<cpf>', methods=['GET', 'POST'])
def criar_senha(cpf):
    cpf = formatar_cpf(cpf)

    # 🔒 checagem de fluxo e correspondência
    if not session.get('matricula_validada') or session.get('cpf_em_validacao') != cpf:
        flash("Valide sua identidade antes de criar a senha.", "warning")
        return redirect(url_for('atualizacao_cadastral'))

    pessoa_tipo = session.get('pessoa_tipo')
    pessoa_id = session.get('pessoa_id')
    if not pessoa_tipo or not pessoa_id:
        flash("Sessão expirada ou inválida. Refaça a identificação.", "warning")
        return redirect(url_for('atualizacao_cadastral'))

    form = CriarSenhaForm()

    # Já tem conta?
    if User.query.filter_by(cpf=cpf).first():
        flash("⚠️ Este CPF já possui uma conta ativa. Tente fazer login.", "warning")
        _limpa_sessao_validacao()
        return redirect(url_for('login'))

    # Carrega a pessoa do tipo correto
    if pessoa_tipo == 'militar':
        pessoa = Militar.query.get(pessoa_id)
        if not pessoa:
            flash("❌ Militar não encontrado para este CPF.", "danger")
            _limpa_sessao_validacao()
            return redirect(url_for('atualizacao_cadastral'))
        nome_user = getattr(pessoa, 'nome_completo',
                            getattr(pessoa, 'nome', ''))
    else:
        pessoa = FichaAlunos.query.get(pessoa_id)
        if not pessoa:
            flash("❌ Aluno não encontrado para este CPF.", "danger")
            _limpa_sessao_validacao()
            return redirect(url_for('atualizacao_cadastral'))
        nome_user = getattr(pessoa, 'nome_completo', '')

    if form.validate_on_submit():
        senha_hash = bcrypt.generate_password_hash(
            form.senha.data).decode('utf-8')

        novo_usuario = User(
            nome=nome_user or '',
            cpf=cpf,  # com máscara
            email=session.get('email_atualizacao'),
            senha=senha_hash,
            funcao_user_id=12  # USUÁRIO COMUM
        )

        if pessoa_tipo == 'militar':
            # Mantém suas atribuições extras para militar
            obm_id_1, obm_id_2 = _obms_ativas_do_militar(pessoa.id)
            if hasattr(novo_usuario, 'obm_id_1'):
                novo_usuario.obm_id_1 = obm_id_1
            if hasattr(novo_usuario, 'obm_id_2'):
                novo_usuario.obm_id_2 = obm_id_2
            if hasattr(novo_usuario, 'localidade_id'):
                novo_usuario.localidade_id = getattr(
                    pessoa, 'localidade_id', None)
        else:
            # 👉 Requisito: todos os alunos são da OBM 26
            if hasattr(novo_usuario, 'obm_id_1'):
                novo_usuario.obm_id_1 = 26
            if hasattr(novo_usuario, 'obm_id_2'):
                novo_usuario.obm_id_2 = None
            if hasattr(novo_usuario, 'localidade_id'):
                novo_usuario.localidade_id = None  # ajuste se necessário

        database.session.add(novo_usuario)
        database.session.flush()
        database.session.commit()

        if pessoa_tipo == 'militar':
            pessoa.usuario_id = novo_usuario.id

        _limpa_sessao_validacao()
        flash("✅ Conta criada com sucesso! Agora você pode fazer login.", "success")
        return redirect(url_for('login_atualizacao'))

    return render_template('atualizacao/criar_senha.html', form=form, cpf=cpf)


def _limpa_sessao_validacao():
    for k in ['matricula_validada', 'cpf_em_validacao', 'militar_id_validado', 'email_atualizacao']:
        session.pop(k, None)


@app.route('/formulario-atualizacao-cadastral', methods=['GET', 'POST'])
@login_required
def formulario_atualizacao_cadastral():
    if current_user.funcao_user_id != 12:
        flash("⚠️ Acesso restrito à atualização cadastral.", "danger")
        return redirect(url_for('home'))

    # Busca o militar vinculado ao CPF do usuário logado
    militar = Militar.query.filter_by(cpf=current_user.cpf).first()

    if not militar:
        flash("❌ Dados do militar não encontrados.", "danger")
        return redirect(url_for('home'))

    form = AtualizacaoCadastralForm(obj=militar)

    if form.validate_on_submit():
        militar.celular = form.celular.data
        militar.email = form.email.data
        militar.endereco = form.endereco.data
        militar.complemento = form.complemento.data
        militar.cidade = form.cidade.data
        militar.estado = form.estado.data
        militar.grau_instrucao = form.grau_instrucao.data

        database.session.commit()

        vinculo = SegundoVinculo.query.filter_by(militar_id=militar.id).first()
        if not vinculo:
            vinculo = SegundoVinculo(militar_id=militar.id)

        vinculo.possui_vinculo = form.possui_vinculo.data
        vinculo.quantidade_vinculos = form.quantidade_vinculos.data
        vinculo.descricao_vinculo = form.descricao_vinculo.data
        vinculo.horario_inicio = form.horario_inicio.data
        vinculo.horario_fim = form.horario_fim.data

        database.session.add(vinculo)
        database.session.commit()
        flash("✅ Dados atualizados com sucesso!", "success")
        return redirect(url_for('ficha_atualizada'))

    return render_template('atualizacao/formulario_cadastro.html', form=form)


@app.route("/login-militar", methods=['GET', 'POST'])
def login_atualizacao():
    if current_user.is_authenticated:
        militar = get_militar_por_user(current_user)
        session['militar_id'] = militar.id
        if not militar:
            flash("Não foi possível localizar seus dados de militar.", "danger")
            return redirect(url_for("home"))
        return redirect(url_for('home_atualizacao'))

    form_login = FormLogin()

    if form_login.validate_on_submit() and 'botao_submit_login' in request.form:
        cpf_formatado = form_login.cpf.data.strip()
        usuario = User.query.filter_by(cpf=cpf_formatado).first()

        if usuario and bcrypt.check_password_hash(usuario.senha, form_login.senha.data):
            if usuario.funcao_user_id == 12:
                login_user(usuario, remember=form_login.lembrar_dados.data)
                militar = get_militar_por_user(usuario)
                if not militar:
                    flash("Não foi possível localizar seus dados de militar.", "danger")
                    return redirect(url_for("home"))
                flash('Login realizado com sucesso.', 'success')
                return redirect(url_for('home_atualizacao'))
            else:
                flash(
                    'Este usuário não tem permissão para acessar a atualização cadastral.', 'danger')
        else:
            flash('CPF ou senha incorretos.', 'danger')

    return render_template("atualizacao/login_atualizacao.html", form_login=form_login)

# --------------------------------------------------------------- PARTE DOS CADETES -----------------------------------------------------------------#


@app.route("/login-cadete", methods=['GET', 'POST'])
def login_cadete():
    if current_user.is_authenticated:
        if is_cadete(current_user):
            return redirect(url_for("home_cadete"))
        flash("Seu perfil não é de cadete.", "warning")
        return redirect(url_for("home"))

    form = FormLogin()
    if form.validate_on_submit() and 'botao_submit_login' in request.form:
        cpf_formatado = form.cpf.data.strip()
        usuario = User.query.filter_by(cpf=cpf_formatado).first()

        if usuario and bcrypt.check_password_hash(usuario.senha, form.senha.data):
            login_user(usuario, remember=form.lembrar_dados.data)
            if not is_cadete(usuario):
                logout_user()
                flash("Acesso exclusivo para cadetes.", "danger")
                return redirect(url_for("login_cadete"))
            flash("Login de cadete realizado.", "success")
            return redirect(url_for("home_cadete"))
        else:
            flash("CPF ou senha incorretos.", "danger")

    return render_template("cadete/login_cadete.html", form_login=form)


@app.route("/cadete", methods=['GET'])
@login_required
def home_cadete():
    if not is_cadete(current_user):
        flash("Acesso restrito aos cadetes.", "danger")
        return redirect(url_for("home"))

    # Fechamento automático: se zerou pendentes, mostra tela de conclusão.
    pendentes = (TarefaAtualizacaoCadete.query
                 .filter_by(cadete_user_id=current_user.id, status="PENDENTE")
                 .count())

    if pendentes == 0 and request.args.get("lista") != "1":
        return render_template("cadete/finalizado.html")

    # lista para a página (paginável, se quiser)
    tarefas = (TarefaAtualizacaoCadete.query
               .filter_by(cadete_user_id=current_user.id)
               .order_by(TarefaAtualizacaoCadete.status.desc(), TarefaAtualizacaoCadete.id.asc())
               .all())

    return render_template("cadete/home.html", tarefas=tarefas, pendentes=pendentes)


@app.route("/cadete/exibir-militar/<int:militar_id>", methods=['GET', 'POST'])
@login_required
def cadete_exibir_militar(militar_id):
    if not is_cadete(current_user):
        flash("Acesso restrito aos cadetes.", "danger")
        return redirect(url_for("home"))

    # Confere a tarefa
    tarefa = TarefaAtualizacaoCadete.query.filter_by(
        cadete_user_id=current_user.id, militar_id=militar_id
    ).first()

    if not tarefa:
        flash("Este militar não está atribuído a você.", "danger")
        return redirect(url_for("home_cadete"))

    militar = Militar.query.get_or_404(militar_id)

    # Regra: alunos não podem aparecer para o cadete
    if militar.situacao_id == 9:
        flash("Este registro não é permitido para atualização pelo cadete.", "warning")
        return redirect(url_for("home_cadete"))

    # (Opcional) trava simples
    if tarefa.status == "PENDENTE":
        tarefa.status = "EM_EDICAO"
        tarefa.locked_by_user_id = current_user.id
        tarefa.locked_at = datetime.utcnow()
        database.session.commit()

    obm_funcao_tipo_1 = MilitarObmFuncao.query.filter_by(militar_id=militar_id, tipo=1) \
        .filter(MilitarObmFuncao.data_fim == None).first()

    obm_funcao_tipo_2 = MilitarObmFuncao.query.filter_by(militar_id=militar_id, tipo=2) \
        .filter(MilitarObmFuncao.data_fim == None).first()

    # --- AQUI você pode reaproveitar quase tudo da sua rota exibir_militar ---
    # Reaproveite o setup do form e choices:
    form_militar = FormMilitar(obj=militar)
    if militar.completa_25_inclusao:
        form_militar.completa_25_inclusao.data = militar.completa_25_inclusao.strftime(
            '%d/%m/%Y')
    if militar.completa_30_inclusao:
        form_militar.completa_30_inclusao.data = militar.completa_30_inclusao.strftime(
            '%d/%m/%Y')
    if militar.completa_25_anos_sv:
        form_militar.completa_25_anos_sv.data = militar.completa_25_anos_sv.strftime(
            '%d/%m/%Y')
    if militar.completa_30_anos_sv:
        form_militar.completa_30_anos_sv.data = militar.completa_30_anos_sv.strftime(
            '%d/%m/%Y')

    form_militar.funcao_gratificada_id.choices = [
        (funcao_gratificada.id, funcao_gratificada.gratificacao) for funcao_gratificada in FuncaoGratificada.query.all()
    ]
    form_militar.posto_grad_id.choices = [
        (posto.id, posto.sigla) for posto in PostoGrad.query.all()
    ]
    form_militar.quadro_id.choices = [
        (quadro.id, quadro.quadro) for quadro in Quadro.query.all()
    ]
    form_militar.localidade_id.choices = [
        (localidade_id.id, localidade_id.sigla) for localidade_id in Localidade.query.all()
    ]

    form_militar.obm_ids_1.choices = ([('', '-- Selecione uma opção --')] +
                                      [(obm.id, obm.sigla) for obm in Obm.query.all()])

    form_militar.funcao_ids_1.choices = ([('', '-- Selecione uma opção --')] +
                                         [(funcao.id, funcao.ocupacao) for funcao in Funcao.query.all()])

    form_militar.obm_ids_2.choices = ([('', '-- Selecione uma opção --')] +
                                      [(obm.id, obm.sigla) for obm in Obm.query.all()])

    form_militar.funcao_ids_2.choices = ([('', '-- Selecione uma opção --')] +
                                         [(funcao.id, funcao.ocupacao) for funcao in Funcao.query.all()])

    form_militar.situacao_id.choices = [
        (situacao.id, situacao.condicao) for situacao in Situacao.query.all()
    ]
    form_militar.estado_civil.choices = [
        (estado.id, estado.estado) for estado in EstadoCivil.query.all()
    ]
    form_militar.especialidade_id.choices = [
        (especialidade.id, especialidade.ocupacao) for especialidade in Especialidade.query.all()
    ]
    form_militar.destino_id.choices = [
        (destino.id, destino.local) for destino in Destino.query.all()
    ]
    form_militar.agregacoes_id.choices = [
        (agregacoes.id, agregacoes.tipo) for agregacoes in Agregacoes.query.all()
    ]
    form_militar.punicao_id.choices = [
        (punicao.id, punicao.sancao) for punicao in Punicao.query.all()
    ]
    form_militar.comportamento_id.choices = ([('', '-- Selecione uma opção --')] +
                                             [(comportamento.id, comportamento.conduta) for comportamento in
                                              Comportamento.query.all()
                                              ])

    if obm_funcao_tipo_1:
        form_militar.obm_ids_1.data = obm_funcao_tipo_1.obm_id
        form_militar.funcao_ids_1.data = obm_funcao_tipo_1.funcao_id

    if obm_funcao_tipo_2:
        form_militar.obm_ids_2.data = obm_funcao_tipo_2.obm_id
        form_militar.funcao_ids_2.data = obm_funcao_tipo_2.funcao_id

    form_militar.sexo.data = militar.sexo if militar.sexo else None
    form_militar.raca.data = militar.raca if militar.raca else None

    hoje = date.today()

    dn = militar.data_nascimento
    # se vier datetime, vira date; se None, fica None
    if isinstance(dn, datetime):
        dn = dn.date()

    if dn:
        idade = hoje.year - dn.year - \
            ((hoje.month, hoje.day) < (dn.month, dn.day))
    else:
        idade = None  # ou 0, se preferir

    form_militar.idade_atual.data = idade
    campos_bg = [
        'transferencia', 'situacao_militar', 'cfsd', 'cfc', 'cfs', 'cas',
        'choa', 'cfo', 'cbo', 'cao', 'csbm', 'soldado_tres',
        'soldado_dois', 'soldado_um', 'cabo', 'terceiro_sgt',
        'segundo_sgt', 'primeiro_sgt', 'subtenente',
        'publicidade_segundo_tenente', 'publicidade_primeiro_tenente',
        'pub_cap', 'pub_maj', 'pub_tc', 'pub_cel', 'pub_alteracao'
    ]

    # Recupera as publicações do BG e define campos para publicação do militar
    publicacoes_bg = PublicacaoBg.query.filter_by(militar_id=militar.id).all()

    # Inicializa todos os campos de publicações BG com valor vazio
    for campo in campos_bg:
        if hasattr(form_militar, campo):
            getattr(form_militar, campo).data = ""

    # Itera sobre as publicações do BG e verifica os campos
    for pub in publicacoes_bg:
        if pub.tipo_bg in campos_bg:
            if hasattr(form_militar, pub.tipo_bg):
                # Se o boletim_geral existir, utiliza o valor; caso contrário, mantém como vazio
                getattr(form_militar, pub.tipo_bg).data = pub.boletim_geral or ""

    def parse_date(d):
        """Aceita date, string 'YYYY-MM-DD' ou 'DD/MM/YYYY'. Retorna date ou None."""
        if not d:
            return None
        if isinstance(d, datetime):
            return d.date()
        if hasattr(d, 'strftime'):  # já é date
            return d
        s = str(d).strip()
        for fmt in ('%Y-%m-%d', '%d/%m/%Y'):
            try:
                return datetime.strptime(s, fmt).date()
            except ValueError:
                pass
        return None  # não converteu

    def safe_bg_id(militar_id):
        """Tenta pegar BG de situacao_militar; se não achar, retorna None (não bloqueia criação)."""
        bg = PublicacaoBg.query.filter_by(
            militar_id=militar_id, tipo_bg='situacao_militar').first()

    if form_militar.validate_on_submit():
        form_militar.process(request.form)

        militar.nome_completo = form_militar.nome_completo.data
        militar.nome_guerra = form_militar.nome_guerra.data
        militar.cpf = form_militar.cpf.data
        militar.rg = form_militar.rg.data
        militar.nome_pai = form_militar.nome_pai.data
        militar.nome_mae = form_militar.nome_mae.data
        militar.matricula = form_militar.matricula.data
        militar.pis_pasep = form_militar.pis_pasep.data
        militar.num_titulo_eleitor = form_militar.num_titulo_eleitor.data
        militar.digito_titulo_eleitor = form_militar.digito_titulo_eleitor.data
        militar.zona = form_militar.zona.data
        militar.secao = form_militar.secao.data
        militar.posto_grad_id = form_militar.posto_grad_id.data
        militar.quadro_id = form_militar.quadro_id.data
        militar.localidade_id = form_militar.localidade_id.data
        militar.antiguidade = form_militar.antiguidade.data
        militar.sexo = form_militar.sexo.data
        militar.raca = form_militar.raca.data
        militar.data_nascimento = form_militar.data_nascimento.data
        militar.inclusao = form_militar.inclusao.data

        militar.completa_25_inclusao = datetime.strptime(
            str(form_militar.completa_25_inclusao.data), '%d/%m/%Y'
        ).date() if form_militar.completa_25_inclusao.data else None

        militar.completa_30_inclusao = datetime.strptime(
            str(form_militar.completa_30_inclusao.data), '%d/%m/%Y'
        ).date() if form_militar.completa_30_inclusao.data else None

        militar.completa_25_anos_sv = datetime.strptime(
            str(form_militar.completa_25_anos_sv.data), '%d/%m/%Y'
        ).date() if form_militar.completa_25_anos_sv.data else None

        militar.completa_30_anos_sv = datetime.strptime(
            str(form_militar.completa_30_anos_sv.data), '%d/%m/%Y'
        ).date() if form_militar.completa_30_anos_sv.data else None

        militar.inicio_periodo = datetime.strptime(
            str(form_militar.inicio_periodo.data), '%Y-%m-%d'
        ).date() if form_militar.inicio_periodo.data else None

        militar.fim_periodo = datetime.strptime(
            str(form_militar.fim_periodo.data), '%Y-%m-%d'
        ).date() if form_militar.fim_periodo.data else None

        militar.punicao_id = form_militar.punicao_id.data
        militar.comportamento_id = form_militar.comportamento_id.data
        militar.efetivo_servico = form_militar.efetivo_servico.data
        militar.anos = form_militar.anos.data
        militar.meses = form_militar.meses.data
        militar.dias = form_militar.dias.data
        militar.total_dias = form_militar.total_dias.data
        militar.idade_reserva_grad = 0
        militar.estado_civil = form_militar.estado_civil.data
        militar.especialidade_id = form_militar.especialidade_id.data
        militar.pronto = form_militar.pronto.data
        militar.situacao_id = form_militar.situacao_id.data
        militar.agregacoes_id = form_militar.agregacoes_id.data
        militar.destino_id = form_militar.destino_id.data
        militar.inicio_periodo = parse_date(form_militar.inicio_periodo.data)
        militar.fim_periodo = parse_date(form_militar.fim_periodo.data)
        militar.ltip_afastamento_cargo_eletivo = form_militar.ltip_afastamento_cargo_eletivo.data
        militar.periodo_ltip = form_militar.periodo_ltip.data
        militar.total_ltip = form_militar.total_ltip.data
        militar.completa_25_anos_ltip = form_militar.completa_25_anos_ltip.data
        militar.completa_30_anos_ltip = form_militar.completa_30_anos_ltip.data
        militar.cursos = form_militar.cursos.data
        militar.grau_instrucao = form_militar.grau_instrucao.data
        militar.graduacao = form_militar.graduacao.data
        militar.pos_graduacao = form_militar.pos_graduacao.data
        militar.mestrado = form_militar.mestrado.data
        militar.doutorado = form_militar.doutorado.data
        militar.cfsd = form_militar.cfsd.data
        militar.cfc = form_militar.cfc.data
        militar.cfs = form_militar.cfs.data
        militar.cas = form_militar.cas.data
        militar.choa = form_militar.choa.data
        militar.cfo = form_militar.cfo.data
        militar.cbo = form_militar.cbo.data
        militar.cao = form_militar.cao.data
        militar.csbm = form_militar.csbm.data
        militar.cursos_civis = form_militar.cursos_civis.data
        militar.endereco = form_militar.endereco.data
        militar.complemento = form_militar.complemento.data
        militar.cidade = form_militar.cidade.data
        militar.estado = form_militar.estado.data
        militar.cep = form_militar.cep.data
        militar.celular = form_militar.celular.data
        militar.email = form_militar.email.data
        militar.inclusao_bg = form_militar.inclusao_bg.data
        militar.soldado_tres = form_militar.soldado_tres.data
        militar.soldado_dois = form_militar.soldado_dois.data
        militar.soldado_um = form_militar.soldado_um.data
        militar.cabo = form_militar.cabo.data
        militar.terceiro_sgt = form_militar.terceiro_sgt.data
        militar.segundo_sgt = form_militar.segundo_sgt.data
        militar.primeiro_sgt = form_militar.primeiro_sgt.data
        militar.subtenente = form_militar.subtenente.data
        militar.segundo_tenente = form_militar.segundo_tenente.data
        militar.primeiro_tenente = form_militar.primeiro_tenente.data
        militar.cap = form_militar.cap.data
        militar.maj = form_militar.maj.data
        militar.tc = form_militar.tc.data
        militar.cel = form_militar.cel.data
        militar.funcao_gratificada_id = form_militar.funcao_gratificada_id.data
        militar.alteracao_nome_guerra = form_militar.alteracao_nome_guerra.data

        # Verifica se OBM 1 e Função 1 foram modificadas ou removidas (selecionada opção vazia)
        if form_militar.obm_ids_1.data and form_militar.funcao_ids_1.data and form_militar.obm_ids_1.data != '' and form_militar.funcao_ids_1.data != '':
            # Somente faz a alteração se houver diferença entre os dados do banco e os do formulário
            obm_id_1 = int(form_militar.obm_ids_1.data)
            funcao_id_1 = int(form_militar.funcao_ids_1.data)

            if not obm_funcao_tipo_1 or (
                    obm_id_1 != obm_funcao_tipo_1.obm_id or funcao_id_1 != obm_funcao_tipo_1.funcao_id):
                # Feche o registro antigo se houver um
                if obm_funcao_tipo_1:
                    obm_funcao_tipo_1.data_fim = datetime.now()

                nova_relacao_1 = MilitarObmFuncao(
                    militar_id=militar_id,
                    obm_id=obm_id_1,
                    funcao_id=funcao_id_1,
                    tipo=1,
                    data_criacao=datetime.now(),
                    data_fim=None
                )
                database.session.add(nova_relacao_1)
        else:
            # Se a opção 'Selecione uma opção' foi escolhida, finalize a relação atual
            if obm_funcao_tipo_1:
                obm_funcao_tipo_1.data_fim = datetime.now()

        # Verifica se OBM 2 e Função 2 foram modificadas ou removidas
        if form_militar.obm_ids_2.data and form_militar.funcao_ids_2.data and form_militar.obm_ids_2.data != '' and form_militar.funcao_ids_2.data != '':
            # Somente faz a alteração se houver diferença entre os dados do banco e os do formulário
            obm_id_2 = int(form_militar.obm_ids_2.data)
            funcao_id_2 = int(form_militar.funcao_ids_2.data)

            if not obm_funcao_tipo_2 or (
                    obm_id_2 != obm_funcao_tipo_2.obm_id or funcao_id_2 != obm_funcao_tipo_2.funcao_id):
                # Feche o registro antigo se houver um
                if obm_funcao_tipo_2:
                    obm_funcao_tipo_2.data_fim = datetime.now()

                nova_relacao_2 = MilitarObmFuncao(
                    militar_id=militar_id,
                    obm_id=obm_id_2,
                    funcao_id=funcao_id_2,
                    tipo=2,
                    data_criacao=datetime.now(),
                    data_fim=None
                )
                database.session.add(nova_relacao_2)
        else:
            # Se a opção 'Selecione uma opção' foi escolhida, finalize a relação atual
            if obm_funcao_tipo_2:
                obm_funcao_tipo_2.data_fim = datetime.now()

        for campo_bg in campos_bg:
            if hasattr(form_militar, campo_bg):
                boletim_geral_data = getattr(form_militar, campo_bg).data
                if boletim_geral_data:
                    publicacao_existente = PublicacaoBg.query.filter_by(
                        militar_id=militar.id, tipo_bg=campo_bg).first()
                    if publicacao_existente:
                        if publicacao_existente.boletim_geral != boletim_geral_data:
                            nova_publicacao = PublicacaoBg(
                                militar_id=militar.id,
                                tipo_bg=campo_bg,
                                boletim_geral=boletim_geral_data
                            )
                            database.session.add(nova_publicacao)
                    else:
                        nova_publicacao = PublicacaoBg(
                            militar_id=militar.id,
                            tipo_bg=campo_bg,
                            boletim_geral=boletim_geral_data
                        )
                        database.session.add(nova_publicacao)

        # Verifica se a situação selecionada é "AGREGADO"
        situacao_selecionada = Situacao.query.get(
            form_militar.situacao_id.data)
        if situacao_selecionada and situacao_selecionada.condicao == 'AGREGADO':
            bg_id = safe_bg_id(militar.id)  # pode ser None
            militar_agregado = MilitaresAgregados.query.filter_by(
                militar_id=militar.id).first()
            if not militar_agregado:
                militar_agregado = MilitaresAgregados(militar_id=militar.id)
                database.session.add(militar_agregado)

            militar_agregado.posto_grad_id = form_militar.posto_grad_id.data
            militar_agregado.quadro_id = form_militar.quadro_id.data
            militar_agregado.destino_id = form_militar.destino_id.data
            militar_agregado.situacao_id = situacao_selecionada.id
            militar_agregado.inicio_periodo = parse_date(
                form_militar.inicio_periodo.data)
            militar_agregado.fim_periodo_agregacao = parse_date(
                form_militar.fim_periodo.data)
            militar_agregado.publicacao_bg_id = bg_id
            militar_agregado.atualizar_status()

        # À DISPOSIÇÃO
        if situacao_selecionada and situacao_selecionada.condicao == 'À DISPOSIÇÃO':
            bg_id = safe_bg_id(militar.id)
            militar_a_disposicao = MilitaresADisposicao.query.filter_by(
                militar_id=militar.id).first()
            if not militar_a_disposicao:
                militar_a_disposicao = MilitaresADisposicao(
                    militar_id=militar.id)
                database.session.add(militar_a_disposicao)

            militar_a_disposicao.posto_grad_id = form_militar.posto_grad_id.data
            militar_a_disposicao.quadro_id = form_militar.quadro_id.data
            militar_a_disposicao.destino_id = form_militar.destino_id.data
            militar_a_disposicao.situacao_id = situacao_selecionada.id
            militar_a_disposicao.inicio_periodo = parse_date(
                form_militar.inicio_periodo.data)
            militar_a_disposicao.fim_periodo_disposicao = parse_date(
                form_militar.fim_periodo.data)
            militar_a_disposicao.publicacao_bg_id = bg_id
            militar_a_disposicao.atualizar_status()

        # LICENÇA ESPECIAL
        if situacao_selecionada and situacao_selecionada.condicao == 'LICENÇA ESPECIAL':
            bg_id = safe_bg_id(militar.id)
            militar_le = LicencaEspecial.query.filter_by(
                militar_id=militar.id).first()
            if not militar_le:
                militar_le = LicencaEspecial(militar_id=militar.id)
                database.session.add(militar_le)

            militar_le.posto_grad_id = form_militar.posto_grad_id.data
            militar_le.quadro_id = form_militar.quadro_id.data
            militar_le.destino_id = form_militar.destino_id.data
            militar_le.situacao_id = situacao_selecionada.id
            militar_le.inicio_periodo_le = parse_date(
                form_militar.inicio_periodo.data)
            militar_le.fim_periodo_le = parse_date(
                form_militar.fim_periodo.data)
            militar_le.publicacao_bg_id = bg_id
            militar_le.atualizar_status()

        # LTS
        if situacao_selecionada and situacao_selecionada.condicao == 'LTS':
            bg_id = safe_bg_id(militar.id)
            militar_lts = LicencaParaTratamentoDeSaude.query.filter_by(
                militar_id=militar.id).first()
            if not militar_lts:
                militar_lts = LicencaParaTratamentoDeSaude(
                    militar_id=militar.id)
                database.session.add(militar_lts)

            militar_lts.posto_grad_id = form_militar.posto_grad_id.data
            militar_lts.quadro_id = form_militar.quadro_id.data
            militar_lts.destino_id = form_militar.destino_id.data
            militar_lts.situacao_id = situacao_selecionada.id
            militar_lts.inicio_periodo_lts = parse_date(
                form_militar.inicio_periodo.data)
            militar_lts.fim_periodo_lts = parse_date(
                form_militar.fim_periodo.data)
            militar_lts.publicacao_bg_id = bg_id
            militar_lts.atualizar_status()
        acao = request.form.get("acao", "salvar")  # padrão: salvar rascunho
        try:
            database.session.commit()
            if acao == "concluir":
                tarefa.status = "CONCLUIDO"
                tarefa.atualizado_em = datetime.utcnow()
                tarefa.locked_by_user_id = None
                tarefa.locked_at = None
            else:
                # só salva rascunho
                tarefa.status = "EM_EDICAO"
                tarefa.locked_by_user_id = current_user.id
                tarefa.locked_at = datetime.utcnow()

            database.session.commit()

            if acao == "concluir" and cadete_restantes(current_user) == 0:
                flash("Você concluiu todas as suas atualizações!", "success")
                return redirect(url_for("home_cadete"))

            flash("Dados salvos.", "success" if acao ==
                  "salvar" else "Atualizado com sucesso.")
            return redirect(url_for("home_cadete"))
        except Exception as e:
            database.session.rollback()
            flash(f"Erro ao atualizar. {e}", "danger")

    # Carregue também documentos, se quiser, mas apenas leitura
    documentos_militar = DocumentoMilitar.query.filter_by(
        militar_id=militar.id
    ).order_by(DocumentoMilitar.criado_em.desc()).all()

    return render_template("cadete/exibir_militar.html",
                           form_militar=form_militar,
                           militar=militar,
                           documentos_militar=documentos_militar,
                           tarefa=tarefa)


@app.route("/cadete/reabrir/<int:militar_id>", methods=["POST"])
@login_required
def cadete_reabrir(militar_id):
    if not is_cadete(current_user):
        flash("Acesso restrito aos cadetes.", "danger")
        return redirect(url_for("home"))

    tarefa = (TarefaAtualizacaoCadete.query
              .filter_by(cadete_user_id=current_user.id, militar_id=militar_id)
              .first_or_404())

    if tarefa.status != "CONCLUIDO":
        flash("Esta tarefa não está concluída.", "warning")
        return redirect(url_for("home_cadete"))

    tarefa.status = "EM_EDICAO"
    tarefa.locked_by_user_id = current_user.id
    tarefa.locked_at = datetime.utcnow()
    database.session.commit()

    return redirect(url_for("cadete_exibir_militar", militar_id=militar_id))


@app.route('/ficha-atualizada')
@login_required
def ficha_atualizada():
    militar = Militar.query.filter_by(cpf=current_user.cpf).first_or_404()
    segundo_vinculo = SegundoVinculo.query.filter_by(
        militar_id=militar.id).first()

    return render_template(
        'atualizacao/ficha_atualizada.html',
        militar=militar,
        segundo_vinculo=segundo_vinculo
    )


def _norm_cpf(cpf: str) -> str:
    return re.sub(r'\D', '', cpf or '')


@app.route("/admin/distribuir-atualizacao", methods=["GET", "POST"])
@login_required
@checar_ocupacao('DIRETOR', 'CHEFE', 'SUPER USER', 'DIRETOR DRH', 'CHEFE DRH', 'DRH')
def distribuir_atualizacao():
    form = DistribuirAtualizacaoForm()

    # 1) Cadetes na tabela MILITAR (por posto_grad_id = 7)
    cadetes_mil = (Militar.query
                   .filter(Militar.posto_grad_id == 7)
                   .all())

    # Conjunto de CPFs normalizados desses cadetes
    cadete_cpfs_norm = {_norm_cpf(m.cpf) for m in cadetes_mil if m.cpf}

    # 2) Users casados por CPF (normalizado)
    #     Aqui faço robusto: pego todos e filtro em Python pelo CPF normalizado.)
    all_users = User.query.all()
    cadetes_users_por_cpf = [
        u for u in all_users if _norm_cpf(u.cpf) in cadete_cpfs_norm]

    # 3) Users casados por FK usuario_id (para quem já está linkado)
    cadetes_users_por_fk = (database.session.query(User)
                            .join(Militar, Militar.usuario_id == User.id)
                            .filter(Militar.posto_grad_id == 7)
                            .all())

    # 4) União dos dois métodos (distinct por id)
    cadetes_users_dict = {u.id: u for u in cadetes_users_por_cpf}
    cadetes_users_dict.update({u.id: u for u in cadetes_users_por_fk})
    cadetes_users = list(cadetes_users_dict.values())

    alvos = (Militar.query
             .filter(Militar.situacao_id != 9)
             .filter((Militar.posto_grad_id != 15) | (Militar.posto_grad_id.is_(None)))
             .order_by(Militar.id.asc())
             .all())

    tot_cadetes_mil = len(cadetes_mil)
    tot_cadetes_cpf = len(cadetes_users_por_cpf)
    tot_cadetes_fk = len(cadetes_users_por_fk)
    tot_cadetes = len(cadetes_users)
    tot_alvos = len(alvos)
    tot_tarefas = database.session.query(
        func.count(TarefaAtualizacaoCadete.id)).scalar() or 0
    tarefas_pend = (TarefaAtualizacaoCadete.query
                    .filter(TarefaAtualizacaoCadete.status.in_(["PENDENTE", "EM_EDICAO"]))
                    .count())
    estim_por_cadete = math.ceil(tot_alvos / tot_cadetes) if tot_cadetes else 0

    if request.method == "GET":
        return render_template(
            "admin/distribuir_atualizacao.html",
            form=form,
            # métricas exibidas
            tot_cadetes=tot_cadetes,
            tot_alvos=tot_alvos,
            tot_tarefas=tot_tarefas,
            tarefas_pend=tarefas_pend,
            estim_por_cadete=estim_por_cadete,
            tot_cadetes_mil=tot_cadetes_mil,
            tot_cadetes_cpf=tot_cadetes_cpf,
            tot_cadetes_fk=tot_cadetes_fk,
        )

    if form.validate_on_submit():
        modo_limpeza = form.limpeza.data
        try:
            # limpeza...
            if modo_limpeza == "pendentes":
                (TarefaAtualizacaoCadete.query
                 .filter(TarefaAtualizacaoCadete.status.in_(["PENDENTE", "EM_EDICAO"]))
                 .delete(synchronize_session=False))
                database.session.commit()
            elif modo_limpeza == "todas":
                TarefaAtualizacaoCadete.query.delete()
                database.session.commit()

            # --- Recarregar bases -----------------------------
            cadetes_mil = Militar.query.filter(
                Militar.posto_grad_id == 7).all()
            cadete_cpfs_norm = {_norm_cpf(m.cpf) for m in cadetes_mil if m.cpf}

            # Users que são cadetes (por CPF normalizado)
            all_users = User.query.options(load_only(User.id, User.cpf)).all()
            cadetes_users = [u for u in all_users if _norm_cpf(
                u.cpf) in cadete_cpfs_norm]
            if not cadetes_users:
                flash("Nenhum cadete encontrado via CPF.", "warning")
                return redirect(url_for("distribuir_atualizacao"))

            # Alvos: não alunos e não civis
            alvos = (Militar.query
                     .filter(Militar.situacao_id != 9)
                     .filter((Militar.posto_grad_id != 15) | (Militar.posto_grad_id.is_(None)))
                     .all())
            if not alvos:
                flash("Não há militares elegíveis.", "warning")
                return redirect(url_for("distribuir_atualizacao"))

            # >>> EMBARALHA TUDO PARA FICAR ALEATÓRIO A CADA POST <<<
            cad_user_ids = [u.id for u in cadetes_users]
            shuffle(cad_user_ids)     # ordem aleatória dos cadetes
            shuffle(alvos)            # ordem aleatória dos alvos

            # --- Mapeia cadete_user_id -> cadete_militar_id por CPF normalizado -------
            cpf_user_norm = func.regexp_replace(
                cast(User.cpf, String), r'[^0-9]', '', 'g')
            cpf_mil_norm = func.regexp_replace(
                cast(Militar.cpf, String), r'[^0-9]', '', 'g')

            pares_user_mil = (database.session.query(User.id, Militar.id)
                              .join(Militar, cpf_user_norm == cpf_mil_norm)
                              .filter(Militar.posto_grad_id == 7,
                                      User.id.in_(cad_user_ids))
                              .all())
            # {cadete_user_id: cadete_militar_id}
            cadete_map = dict(pares_user_mil)

            # --- Evita recriar pares já existentes ------------------------------------
            existentes = set(database.session.query(
                TarefaAtualizacaoCadete.cadete_user_id,
                TarefaAtualizacaoCadete.militar_id
            ).all())

            # --- Round-robin em cima das listas embaralhadas --------------------------
            to_insert = []
            ccount = len(cad_user_ids)

            for i, alvo in enumerate(alvos):
                cad_uid = cad_user_ids[i % ccount]
                cad_mil_id = cadete_map.get(cad_uid)
                if not cad_mil_id:
                    continue
                par = (cad_uid, alvo.id)
                if par in existentes:
                    continue
                to_insert.append({
                    "cadete_user_id": cad_uid,
                    "cadete_militar_id": cad_mil_id,
                    "militar_id": alvo.id,
                    "status": "PENDENTE",
                })

            if to_insert:
                database.session.bulk_insert_mappings(
                    TarefaAtualizacaoCadete, to_insert)

            database.session.commit()
            flash(
                f"Distribuição concluída. Criadas: {len(to_insert)}.", "success")
            return redirect(url_for("distribuir_atualizacao"))

        except Exception as e:
            database.session.rollback()
            flash(f"Erro na distribuição: {e}", "danger")
            return redirect(url_for("distribuir_atualizacao"))

    # fallback de validação
    return render_template(
        "admin/distribuir_atualizacao.html",
        form=form,
        tot_cadetes=tot_cadetes,
        tot_alvos=tot_alvos,
        tot_tarefas=tot_tarefas,
        tarefas_pend=tarefas_pend,
        estim_por_cadete=estim_por_cadete,
        tot_cadetes_mil=tot_cadetes_mil,
        tot_cadetes_cpf=tot_cadetes_cpf,
        tot_cadetes_fk=tot_cadetes_fk,
    )


@app.route("/relatorio/cadetes-militares")
def relatorio_cadetes_militares():
    fmt = (request.args.get("format") or "html").lower()
    status = (request.args.get("status") or "TODOS").upper()

    # SUBQUERY: escolhe 1 OBM por militar: prioridade (data_fim IS NULL) e mais recente por data_criacao
    obm_ranked = (
        database.session.query(
            MilitarObmFuncao.militar_id.label("m_id"),
            MilitarObmFuncao.obm_id.label("obm_id"),
            func.row_number().over(
                partition_by=MilitarObmFuncao.militar_id,
                order_by=(
                    case((MilitarObmFuncao.data_fim.is_(None), 0), else_=1),
                    MilitarObmFuncao.data_criacao.desc(),
                )
            ).label("rn")
        )
    ).subquery()

    q = (
        database.session.query(
            User.id.label("cadete_id"),
            User.nome.label("cadete"),
            Militar.id.label("militar_id"),
            Militar.nome_completo.label("militar"),
            PostoGrad.sigla.label("posto_grad"),
            func.coalesce(Obm.sigla, "").label("obm"),
        )
        .join(TarefaAtualizacaoCadete, TarefaAtualizacaoCadete.cadete_user_id == User.id)
        .join(Militar, Militar.id == TarefaAtualizacaoCadete.militar_id)
        .outerjoin(PostoGrad, PostoGrad.id == Militar.posto_grad_id)
        # junta a OBM "escolhida" (rn = 1) sem perder quem não tem OBM
        .outerjoin(obm_ranked, and_(obm_ranked.c.m_id == Militar.id, obm_ranked.c.rn == 1))
        .outerjoin(Obm, Obm.id == obm_ranked.c.obm_id)
    )

    if status != "TODOS":
        q = q.filter(TarefaAtualizacaoCadete.status == status)

    q = q.order_by(User.nome.asc(), Militar.nome_completo.asc())

    rows = [{
        "cadete_id":  r.cadete_id,
        "cadete":     r.cadete,
        "militar_id": r.militar_id,
        "militar":    r.militar,
        "posto_grad": r.posto_grad or "",
        "obm":        r.obm or "",
    } for r in q.all()]

    # ===== formatos =====
    if fmt == "json":
        return jsonify(rows)

    if fmt == "xlsx":
        wb = Workbook()
        ws = wb.active
        ws.title = "Militares por Cadete"
        headers = ["Cadete", "Militar", "Posto/Grad", "OBM", "Militar ID"]
        ws.append(headers)
        for r in rows:
            ws.append([r["cadete"], r["militar"],
                      r["posto_grad"], r["obm"], r["militar_id"]])
        for col_idx in range(1, len(headers) + 1):
            col = get_column_letter(col_idx)
            maxlen = max((len(str(c.value)) if c.value else 0)
                         for c in ws[col])
            ws.column_dimensions[col].width = min(max(12, maxlen + 2), 50)
        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="cadetes_militares.xlsx"
        )

    # HTML agrupado por cadete
    grupos = defaultdict(list)
    for r in rows:
        grupos[(r["cadete_id"], r["cadete"])].append(r)
    grupos_sorted = sorted(grupos.items(), key=lambda k: k[0][1].upper())

    return render_template(
        "relatorio_cadetes_militares.html",
        grupos=grupos_sorted,
        status=status
    )


@app.route('/fichas/<int:aluno_id>/recompensa', methods=['GET', 'POST'])
@login_required
def registrar_recompensa(aluno_id):
    aluno = FichaAlunos.query.get_or_404(aluno_id)
    form = RecompensaAlunoForm()

    if form.validate_on_submit():
        nova = RecompensaAluno(
            ficha_aluno_id=aluno.id,
            natureza=form.natureza.data,
            autoridade=form.autoridade.data,
            boletim=form.boletim.data,
            discriminacao=form.discriminacao.data,
            usuario_id=current_user.id
        )
        database.session.add(nova)
        database.session.commit()
        flash('Recompensa registrada com sucesso!', 'success')
        return redirect(url_for('editar_ficha', aluno_id=aluno.id))

    return render_template('registrar_recompensa.html', form=form, aluno=aluno)


@app.route('/fichas/<int:aluno_id>/sancao', methods=['GET', 'POST'])
@login_required
def registrar_sancao(aluno_id):
    aluno = FichaAlunos.query.get_or_404(aluno_id)
    form = SancaoAlunoForm()

    if form.validate_on_submit():
        nova = SancaoAluno(
            ficha_aluno_id=aluno.id,
            natureza=form.natureza.data,
            numero_dias=form.numero_dias.data,
            boletim=form.boletim.data,
            data_inicio=form.data_inicio.data,
            data_fim=form.data_fim.data,
            discriminacao=form.discriminacao.data,
            usuario_id=current_user.id
        )
        database.session.add(nova)
        database.session.commit()
        flash('Sanção registrada com sucesso!', 'success')
        return redirect(url_for('editar_ficha', aluno_id=aluno.id))

    return render_template('registrar_sancao.html', form=form, aluno=aluno)


@app.route('/quiz')
def quiz():
    return render_template('quiz.html')


@app.route('/gerar-le', methods=['GET', 'POST'])
@login_required
def gerar_le():
    if request.method == 'POST':
        nome = request.form['nome']

        dados = {
            'nota_bg': request.form['nota_bg'],
            'data_do_requerimento': formatar_data_extenso(request.form['data_do_requerimento']),
            # <-- corrigido aqui
            'POSTO/GRADUACAO': request.form['posto_grad'],
            'QUADRO': request.form['quadro'],
            'NOME do MILITAR': nome,
            'OBM do militar': request.form['obm'],
            'tipo_licenca_especial': request.form['tipo_licenca_especial'],
            'data_inicio_licenca_especial a data_fim_licenca_especial': request.form['periodo_licenca'],
            'data_inicio_pedido': formatar_data_sem_zero(request.form['data_inicio_pedido']),
            'data_de_apresentacao': formatar_data_sem_zero(request.form['data_apresentacao']),
            'numero_siged': request.form['numero_siged'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nota_bg', 'POSTO/GRADUACAO', 'QUADRO',
                   'NOME do MILITAR', 'data_inicio_pedido']
        ITALICO = ['numero_siged']

        doc = Document('src/template/nota_tecnica.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                    if chave in ITALICO:
                        novo_run.italic = True
                        if chave == 'numero_siged':
                            novo_run.text = f"({valor})"
                            novo_run.font.size = Pt(10)
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name='nota_le_gerada.docx')

    return render_template('gerar_le.html')


@app.route('/indeferimento-le', methods=['GET', 'POST'])
@login_required
def indeferimento_le():
    if request.method == 'POST':
        nome = request.form['nome']

        dados = {
            'nota_bg': request.form['nota_bg'],
            'POSTO/GRADUACAO': request.form['posto_grad'],
            'QUADRO': request.form['quadro'],
            'NOME do MILITAR': nome,
            'OBM do militar': request.form['obm'],
            'tipo_licenca_especial': request.form['tipo_licenca_especial'],
            'data_inicio_licenca_especial a data_fim_licenca_especial': request.form['periodo_licenca'],
            'numero_siged': request.form['numero_siged'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nota_bg', 'POSTO/GRADUACAO', 'QUADRO',
                   'NOME do MILITAR', 'data_inicio_licenca_especial a data_fim_licenca_especial']
        ITALICO = ['numero_siged']
        doc = Document('src/template/indeferimento_le.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                    if chave in ITALICO:
                        novo_run.italic = True
                        if chave == 'numero_siged':
                            novo_run.text = f"({valor})"
                            novo_run.font.size = Pt(10)
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name='indeferimento_le_gerada.docx')

    return render_template('indeferimento_le.html')


@app.route('/gerar-lp', methods=['GET', 'POST'])
@login_required
def gerar_lp():
    if request.method == 'POST':
        dados = {
            'nota_bg': request.form['nota_bg'],
            'matricula_certidao': request.form['matricula_certidao'],
            'cartorio': request.form['cartorio'],
            'nome_filho': request.form['nome_filho'],
            'cidade_natal': request.form['cidade_natal'],
            'pai': request.form['pai'],
            'mae': request.form['mae'],
            'data_certidao': formatar_data_extenso(request.form['data_certidao']),
            'oficial_responsavel': request.form['oficial_responsavel'],
            'data_inicio_lp': formatar_data_extenso(request.form['data_inicio_lp']),
            'data_apresentacao': formatar_data_extenso(request.form['data_apresentacao']),
            'numero_siged': request.form['numero_siged'],
            'posto_graduacao': request.form['posto_grad'],
            'quadro': request.form['quadro'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }
        NEGRITO = ['nota_bg', 'pai', 'posto_graduacao', 'quadro',
                   'data_inicio_lp', 'data_apresentacao', 'matricula_certidao']

        ITALICO = ['numero_siged']

        doc = Document('src/template/certidao_lp.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                    if chave in ITALICO:
                        novo_run.italic = True
                        if chave == 'numero_siged':
                            novo_run.text = f"({valor})"
                            novo_run.font.size = Pt(10)
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name='certidao_lp_gerada.docx')

    return render_template('gerar_lp.html')


@app.route('/gerar-certidao-casamento', methods=['GET', 'POST'])
@login_required
def gerar_certidao_casamento():
    if request.method == 'POST':

        dados = {
            'nota_bg': request.form['nota_bg'],
            'matricula_certidao': request.form['matricula_certidao'],
            'cartorio': request.form['cartorio'],
            'esposo': request.form['esposo'],
            'posto_grad': request.form['posto_grad'],
            'quadro': request.form['quadro'],
            'esposa': request.form['esposa'],
            'posto_grad_esposa': request.form['posto_grad_esposa'],
            'quadro_esposa': request.form['quadro_esposa'],
            'data_casamento': formatar_data_extenso(request.form['data_casamento']),
            'regime_casamento': request.form['regime_casamento'],
            'escrevente': request.form['escrevente'],
            'data_registro': formatar_data_extenso(request.form['data_registro']),
            'data_inicio_licenca': formatar_data_sem_zero(request.form['data_inicio_licenca']),
            'data_apresentacao': formatar_data_sem_zero(request.form['data_apresentacao']),
            'numero_siged': request.form['numero_siged'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nota_bg', 'matricula_certidao', 'cartorio',
                   'esposo', 'esposa', 'data_casamento', 'regime_casamento', 'data_inicio_licenca', 'data_apresentacao']

        ITALICO = ['numero_siged']

        doc = Document('src/template/certidao_casamento.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                    if chave in ITALICO:
                        novo_run.italic = True
                        if chave == 'numero_siged':
                            novo_run.text = f"({valor})"
                            novo_run.font.size = Pt(10)
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name='certidao_casamento_gerada.docx')

    return render_template('gerar_certidao_casamento.html')


@app.route('/gerar=certidao-obito', methods=['GET', 'POST'])
@login_required
def gerar_certidao_obito():
    if request.method == 'POST':
        dados = {
            'nota_bg': request.form['nota_bg'],
            'numero_certidao': request.form['numero_certidao'],
            'cartorio': request.form['cartorio'],
            'cidade_estado': request.form['cidade_estado'],
            'nome_falecido': request.form['nome_falecido'],
            'cidade_falecido': request.form['cidade_falecido'],
            'data_falecimento': formatar_data_extenso(request.form['data_falecimento']),
            'posto_grad': request.form['posto_grad'],
            'quadro': request.form['quadro'],
            'nome_militar': request.form['nome_militar'],
            'escrevente': request.form['escrevente'],
            'data_inicio_licenca': formatar_data_sem_zero(request.form['data_inicio_licenca']),
            'data_apresentacao': formatar_data_sem_zero(request.form['data_apresentacao']),
            'numero_siged': request.form['numero_siged'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nota_bg', 'data_falecimento', 'posto_grad', 'quadro',
                   'nome_militar', 'data_inicio_licenca', 'data_apresentacao']

        ITALICO = ['numero_siged']

        doc = Document('src/template/certidao_obito.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                    if chave in ITALICO:
                        novo_run.italic = True
                        if chave == 'numero_siged':
                            novo_run.text = f"({valor})"
                            novo_run.font.size = Pt(10)
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name='certidao_obito.docx')

    return render_template('gerar_certidao_obito.html')


@app.route('/gerar-certidao-tempo-servico', methods=['GET', 'POST'])
@login_required
def certidao_tempo_servico():
    if request.method == 'POST':
        nome = request.form['nome_completo'].replace(" ", "_")
        dados = {
            'nome_completo': request.form['nome_completo'],
            'posto_grad': request.form['posto_grad'],
            'cpf': request.form['cpf'],
            'dia_ingresso': request.form['dia_ingresso'],
            'mes_ingresso': request.form['mes_ingresso'],
            'ano_ingresso': request.form['ano_ingresso'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nome_completo', 'posto_grad', 'cpf',
                   'dia_ingresso', 'mes_ingresso', 'ano_ingresso', 'data_atual']

        doc = Document('src/template/declaracao_tempo_de_servico.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name=f'declaracao_tempo_de_servico{nome}.docx')

    return render_template('gerar_certidao_tempo_servico.html')


@app.route('/gerar-certidao-exerc-atp', methods=['GET', 'POST'])
@login_required
def certidao_exercicio_atv_atipica():
    if request.method == 'POST':
        nome = request.form['nome_completo'].replace(" ", "_")
        dados = {
            'nome_completo': request.form['nome_completo'],
            'posto_grad': request.form['posto_grad'],
            'cpf': request.form['cpf'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nome_completo', 'posto_grad', 'cpf']

        ITALICO = ['data_atual']

        doc = Document('src/template/exercicio_atividade_atipica.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True

                    if chave in ITALICO:
                        novo_run.italic = True
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name=f'declaracao_exercicio_atividade_{nome}.docx')

    return render_template('gerar_exc_atv.html')


@app.route('/gerar-declaracao', methods=['GET', 'POST'])
@login_required
def declaracao():
    if request.method == 'POST':
        nome = request.form['nome_militar'].replace(" ", "_")
        dados = {
            'nota_declaracao': request.form['nota_declaracao'],
            'nome_militar': request.form['nome_militar'],
            'orgao': request.form['orgao'],
            'posto_graduacao': request.form['posto_graduacao'],
            'quadro': request.form['quadro'],
            'cpf': request.form['cpf'],
            'rg_cbmam': request.form['rg_cbmam'],
            'matricula': request.form['matricula'],
            'especialidade': request.form['especialidade'],
            'data_concurso': formatar_data_sem_zero(request.form['data_concurso']),
            'numero_bg': request.form['numero_bg'],
            'data_bg': formatar_data_sem_zero(request.form['data_bg']),
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nome_militar', 'posto_graduacao',
                   'quadro', 'cpf', 'nota_declaracao']

        doc = Document('src/template/declaracao1.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name=f'declaracao_{nome}.docx')

    return render_template('declaracao.html')


@app.route('/nota-elogio', methods=['GET', 'POST'])
@login_required
def nota_elogio():
    if request.method == 'POST':
        nome = request.form['nome_militar']

        dados = {
            'nota_bg': request.form['nota_bg'],
            # <-- corrigido aqui
            'posto_graduacao': request.form['posto_graduacao'],
            'quadro': request.form['quadro'],
            'nome_militar': nome,
            'atestador': request.form['atestador'],
            'data_doacao': formatar_data_sem_zero(request.form['data_doacao']),
            'numero_siged': request.form['numero_siged'],
            'numero_coren': request.form['numero_coren'],
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
        }

        NEGRITO = ['nota_bg', 'posto_graduacao', 'quadro',
                   'nome_militar', 'data_doacao']
        ITALICO = ['numero_siged']

        doc = Document('src/template/doacao_sangue.docx')

        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove todos os runs do parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Regex para encontrar todos os campos do tipo {chave}
            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True
                    if chave in ITALICO:
                        novo_run.italic = True
                        if chave == 'numero_siged':
                            novo_run.text = f"({valor})"
                            novo_run.font.size = Pt(10)
                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name=f'elogio_{nome}.docx')

    return render_template('elogio_doacao.html')


@app.route('/portaria-gratificacao', methods=['GET', 'POST'])
@login_required
def portaria_gratificacao():
    if request.method == 'POST':
        # Nome para gerar o arquivo final
        nome_servidor = request.form['nome_servidor'].replace(" ", "_")

        # DADOS PARA PREENCHER O DOCUMENTO
        dados = {
            'data_atual': formatar_data_extenso(datetime.today().strftime('%Y-%m-%d')),
            'numero_processo': request.form['numero_processo'],
            'nome_juiz': request.form['nome_juiz'],
            'porcentagem': request.form['porcentagem'],
            'porcentagem_por_extenso': request.form['porcentagem_por_extenso'],
            'nome_servidor': request.form['nome_servidor'],
            'matricula': request.form['matricula'],
            'data_a_contar': formatar_data_sem_zero(request.form['data_a_contar']),
            'memo': request.form['memo']
        }

        # Campos que ficarão em negrito (opcional, pode editar)
        NEGRITO = [
            'nome_servidor', 'matricula', 'porcentagem',
            'numero_processo', 'data_atual'
        ]

        doc = Document('src/template/portaria_gc.docx')

        # PERCORRER TODOS OS PARÁGRAFOS
        for p in doc.paragraphs:
            texto = p.text
            if not any(f"{{{k}}}" in texto for k in dados):
                continue

            # Remove runs para reconstruir o parágrafo
            for run in p.runs:
                p._element.remove(run._element)

            partes = re.split(r'(\{.*?\})', texto)

            for parte in partes:
                if re.match(r'\{.*?\}', parte):
                    chave = parte.strip('{}')
                    valor = dados.get(chave, parte)

                    novo_run = p.add_run(str(valor))
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

                    if chave in NEGRITO:
                        novo_run.bold = True

                else:
                    novo_run = p.add_run(parte)
                    novo_run.font.name = 'Times New Roman'
                    novo_run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman')
                    novo_run.font.size = Pt(12)

        # EXPORTA ARQUIVO FINAL
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(
            output,
            as_attachment=True,
            download_name=f'portaria_gratificacao_{nome_servidor}.docx'
        )

    return render_template('portaria_gratificacao.html')
