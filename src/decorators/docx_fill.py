from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

def _apply_table_style(table, doc):
    for name in ("Table Grid", "Table Normal", "Normal Table", "TableNormal"):
        try:
            table.style = name
            return
        except KeyError:
            pass



def set_table_font(table, font_name="Times New Roman", font_size=11):
    """
    Força fonte e tamanho em TODAS as células da tabela.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

                    # garante Times New Roman mesmo no Word (XML-level)
                    r = run._element
                    rPr = r.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn("w:ascii"), font_name)
                    rFonts.set(qn("w:hAnsi"), font_name)


def _insert_table_after_paragraph(paragraph, rows, cols):
    """
    Insere uma tabela logo após o parágrafo (posição correta no DOCX).
    """
    tbl = OxmlElement('w:tbl')
    paragraph._p.addnext(tbl)

    # cria tabela via API e injeta o elemento xml no lugar certo
    doc = paragraph.part.document
    table = doc.add_table(rows=rows, cols=cols)
    new_tbl = table._tbl

    # substitui o tbl vazio pelo tbl real
    tbl.getparent().replace(tbl, new_tbl)
    return table

def docx_fill_template_dependentes(template_path: str, mapping: dict, dependentes: list[dict]) -> bytes:
    doc = Document(template_path)

    # 1) replace de textos simples (parágrafos)
    #    (se você já tem uma função docx_fill_template, pode usar ela ANTES e reabrir)
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                p.text = p.text.replace(k, str(v))

    # 2) procura âncora
    anchor = None
    for p in doc.paragraphs:
        if "{tabela_dependentes}" in p.text:
            anchor = p
            break
    if not anchor:
        raise ValueError("Não achei a âncora {tabela_dependentes} no template DOCX.")

    # limpa a âncora
    anchor.text = ""

    # 3) cria tabela no lugar certo
    # header + N linhas
    table = _insert_table_after_paragraph(anchor, rows=1, cols=4)
    _apply_table_style(table, doc)

    hdr = table.rows[0].cells
    hdr[0].text = "Nº"
    hdr[1].text = "Nome"
    hdr[2].text = "Grau de parentesco"
    hdr[3].text = "Idade"

    for i, d in enumerate(dependentes, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = (d.get("nome") or d.get("nome_dependente") or "").strip()
        row[2].text = (d.get("grau") or d.get("grau_parentesco") or "").strip()
        row[3].text = str((d.get("idade") or d.get("idade_dependente") or "")).strip()
    
    set_table_font(table, font_name="Times New Roman", font_size=11)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
