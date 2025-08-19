# helpers_docx.py (versão com python-docx, sem regex no XML)
import io
from typing import Dict
from docx import Document

def _replace_in_runs_of_paragraph(paragraph, mapping: Dict[str, str]):
    # Concatena o texto de todos os runs do parágrafo
    full = "".join(run.text for run in paragraph.runs)
    if not full:
        return

    # Aplica todos os replaces {chave} -> valor
    for k, v in mapping.items():
        full = full.replace("{%s}" % k, str(v))

    # Zera os runs e escreve tudo em um único run
    if paragraph.runs:
        # mantém o estilo do primeiro run
        first = paragraph.runs[0]
        first.text = full
        # apaga os demais runs
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(full)

def _walk_table(table, mapping: Dict[str, str]):
    for row in table.rows:
        for cell in row.cells:
            # parágrafos dentro da célula
            for p in cell.paragraphs:
                _replace_in_runs_of_paragraph(p, mapping)
            # tabelas aninhadas (se houver)
            for t in cell.tables:
                _walk_table(t, mapping)

def render_docx_from_template(template_path: str, mapping: Dict[str, str]) -> io.BytesIO:
    doc = Document(template_path)

    # Corpo: parágrafos
    for p in doc.paragraphs:
        _replace_in_runs_of_paragraph(p, mapping)

    # Corpo: tabelas
    for t in doc.tables:
        _walk_table(t, mapping)

    # Cabeçalhos e rodapés de todas as seções
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                _replace_in_runs_of_paragraph(p, mapping)
            for t in section.header.tables:
                _walk_table(t, mapping)
        if section.footer:
            for p in section.footer.paragraphs:
                _replace_in_runs_of_paragraph(p, mapping)
            for t in section.footer.tables:
                _walk_table(t, mapping)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
